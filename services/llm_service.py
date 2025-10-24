"""
Enhanced LLM Service with Robust JSON Parsing

This service provides intelligent document analysis and content generation
with bulletproof JSON parsing and fallback mechanisms.
"""

import json
import re
import logging
from typing import Dict, List, Any, Optional
from services.providers.llm_providers import OllamaProvider
from services.providers.model_manager import ModelManager
from docx import Document

logger = logging.getLogger(__name__)


def capture_detailed_formatting(paragraph):
    """
    Captures formatting for each run in detail, including word-level information.
    
    Args:
        paragraph: Paragraph to analyze
        
    Returns:
        List of formatting info for each run with detailed properties
    """
    run_formats = []
    
    for idx, run in enumerate(paragraph.runs):
        if not run.text:
            continue
            
        format_info = {
            'run_index': idx,
            'text': run.text,
            'length': len(run.text),
            'font_name': run.font.name,
            'font_size': run.font.size,
            'bold': run.font.bold,
            'italic': run.font.italic,
            'underline': run.font.underline,
            'color': run.font.color.rgb if run.font.color.rgb else None,
            'highlight': run.font.highlight_color,
            'strikethrough': getattr(run.font, 'strike', None)
        }
        run_formats.append(format_info)
    
    return run_formats


def apply_proportional_formatting_safe(paragraph, new_text, original_formats):
    """
    Applies formatting proportionally while reusing existing runs (SAFE for tables/lists).
    
    Strategy: Reuses cleared runs where possible, only adds new runs if needed.
    
    Args:
        paragraph: Paragraph to modify (with cleared runs)
        new_text: New text to apply
        original_formats: List of format info from capture_detailed_formatting
    """
    if not original_formats:
        logger.warning("No original formats, using plain text")
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)
        return
    
    existing_runs = list(paragraph.runs)
    
    # Calculate total original length
    total_original_length = sum(fmt['length'] for fmt in original_formats)
    new_text_length = len(new_text)
    
    # Check if formatting is uniform
    is_uniform = True
    if len(original_formats) > 1:
        first_fmt = original_formats[0]
        
        def normalize_bool(val):
            return bool(val) if val is not None else False
        
        for idx, fmt in enumerate(original_formats[1:], 1):
            bold_same = normalize_bool(fmt['bold']) == normalize_bool(first_fmt['bold'])
            italic_same = normalize_bool(fmt['italic']) == normalize_bool(first_fmt['italic'])
            underline_same = normalize_bool(fmt['underline']) == normalize_bool(first_fmt['underline'])
            color_same = fmt['color'] == first_fmt['color']
            highlight_same = fmt['highlight'] == first_fmt['highlight']
            
            if not (bold_same and italic_same and underline_same and color_same and highlight_same):
                is_uniform = False
                break
    
    if is_uniform:
        # Uniform formatting - reuse first run
        logger.info("Applying uniform formatting (reusing first run)")
        if existing_runs:
            run = existing_runs[0]
            run.text = new_text
        else:
            run = paragraph.add_run(new_text)
        
        fmt = original_formats[0]
        if fmt['font_name']:
            run.font.name = fmt['font_name']
        if fmt['font_size']:
            run.font.size = fmt['font_size']
        if fmt['bold'] is not None:
            run.font.bold = fmt['bold']
        if fmt['italic'] is not None:
            run.font.italic = fmt['italic']
        if fmt['underline'] is not None:
            run.font.underline = fmt['underline']
        if fmt['color']:
            run.font.color.rgb = fmt['color']
        if fmt['highlight']:
            run.font.highlight_color = fmt['highlight']
    else:
        # Mixed formatting - distribute proportionally, reuse runs
        logger.info(f"Applying mixed formatting across {len(original_formats)} runs (reusing existing)")
        
        position = 0
        for idx, fmt in enumerate(original_formats):
            # Calculate segment
            if idx == len(original_formats) - 1:
                segment_text = new_text[position:]
            else:
                proportion = fmt['length'] / total_original_length
                segment_length = int(new_text_length * proportion)
                
                # Try to split at word boundary
                if segment_length > 0 and position + segment_length < new_text_length:
                    space_pos = new_text.find(' ', position + segment_length)
                    if space_pos != -1 and space_pos - position < segment_length + 20:
                        segment_length = space_pos - position + 1
                
                segment_text = new_text[position:position + segment_length]
                position += segment_length
            
            if not segment_text:
                continue
            
            # Reuse existing run or create new
            if idx < len(existing_runs):
                run = existing_runs[idx]
                run.text = segment_text
            else:
                run = paragraph.add_run(segment_text)
            
            # Apply formatting
            if fmt['font_name']:
                run.font.name = fmt['font_name']
            if fmt['font_size']:
                run.font.size = fmt['font_size']
            if fmt['bold'] is not None:
                run.font.bold = fmt['bold']
            if fmt['italic'] is not None:
                run.font.italic = fmt['italic']
            if fmt['underline'] is not None:
                run.font.underline = fmt['underline']
            if fmt['color']:
                run.font.color.rgb = fmt['color']
            if fmt['highlight']:
                run.font.highlight_color = fmt['highlight']
            
            logger.info(f"  Run {idx+1}: '{segment_text[:30]}...' with bold={fmt['bold']}, "
                       f"italic={fmt['italic']}, color={fmt['color']}")


def replace_paragraph_text_with_formatting(paragraph, new_text, original_formats):
    """
    Replaces paragraph text while preserving word-level formatting.
    
    Args:
        paragraph: Paragraph to modify
        new_text: New text to apply
        original_formats: Original formatting captured
    """
    if not paragraph.runs:
        logger.warning("No runs found, adding plain text")
        paragraph.add_run(new_text)
        return
    
    # SAFE REMOVAL: Clear text from existing runs instead of removing XML elements
    # This preserves document structure for tables, lists, headers, footers
    logger.info(f"Clearing text from {len(paragraph.runs)} existing runs")
    for run in paragraph.runs:
        run.text = ""
    
    logger.info(f"After clearing, paragraph has {len(paragraph.runs)} runs")
    
    # Apply proportional formatting distribution
    logger.info("Applying proportional formatting distribution")
    apply_proportional_formatting_safe(paragraph, new_text, original_formats)
    
    logger.info(f"Paragraph now has {len(paragraph.runs)} runs with text: '{paragraph.text}'")


class LLMService:
    """
    Enhanced LLM service with robust JSON parsing and intelligent fallbacks.
    """
    
    def __init__(self, ai_config: Dict[str, Any]):
        self.ai_config = ai_config
        self.provider = None
        self.model_manager = ModelManager()
        
        # Initialize provider based on config
        self._initialize_provider()
    
    def _initialize_provider(self):
        """Initialize the LLM provider based on configuration."""
        try:
            provider_name = self.ai_config.get('provider', 'ollama')
            model_name = self.ai_config.get('model', 'qwen2.5:7b')
            
            if provider_name == 'ollama':
                self.provider = OllamaProvider(self.ai_config)
            else:
                # Fallback to Ollama if other providers fail
                logger.warning(f"Provider {provider_name} not implemented, falling back to Ollama")
                self.provider = OllamaProvider(self.ai_config)
            
            if self.provider and self.provider.is_available():
                logger.info(f"LLM provider initialized: {provider_name} with model {model_name}")
            else:
                logger.error(f"LLM provider {provider_name} is not available")
                
        except Exception as e:
            logger.error(f"Error initializing LLM provider: {e}")
            # Create a fallback provider
            self.provider = OllamaProvider(self.ai_config)
    
    def analyze_commit_impact(self, commit_context: Dict[str, Any]) -> Dict[str, Any]:
        """
        Analyze the impact of a commit on documentation.
        
        Args:
            commit_context: Commit information and context
            
        Returns:
            Analysis results with impact level and recommendations
        """
        try:
            commit_message = commit_context.get('message', '')
            files_changed = commit_context.get('files', [])
            diff_content = commit_context.get('diff', '')
            
            prompt = f"""
            Analyze this commit to determine its impact on documentation:
            
            COMMIT MESSAGE: {commit_message}
            FILES CHANGED: {', '.join(files_changed)}
            DIFF PREVIEW: {diff_content[:500] if diff_content else 'No diff available'}
            
            Determine:
            1. Impact level: low, medium, high
            2. Key changes that need documentation
            3. Recommended documentation types to update
            
            Return ONLY valid JSON in this exact format:
            {{
                "impact_level": "high|medium|low",
                "key_changes": ["change1", "change2"],
                "reasoning": "explanation",
    "recommended_docs": [
                    {{"doc_type": "user_guide|api_reference|examples", "priority": "high|medium|low", "reason": "explanation"}}
    ]
}}
"""
            
            response = self.provider.generate_response(prompt)
            analysis = self._parse_json_safely(response)
            
            if analysis:
                logger.info(f"Commit analysis completed: {analysis.get('impact_level', 'unknown')} impact")
                return analysis
            else:
                return self._create_fallback_commit_analysis(commit_message, files_changed)
            
        except Exception as e:
            logger.error(f"Error analyzing commit impact: {e}")
            return self._create_fallback_commit_analysis(commit_context.get('message', ''), commit_context.get('files', []))
    
    def select_documentation_files(self, analysis_result: Dict[str, Any], available_files: List[str]) -> List[str]:
        """
        Select which documentation files should be updated based on analysis.
        
        Args:
            analysis_result: Commit analysis results
            available_files: List of available documentation files
            
        Returns:
            List of selected file paths
        """
        try:
            impact_level = analysis_result.get('impact_level', 'low')
            key_changes = analysis_result.get('key_changes', [])
            recommended_docs = analysis_result.get('recommended_docs', [])
            
            # Simple logic: if high impact, select all available files
            if impact_level == 'high':
                logger.info(f"High impact commit - selecting all {len(available_files)} files")
                return available_files
            elif impact_level == 'medium':
                # Select first half of files
                selected_count = max(1, len(available_files) // 2)
                selected_files = available_files[:selected_count]
                logger.info(f"Medium impact commit - selecting {len(selected_files)} files")
                return selected_files
            else:
                # Low impact - select first file only
                selected_files = available_files[:1] if available_files else []
                logger.info(f"Low impact commit - selecting {len(selected_files)} files")
                return selected_files
            
        except Exception as e:
            logger.error(f"Error selecting documentation files: {e}")
            return available_files
    
    def analyze_document_structure_and_placement(self, commit_context: Dict[str, Any], doc_content: str, rag_context: str = "") -> Dict[str, Any]:
        """
        Analyze document structure and determine where to place new content.
        
        Args:
            commit_context: Commit information including analysis results
            doc_content: Document content
            rag_context: Additional context from RAG
            
        Returns:
            Analysis results with placement recommendations
        """
        try:
            commit_message = commit_context.get('message', '')
            files_changed = commit_context.get('files', [])
            diff_content = commit_context.get('diff', '')
            analysis_result = commit_context.get('analysis', {})
            
            # Generate LLM summary of the diff
            diff_summary = self._generate_diff_summary(commit_message, diff_content, files_changed)
            
            # Extract key information from analysis
            impact_level = analysis_result.get('impact_level', 'low')
            key_changes = analysis_result.get('key_changes', [])
            reasoning = analysis_result.get('reasoning', '')
            recommended_docs = analysis_result.get('recommended_docs', [])
            
            # Parse document into sections
            sections = self._parse_document_sections(doc_content)
            
            prompt = f"""
            You are an expert documentation analyst. Analyze this document structure and determine the optimal placement for new content based on the commit information.

            COMMIT SUMMARY:
            - Message: {commit_message}
            - Impact Level: {impact_level}
            - Key Changes: {', '.join(key_changes)}
            - Reasoning: {reasoning}
            - Files Changed: {', '.join(files_changed)}
            
            LLM-GENERATED DIFF SUMMARY:
            {diff_summary}
            
            RAG CONTEXT:
            {rag_context[:500] if rag_context else 'No additional context'}
            
            DOCUMENT STRUCTURE:
            {self._format_sections_for_analysis(sections)}
            
            ANALYSIS REQUIREMENTS:
            1. Use the diff summary to understand exactly what changed
            2. Identify which sections are most relevant to the changes
            3. Determine the best section for placing new content
            4. Consider the commit's impact level and key changes
            5. Respect document structure and avoid protected sections (title, TOC, etc.)
            6. Match content type to section purpose
            
            Return ONLY valid JSON in this exact format:
            {{
                "section_analysis": [
                    {{
                        "heading": "section_name",
                        "relevance_score": 1-10,
                        "should_update": true/false,
                        "update_reason": "detailed explanation based on diff summary",
                        "content_type": "paragraph|table|list|example",
                        "protected": true/false
                    }}
                ],
                "placement_decision": {{
                    "target_section": "section_name",
                    "placement_type": "append|insert|replace",
                    "content_type": "description|example|table|feature_list",
                    "reasoning": "detailed explanation of why this placement is optimal based on diff summary",
                    "confidence": 1-10
                }},
                "content_specification": {{
                    "style": "formal|informal|technical",
                    "length": "short|medium|long",
                    "focus": "what|how|why|examples",
                    "examples_needed": true/false
                }}
}}
"""
            
            response = self.provider.generate_response(prompt)
            analysis = self._parse_json_safely(response)
            
            if analysis:
                logger.info(f"Document structure analysis completed with diff summary context")
                return analysis
            else:
                return self._create_fallback_document_analysis(commit_context, doc_content)
                
        except Exception as e:
            logger.error(f"Error analyzing document structure: {e}")
            return self._create_fallback_document_analysis(commit_context, doc_content)

    def _generate_diff_summary(self, commit_message: str, diff_content: str, files_changed: List[str]) -> str:
        """
        Generate an LLM summary of the diff content to provide context for document analysis.
        
        Args:
            commit_message: Commit message
            diff_content: Full diff content
            files_changed: List of changed files
            
        Returns:
            LLM-generated summary of the changes
        """
        try:
            # Truncate diff if too long to avoid token limits
            diff_preview = diff_content[:2000] if len(diff_content) > 2000 else diff_content
            
            prompt = f"""
                You are an expert code-change summarizer for an automated documentation update system.
                You will be given details of a Git commit and its diff.

                Your goal is to:
                1. Analyze the commit message and diff precisely.
                2. Extract a structured summary of what changed.
                3. Focus on information relevant to documentation updates.

                ---

                COMMIT MESSAGE:
                {commit_message}

                FILES CHANGED:
                {', '.join(files_changed)}
            
            DIFF CONTENT:
            {diff_preview}
            
                ---

                ### OUTPUT REQUIREMENTS

                Provide a clear, structured **natural-language summary** with the following sections:

                #### 1. Overall Summary
                A concise 2–3 sentence overview of what the commit does and its intent.

                #### 2. Key Changes by File
                For each affected file, list:
                - Functions, classes, or methods added, modified, or removed  
                - Each entry should include:
                - **Name** of the function/class/method  
                - **Parameters** (if any)  
                - **Return type** (if inferable)  
                - **Brief description** of what changed or what it does  

                #### 3. Technical Notes
                Include any:
                - New public APIs, CLI flags, config options, or environment variables  
                - Changes in function behavior or side effects  
                - Breaking changes, deprecations, or migration notes  

                ---

                ### STYLE NOTES
                - Keep it concise but informative — think “developer-facing changelog.”  
                - Focus on documentation-impacting changes.  
                - Ignore trivial edits (comments, formatting, etc.).  
                - Use bullet points or subheadings for readability.

                Now analyze the commit and produce the structured summary.
                """

            
            if self.provider and self.provider.is_available():
                response = self.provider.generate_response(prompt)
                logger.info("Generated LLM diff summary for document analysis")
                return response.strip()
            else:
                # Fallback summary
                return f"Commit '{commit_message}' modified {len(files_changed)} files: {', '.join(files_changed)}"
            
        except Exception as e:
            logger.error(f"Error generating diff summary: {e}")
            return f"Commit '{commit_message}' modified {len(files_changed)} files: {', '.join(files_changed)}"

    def generate_contextual_content(self, commit_context: Dict[str, Any], analysis: Dict[str, Any], target_section_content: str) -> str:
        """
        Generate contextual content that matches the target section's style.
        
        Args:
            commit_context: Commit information
            analysis: Document analysis results
            target_section_content: Content of the target section
            
        Returns:
            Generated content
        """
        try:
            commit_message = commit_context.get('message', '')
            files_changed = commit_context.get('files', [])
            
            prompt = f"""
            Generate content for a documentation update based on this commit:
            
            COMMIT MESSAGE: {commit_message}
            FILES CHANGED: {', '.join(files_changed)}
            
            TARGET SECTION STYLE (for matching):
            {target_section_content[:500]}
            
            ANALYSIS:
            {json.dumps(analysis, indent=2)}
            
            Generate content that:
            1. Matches the style and tone of the target section
            2. Is relevant to the commit changes
            3. Provides useful information for users
            4. Is concise but informative
            
            Return ONLY the content text, no explanations or metadata.
            """
            
            response = self.provider.generate_response(prompt)
            return response.strip()
            
        except Exception as e:
            logger.error(f"Error generating contextual content: {e}")
            return f"New feature: {commit_context.get('message', '')}"

    def determine_docx_edit_type(self, commit_context: Dict[str, Any], doc_content: str) -> str:
        """
        Determine what type of DOCX edit is needed.
        
        Args:
            commit_context: Commit information
            doc_content: Document content
            
        Returns:
            Edit type: "edit_line", "add_paragraph", "edit_table", or "no_change"
        """
        try:
            commit_message = commit_context.get('message', '').lower()
            
            # Simple keyword-based logic
            if any(keyword in commit_message for keyword in ['feat', 'feature', 'new', 'add', 'introduce']):
                return "add_paragraph"
            elif any(keyword in commit_message for keyword in ['fix', 'bug', 'update', 'modify']):
                return "edit_line"
            elif any(keyword in commit_message for keyword in ['table', 'data', 'structure']):
                return "edit_table"
            else:
                return "no_change"
                
        except Exception as e:
            logger.error(f"Error determining DOCX edit type: {e}")
            return "no_change"
    
    def generate_documentation_update(self, current_content: str, commit_context: Dict[str, Any], rag_context: str) -> str:
        """
        Generate updated documentation content based on commit changes.
        
        Args:
            current_content: Current document content
            commit_context: Commit information
            rag_context: RAG context
            
        Returns:
            Updated content
        """
        try:
            commit_message = commit_context.get('message', '')
            files_changed = commit_context.get('files', [])
            
            prompt = f"""
            Update this documentation based on the commit changes:
            
            COMMIT MESSAGE: {commit_message}
            FILES CHANGED: {', '.join(files_changed)}
            
            RAG CONTEXT: {rag_context}
            
            CURRENT CONTENT:
            {current_content[:1000]}
            
            Generate updated content that incorporates the new changes while maintaining
            the existing structure and style. Return only the updated content.
            """
            
            if self.provider and self.provider.is_available():
                response = self.provider.generate_response(prompt)
                return response.strip()
            else:
                # Fallback content
                return f"{current_content}\n\n## New Feature: {commit_message}\n\nThis feature was added in the latest update."
                
        except Exception as e:
            logger.error(f"Error generating documentation update: {e}")
            return f"{current_content}\n\n## Update: {commit_context.get('message', 'New changes')}\n\nDocumentation updated based on recent changes."
    
    def analyze_document_structure_and_placement_enhanced(self, commit_context: Dict[str, Any], doc_content: str, rag_context: str = "") -> Dict[str, Any]:
        """
        Enhanced document structure analysis using heading-by-heading approach with formatting preservation.
        
        Args:
            commit_context: Commit information including analysis results
            doc_content: Document content (for text-based analysis)
            rag_context: Additional context from RAG
            
        Returns:
            Analysis results with placement recommendations and formatting info
        """
        try:
            commit_message = commit_context.get('message', '')
            files_changed = commit_context.get('files', [])
            diff_content = commit_context.get('diff', '')
            analysis_result = commit_context.get('analysis', {})
            
            # Generate LLM summary of the diff
            diff_summary = self._generate_diff_summary(commit_message, diff_content, files_changed)
            
            # Extract key information from analysis
            impact_level = analysis_result.get('impact_level', 'low')
            key_changes = analysis_result.get('key_changes', [])
            reasoning = analysis_result.get('reasoning', '')
            recommended_docs = analysis_result.get('recommended_docs', [])
            
            # For now, return a structure that indicates we should use heading-by-heading approach
            # The actual DOCX processing will be handled by the workflow
            analysis = {
                "use_heading_by_heading": True,
                "commit_message": commit_message,
                "files_changed": files_changed,
                "diff_summary": diff_summary,
                "impact_level": impact_level,
                "key_changes": key_changes,
                "reasoning": reasoning,
                "rag_context": rag_context,
                "placement_strategy": "heading_by_heading_with_formatting"
            }
            
            logger.info("Enhanced document structure analysis completed - using heading-by-heading approach")
            return analysis
            
        except Exception as e:
            logger.error(f"Error in enhanced document structure analysis: {e}")
            return self._create_fallback_document_analysis(commit_context, doc_content)
    
    def process_docx_with_whole_document(self, docx_path: str, commit_context: Dict[str, Any], rag_context: str = "", use_instruction_approach: bool = True) -> Dict[str, Any]:
        """
        Process DOCX file directly using python-docx Document object.

        Args:
            docx_path: Path to the DOCX file
            commit_context: Commit information
            rag_context: Additional context from RAG

        Returns:
            Processing results with updated content
        """
        try:
            doc = Document(docx_path)
            commit_message = commit_context.get('message', '')
            files_changed = commit_context.get('files', [])
            diff_summary = commit_context.get('diff_summary', '')

            # Extract structured text content from DOCX
            doc_content = self._extract_docx_text_content(doc)

            if use_instruction_approach:
                # New approach: Generate instruction, then pass to LLM for direct update
                instruction = self._generate_update_instruction(doc_content, commit_message, files_changed, diff_summary, rag_context)

                if instruction and len(instruction.strip()) > 10:
                    logger.info("📄 ORIGINAL DOCUMENT STRUCTURE:")
                    self._log_document_structure(doc, "BEFORE")
                    
                    logger.info(f"📝 Generated instruction: {instruction}")
                    
                    # Parse instruction and apply updates directly to DOCX
                    updates_applied = self._parse_and_apply_instruction_to_docx(doc, instruction, commit_message, files_changed, diff_summary, rag_context)
                
                    if updates_applied > 0:
                        logger.info("📄 UPDATED DOCUMENT STRUCTURE:")
                        self._log_document_structure(doc, "AFTER")
                        
                        logger.info("📋 SUMMARY OF CHANGES:")
                        logger.info(f"   🔢 Instruction applied: {instruction[:100]}...")
                        logger.info(f"   📄 Updates applied: {updates_applied}")

                        return {
                            'success': True,
                            'method': 'instruction_based_docx_update',
                            'updated_doc': doc,
                            'original_length': len(doc_content),
                            'updates_applied': updates_applied,
                            'instruction': instruction,
                            'rag_context': rag_context
                        }
                    else:
                        logger.warning("No updates were applied based on the instruction")
                        return {
                            'success': False,
                            'error': 'No updates were applied based on the instruction',
                            'method': 'instruction_based_docx_update'
                        }
                else:
                    logger.warning("LLM failed to generate meaningful instruction")
                    return {
                        'success': False,
                        'error': 'LLM failed to generate meaningful instruction',
                        'method': 'llm_direct_update'
                    }
            else:
                # Fallback to original approach (if needed)
                logger.info("Using fallback approach - direct document update")
                return {
                    'success': False,
                    'error': 'Fallback approach not implemented',
                    'method': 'fallback_not_implemented'
                }

        except Exception as e:
            logger.error(f"Error processing DOCX with direct approach: {e}")
            return {
                'success': False,
                'error': str(e),
                'method': 'direct_docx_update'
            }
    
    def process_docx_with_heading_by_heading(self, docx_path: str, commit_context: Dict[str, Any], rag_context: str = "") -> Dict[str, Any]:
        """
        Process DOCX file using heading-by-heading approach with formatting preservation.
        
        Args:
            docx_path: Path to the DOCX file
            commit_context: Commit information
            rag_context: Additional context from RAG
            
        Returns:
            Processing results with updates made
        """
        try:
            doc = Document(docx_path)
            commit_message = commit_context.get('message', '')
            files_changed = commit_context.get('files', [])
            diff_summary = commit_context.get('diff_summary', '')
            
            # Step 1: Parse document structure
            headings_with_content = []
            current_section = None
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                
                if not text:
                    continue
                    
                # Check if this is a heading
                is_heading = False
                heading_level = 0
                
                if paragraph.style.name.startswith('Heading'):
                    is_heading = True
                    heading_level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
                elif paragraph.runs and any(run.bold for run in paragraph.runs):
                    if len(text) < 100 and not text.endswith('.'):
                        is_heading = True
                        heading_level = 1
                
                if is_heading:
                    # Save previous section if it has content
                    if current_section and current_section['paragraphs']:
                        headings_with_content.append(current_section)
                    
                    # Start new section
                    current_section = {
                        'heading': text,
                        'level': heading_level,
                        'paragraph_index': i,
                        'paragraphs': []
                    }
                else:
                    # Add paragraph to current section
                    if current_section:
                        current_section['paragraphs'].append({
                            'text': text,
                            'index': i,
                            'paragraph_obj': paragraph
                        })
            
            # Don't forget the last section
            if current_section and current_section['paragraphs']:
                headings_with_content.append(current_section)
            
            logger.info(f"Found {len(headings_with_content)} headings with content")
            
            # Step 2: Use RAG to find relevant sections based on commit diff summary
            relevant_headings = []
            
            # First, generate a summary of the commit changes
            diff_content = commit_context.get('diff', '')
            commit_summary = self._generate_commit_summary(commit_message, files_changed, diff_content)
            logger.info(f"Commit summary: {commit_summary}")
            
            # Use RAG to find relevant sections in the document
            logger.info("Using RAG to find relevant document sections...")
            relevant_sections = self._find_relevant_sections_with_rag(
                commit_summary, headings_with_content, rag_context
            )
            
            # Process the RAG results
            for section_data in relevant_sections:
                section = section_data['section']
                similarity_score = section_data.get('similarity_score', 0)
                
                relevant_headings.append({
                    'heading': section['heading'],
                    'level': section['level'],
                    'relevance_score': int(similarity_score * 10),  # Convert to 1-10 scale
                    'reasons': [f"RAG similarity: {similarity_score:.3f}"],
                    'paragraphs': section['paragraphs'],
                    'paragraph_index': section['paragraph_index']
                })
                
                logger.info(f"RAG found relevant heading: {section['heading']} (similarity: {similarity_score:.3f})")
            
            logger.info(f"Found {len(relevant_headings)} RAG-relevant headings")
            
            # Step 3: Process ALL tables first (before processing headings)
            logger.info("Processing all tables in document")
            table_updates = []
            
            for table_idx, table in enumerate(doc.tables):
                logger.info(f"Analyzing table {table_idx + 1}")
                
                try:
                    # Extract table content
                    table_data = self._extract_table_content(table)
                    
                    if not table_data:
                        logger.info(f"Table {table_idx + 1} has no content, skipping")
                        continue
                    
                    # Analyze table with LLM (using general commit context)
                    table_analysis = self._analyze_table_with_llm(table_data, "Document Tables", "", commit_context)
                    
                    if table_analysis.get('is_relevant', False):
                        confidence = table_analysis.get('confidence', 0.0)
                        
                        # Get confidence threshold from configuration
                        from utils.config import get_settings
                        settings = get_settings()
                        confidence_threshold = settings.document.table_confidence_threshold
                        
                        # Also check if we can actually add a row
                        can_add_row = table_analysis.get('can_add_row', False)
                        
                        if confidence >= confidence_threshold and can_add_row:
                            logger.info(f"Table {table_idx + 1} needs updates (confidence: {confidence:.2f}): {table_analysis.get('relevance_reason', 'No reason provided')}")
                            
                            # Update table data
                            updated_table_data = self._update_table_with_new_data(table, table_data, table_analysis)
                            
                            # Replace table in document
                            success = self._replace_table_in_document(doc, table, updated_table_data, table_idx + 1)
                            
                            if success:
                                logger.info(f"Successfully updated table {table_idx + 1}")
                                table_updates.append({
                                    'table_updated': True,
                                    'table_index': table_idx + 1,
                                    'table_purpose': table_analysis.get('table_purpose', 'Unknown'),
                                    'updates_applied': 1 if can_add_row else 0,
                                    'confidence': confidence
                                })
                            else:
                                logger.error(f"Failed to update table {table_idx + 1}")
                        else:
                            logger.info(f"Table {table_idx + 1} update skipped - confidence {confidence:.2f} below threshold {confidence_threshold}")
                            table_updates.append({
                                'table_skipped': True,
                                'table_index': table_idx + 1,
                                'table_purpose': table_analysis.get('table_purpose', 'Unknown'),
                                'skip_reason': f"Low confidence: {confidence:.2f} < {confidence_threshold}",
                                'confidence': confidence
                            })
                    else:
                        logger.info(f"Table {table_idx + 1} does not need updates")
                        
                except Exception as e:
                    logger.error(f"Error processing table {table_idx + 1}: {e}")
                    continue
            
            logger.info(f"Table processing completed. Updated: {len([u for u in table_updates if u.get('table_updated')])} tables")
            
            # Step 4: Process headings for content updates
            updates_made = []
            
            for analysis in relevant_headings:
                heading = analysis['heading']
                paragraphs = analysis['paragraphs']
                
                logger.info(f"Updating: [{analysis['level']}] {heading}")
                logger.info(f"Relevance: {analysis['relevance_score']}/10")
                logger.info(f"Reasons: {'; '.join(analysis['reasons'])}")
                logger.info(f"Paragraphs available: {len(paragraphs)}")
                
                # Extract heading content from paragraphs for context
                heading_content = ""
                if paragraphs:
                    # Get text from first few paragraphs for context
                    content_parts = []
                    for para_info in paragraphs[:3]:  # Use first 3 paragraphs for context
                        if para_info.get('text'):
                            content_parts.append(para_info['text'])
                    heading_content = " ".join(content_parts)
                
                # Generate content for this heading
                new_content = self._generate_content_for_heading(heading, commit_message, files_changed, diff_summary, heading_content)
                
                # Use LLM to find the best paragraph to duplicate
                best_para_info = self._select_best_paragraph_with_llm(paragraphs, heading, commit_message, files_changed, diff_summary)
                
                if not best_para_info:
                    logger.info("LLM could not determine best paragraph - skipping")
                    continue
                    
                best_para = best_para_info['paragraph_obj']
                
                logger.info(f"LLM selected paragraph: {best_para_info['text'][:80]}...")
                logger.info(f"LLM reasoning: {best_para_info.get('reasoning', 'No reasoning provided')}")
                
                # Capture original formatting before duplication
                original_formats = capture_detailed_formatting(best_para)
                logger.info(f"Captured {len(original_formats)} runs with formatting from original paragraph")
                
                # Find the best insertion point within the section
                insertion_index = self._find_best_insertion_point(paragraphs, best_para_info, heading, commit_message)
                
                # Duplicate the paragraph by copying its element
                parent = best_para._element.getparent()
                new_element = best_para._element.__copy__()
                parent.insert(insertion_index, new_element)
                
                # Find the new paragraph object
                new_paragraph = None
                for para in doc.paragraphs:
                    if para._element == new_element:
                        new_paragraph = para
                        break
                
                if new_paragraph:
                    # Update the duplicated paragraph with new content while preserving formatting
                    logger.info(f"Updating duplicated paragraph with formatting preservation")
                    replace_paragraph_text_with_formatting(new_paragraph, new_content, original_formats)
                    logger.info(f"Added: {new_content[:80]}...")
                    
                    updates_made.append({
                        'heading': heading,
                        'content_added': new_content,
                        'original_paragraph': best_para_info['text'][:100] + "...",
                        'relevance_score': analysis['relevance_score'],
                        'llm_reasoning': best_para_info.get('reasoning', 'No reasoning provided'),
                        'formatting_preserved': len(original_formats)
                    })
                else:
                    logger.error("Failed to find duplicated paragraph")
                
            # Save the updated document and return the bytes
            import io
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            updated_content = doc_bytes.getvalue()
            
            return {
                'success': True,
                'updates_made': updates_made + table_updates,  # Combine heading updates and table updates
                'headings_analyzed': len(headings_with_content),
                'relevant_headings': len(relevant_headings),
                'tables_processed': len(doc.tables),
                'tables_updated': len([u for u in table_updates if u.get('table_updated')]),
                'method': 'heading_by_heading_with_formatting',
                'updated_content': updated_content
            }
                
        except Exception as e:
            logger.error(f"Error processing DOCX with heading-by-heading approach: {e}")
            return {
                'success': False,
                'error': str(e),
                'method': 'heading_by_heading_with_formatting'
            }
    
    def _generate_commit_summary(self, commit_message: str, files_changed: List[str], diff_content: str) -> str:
        """Generate a concise summary of what changed in the commit."""
        try:
            # Truncate diff content for prompt
            diff_preview = diff_content[:2000] + "..." if len(diff_content) > 2000 else diff_content
            
            prompt = f"""
You are an expert code-change summarizer for an automated documentation update system.
You will be given details of a Git commit and its diff.

Your goal is to:
1. Analyze the commit message and diff precisely.
2. Extract a structured summary of what changed.
3. Focus on information relevant to documentation updates.

---

COMMIT MESSAGE:
{commit_message}

FILES CHANGED:
{', '.join(files_changed)}

DIFF CONTENT:
{diff_preview}

---

### OUTPUT REQUIREMENTS

Provide a clear, structured **natural-language summary** with the following sections:

#### 1. Overall Summary
A concise 2–3 sentence overview of what the commit does and its intent.

#### 2. Key Changes by File
For each affected file, list:
- Functions, classes, or methods added, modified, or removed  
- Each entry should include:
- **Name** of the function/class/method  
- **Parameters** (if any)  
- **Return type** (if inferable)  
- **Brief description** of what changed or what it does  

#### 3. Technical Notes
Include any:
- New public APIs, CLI flags, config options, or environment variables  
- Changes in function behavior or side effects  
- Breaking changes, deprecations, or migration notes  

---

### STYLE NOTES
- Keep it concise but informative — think "developer-facing changelog."  
- Focus on documentation-impacting changes.  
- Ignore trivial edits (comments, formatting, etc.).  
- Use bullet points or subheadings for readability.

Now analyze the commit and produce the structured summary.
"""
            
            response = self.provider.generate_response(
                prompt=prompt,
                max_tokens=800,
                temperature=0.2
            )
            return response.strip()
            
        except Exception as e:
            logger.error(f"Error generating commit summary: {e}")
            return f"Commit '{commit_message}' modified {len(files_changed)} files: {', '.join(files_changed)}"
    
    def _generate_heading_summary(self, heading: str, content: str) -> str:
        """Generate a concise summary of what a heading section is about."""
        try:
            prompt = f"""
            Analyze this documentation heading and its content to provide a concise summary:
            
            Heading: {heading}
            Content: {content[:1500]}...
            
            Provide a 1-2 sentence summary focusing on:
            - What this section is about
            - What topics or concepts it covers
            - The main purpose of this section
            
            Summary:
            """
            
            response = self.provider.generate_response(
                prompt=prompt,
                max_tokens=150,
                temperature=0.3
            )
            return response.strip()
            
        except Exception as e:
            logger.error(f"Error generating heading summary: {e}")
            return f"Section about {heading.lower()}"
    
    def _analyze_semantic_relevance(self, commit_summary: str, heading_summary: str, heading: str, commit_message: str) -> Dict[str, Any]:
        """Use LLM to determine if a heading is semantically relevant to the commit changes."""
        try:
            prompt = f"""
            Analyze whether this documentation heading is relevant to the commit changes.
            
            COMMIT CHANGES:
            {commit_summary}
            
            HEADING SECTION:
            Heading: {heading}
            Content Summary: {heading_summary}
            
            Determine if this heading section should be updated based on the commit changes.
            Consider semantic relevance, not just keyword matching.
            
            Respond with JSON:
            {{
                "is_relevant": true/false,
                "relevance_score": 1-10,
                "reasons": ["reason1", "reason2"],
                "reason": "Brief explanation of relevance or why not relevant"
            }}
            
            Scoring guidelines:
            - 9-10: Directly related to the changes (e.g., new feature documentation)
            - 7-8: Strongly related (e.g., usage examples, recommendations)
            - 5-6: Moderately related (e.g., general concepts that apply)
            - 3-4: Weakly related (e.g., tangential concepts)
            - 1-2: Barely related (e.g., unrelated technical details)
            - 0: Not relevant at all
            """
            
            response = self.provider.generate_response(
                prompt=prompt,
                max_tokens=300,
                temperature=0.2
            )
            
            # Parse JSON response
            import json
            try:
                # Extract JSON from response
                json_start = response.find('{')
                json_end = response.rfind('}') + 1
                if json_start != -1 and json_end != -1:
                    json_str = response[json_start:json_end]
                    analysis = json.loads(json_str)
                    
                    # Validate required fields
                    if all(key in analysis for key in ['is_relevant', 'relevance_score', 'reasons', 'reason']):
                        return analysis
                    else:
                        logger.warning("Invalid analysis response structure")
                        return self._create_fallback_relevance_analysis(heading, commit_message)
                else:
                    logger.warning("Could not find JSON in response")
                    return self._create_fallback_relevance_analysis(heading, commit_message)
                    
            except json.JSONDecodeError as e:
                logger.warning(f"Failed to parse JSON response: {e}")
                return self._create_fallback_relevance_analysis(heading, commit_message)
                            
        except Exception as e:
            logger.error(f"Error analyzing semantic relevance: {e}")
            return self._create_fallback_relevance_analysis(heading, commit_message)
    
    def _analyze_batch_semantic_relevance(self, commit_summary: str, heading_summaries: List[Dict], commit_message: str) -> List[Dict[str, Any]]:
        """Use LLM to analyze semantic relevance for all headings in a single batch comparison."""
        try:
            # Create a formatted list of all headings and their summaries
            headings_text = ""
            for i, heading_data in enumerate(heading_summaries):
                headings_text += f"{i+1}. HEADING: {heading_data['heading']}\n"
                headings_text += f"   SUMMARY: {heading_data['summary']}\n\n"
            
            prompt = f"""
            Analyze the semantic relevance between a commit change and multiple documentation headings.
            
            COMMIT SUMMARY:
            {commit_summary}
            
            COMMIT MESSAGE: {commit_message}
            
            AVAILABLE HEADINGS AND THEIR SUMMARIES:
            {headings_text}
            
            For each heading, determine if it's semantically relevant to the commit changes.
            Consider:
            - Does the heading content relate to the changes made?
            - Would updating this section make sense for users?
            - Is there conceptual overlap between the commit and heading?
            
            Respond with JSON only, analyzing ALL headings:
            {{
                "headings": [
                    {{
                        "heading_number": 1,
                        "heading_name": "heading name",
                        "is_relevant": true/false,
                        "relevance_score": 1-10,
                        "reasons": ["reason1", "reason2"],
                        "reason": "brief explanation if not relevant"
                    }},
                    ...
                ]
            }}
            """
            
            response = self.provider.generate_response(
                prompt,
                max_tokens=800,
                temperature=0.3
            )
            
            # Parse JSON response
            try:
                import json
                analysis = json.loads(response.strip())
                
                # Convert the response back to the expected format
                results = []
                for heading_analysis in analysis.get('headings', []):
                    heading_number = heading_analysis.get('heading_number', 1)
                    if 1 <= heading_number <= len(heading_summaries):
                        heading_data = heading_summaries[heading_number - 1]
                        
                        results.append({
                            'heading': heading_data['heading'],
                            'section': heading_data['section'],
                            'is_relevant': heading_analysis.get('is_relevant', False),
                            'relevance_score': heading_analysis.get('relevance_score', 0),
                            'reasons': heading_analysis.get('reasons', []),
                            'reason': heading_analysis.get('reason', 'Not relevant')
                        })
                
                return results
                
            except json.JSONDecodeError:
                logger.warning("Invalid JSON response from LLM, using fallback")
                return self._create_fallback_batch_relevance_analysis(heading_summaries, commit_message)
            
        except Exception as e:
            logger.error(f"Error analyzing batch semantic relevance: {e}")
            return self._create_fallback_batch_relevance_analysis(heading_summaries, commit_message)
    
    def _create_fallback_relevance_analysis(self, heading: str, commit_message: str) -> Dict[str, Any]:
        """Create a fallback relevance analysis when LLM analysis fails."""
        heading_lower = heading.lower()
        commit_lower = commit_message.lower()
        
        # Simple keyword-based fallback
        if any(word in commit_lower for word in ['interactive', 'builder', 'demo']):
            if any(word in heading_lower for word in ['example', 'demo', 'usage', 'recommendation']):
                return {
                    'is_relevant': True,
                    'relevance_score': 7,
                    'reasons': ['Contains interactive/builder keywords'],
                    'reason': 'Fallback: keyword matching suggests relevance'
                }
        
        return {
            'is_relevant': False,
            'relevance_score': 0,
            'reasons': [],
            'reason': 'Fallback: no clear relevance found'
        }
    
    def _create_fallback_batch_relevance_analysis(self, heading_summaries: List[Dict], commit_message: str) -> List[Dict[str, Any]]:
        """Create a fallback batch relevance analysis when LLM analysis fails."""
        results = []
        commit_lower = commit_message.lower()
        
        for heading_data in heading_summaries:
            heading = heading_data['heading']
            heading_lower = heading.lower()
            
            # Simple keyword-based fallback for each heading
            if any(word in commit_lower for word in ['interactive', 'builder', 'demo']):
                if any(word in heading_lower for word in ['example', 'demo', 'usage', 'recommendation']):
                    results.append({
                        'heading': heading,
                        'section': heading_data['section'],
                        'is_relevant': True,
                        'relevance_score': 7,
                        'reasons': ['Contains interactive/builder keywords'],
                        'reason': 'Fallback: keyword matching suggests relevance'
                    })
                    continue
            
            results.append({
                'heading': heading,
                'section': heading_data['section'],
                'is_relevant': False,
                'relevance_score': 0,
                'reasons': [],
                'reason': 'Fallback: no clear relevance found'
            })
        
        return results
    
    def _find_relevant_sections_with_rag(self, commit_summary: str, headings_with_content: List[Dict], rag_context: str) -> List[Dict]:
        """Use RAG to find relevant document sections based on commit summary."""
        try:
            from services.rag_service import RAGService
            from utils.config import get_settings
            
            settings = get_settings()
            rag_service = RAGService(settings.ai)
            
            # Create document chunks from headings
            document_chunks = []
            for i, section in enumerate(headings_with_content):
                heading = section['heading']
                paragraphs = section['paragraphs']
                
                # Create content chunk with ALL content under the heading
                content = f"HEADING: {heading}\n\n"
                content += "\n".join([p['text'] for p in paragraphs])  # All paragraphs under heading
                
                document_chunks.append({
                    'id': f"heading_{i}",
                    'content': content,
                    'heading': heading,
                    'section': section
                })
            
            logger.info(f"Created {len(document_chunks)} document chunks for RAG analysis")
            
            # Use RAG to find similar sections
            query = f"COMMIT SUMMARY: {commit_summary}\n\nFind document sections that should be updated based on this commit."
            
            # Get embeddings for the query
            query_embedding = rag_service._get_query_embedding(query)
            
            # Get embeddings for document chunks and calculate similarities
            relevant_sections = []
            for chunk in document_chunks:
                try:
                    chunk_embedding = rag_service._get_query_embedding(chunk['content'])
                    
                    # Calculate cosine similarity
                    import numpy as np
                    similarity = np.dot(query_embedding, chunk_embedding) / (
                        np.linalg.norm(query_embedding) * np.linalg.norm(chunk_embedding)
                    )
                    
                    # Only include sections with reasonable similarity
                    if similarity > 0.1:  # Threshold for relevance
                        relevant_sections.append({
                            'section': chunk['section'],
                            'similarity_score': float(similarity),
                            'content': chunk['content']
                        })
                        
                except Exception as e:
                    logger.warning(f"Error processing chunk {chunk['heading']}: {e}")
                    continue
            
            # Sort by similarity score (highest first)
            relevant_sections.sort(key=lambda x: x['similarity_score'], reverse=True)
            
            logger.info(f"RAG found {len(relevant_sections)} relevant sections")
            return relevant_sections
            
        except Exception as e:
            logger.error(f"Error in RAG-based section finding: {e}")
            # Fallback to simple keyword matching
            return self._fallback_keyword_matching(commit_summary, headings_with_content)
    
    def _fallback_keyword_matching(self, commit_summary: str, headings_with_content: List[Dict]) -> List[Dict]:
        """Fallback keyword matching when RAG fails."""
        commit_lower = commit_summary.lower()
        relevant_sections = []
        
        for section in headings_with_content:
            heading = section['heading']
            heading_lower = heading.lower()
            
            # Simple keyword matching
            if any(word in commit_lower for word in ['interactive', 'builder', 'demo', 'example']):
                if any(word in heading_lower for word in ['example', 'demo', 'usage', 'recommendation', 'guide']):
                    relevant_sections.append({
                        'section': section,
                        'similarity_score': 0.7,  # Default similarity score
                        'content': f"HEADING: {heading}"
                    })
        
        return relevant_sections
    
    
    
    def _extract_docx_text_content(self, doc: Document) -> str:
        """Extract structured text content from DOCX for LLM analysis."""
        content_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Include paragraph style information for context
                style_info = f"[{paragraph.style.name}]" if paragraph.style.name else "[Normal]"
                content_parts.append(f"{style_info} {paragraph.text.strip()}")
        
        # Include table content
        for table_idx, table in enumerate(doc.tables):
            content_parts.append(f"\n[TABLE {table_idx + 1}]")
            for row_idx, row in enumerate(table.rows):
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    content_parts.append(f"Row {row_idx + 1}: {row_text}")
        
        return "\n".join(content_parts)

    def _generate_structured_docx_updates(self, doc_content: str, commit_message: str, files_changed: List[str], diff_summary: str, rag_context: str) -> List[Dict[str, Any]]:
        """
        Generate structured updates for DOCX using LLM.
        
        Returns:
            List of update instructions for specific paragraphs/sections
        """
        try:
            prompt = f"""
You are a technical documentation editor. Analyze the DOCX document content and provide specific updates based on the commit changes.

CURRENT DOCUMENT CONTENT:
{doc_content}

COMMIT CHANGES:
- Message: {commit_message}
- Files Changed: {', '.join(files_changed)}
- Summary: {diff_summary}

ADDITIONAL CONTEXT:
{rag_context}

TASK:
1. Identify which sections need updates based on the commit changes
2. For each section that needs updating, provide:
   - Section identifier (exact text from document WITHOUT markdown symbols)
   - Action to take (add|modify|insert_after)
   - New content to add
   - Position relative to the section

IMPORTANT: 
- Use the EXACT text from the document content (without [Heading] prefixes or markdown symbols)
- Respond with ONLY valid JSON. No markdown code blocks, no explanations, no additional text.

RESPONSE FORMAT (JSON only):
{{
  "updates": [
    {{
      "section": "exact text from document to identify the section",
      "action": "add|modify|insert_after",
      "content": "new content to add",
      "position": "after|before|replace"
    }}
  ]
}}
"""

            response = self.provider.generate_response(prompt)
            
            if response and len(response.strip()) > 20:
                # Parse JSON response with robust cleaning
                try:
                    # Clean the response to extract JSON
                    cleaned_response = self._clean_json_response(response)
                    result = json.loads(cleaned_response)
                    return result.get('updates', [])
                except json.JSONDecodeError as e:
                    logger.warning(f"Failed to parse JSON from LLM response: {e}")
                    logger.warning(f"Raw response: {response[:500]}...")
            
            return []
            
        except Exception as e:
            logger.error(f"Error generating structured DOCX updates: {e}")
            return []

    def _clean_json_response(self, response: str) -> str:
        """Clean LLM response to extract valid JSON."""
        try:
            logger.info(f"🔍 Cleaning JSON response: {response[:200]}...")
            
            # Remove markdown code blocks
            if '```json' in response:
                start = response.find('```json') + 7
                end = response.find('```', start)
                if end != -1:
                    response = response[start:end].strip()
                    logger.info(f"🔍 After removing ```json: {response[:200]}...")
            elif '```' in response:
                start = response.find('```') + 3
                end = response.find('```', start)
                if end != -1:
                    response = response[start:end].strip()
                    logger.info(f"🔍 After removing ```: {response[:200]}...")
            
            # Find JSON boundaries (object or array)
            json_start = response.find('{')
            json_end = response.rfind('}') + 1
            
            # Also check for JSON arrays
            array_start = response.find('[')
            array_end = response.rfind(']') + 1
            
            logger.info(f"🔍 JSON boundaries - object: {json_start} to {json_end}, array: {array_start} to {array_end}")
            
            # Prioritize arrays over objects when both are present
            if array_start != -1 and array_end > array_start:
                # JSON array - this is what we want for instructions
                json_str = response[array_start:array_end]
                logger.info(f"🔍 Using JSON array: {json_str[:200]}...")
            elif json_start != -1 and json_end > json_start:
                # JSON object - fallback
                json_str = response[json_start:json_end]
                logger.info(f"🔍 Using JSON object: {json_str[:200]}...")
            else:
                logger.warning(f"🔍 No valid JSON boundaries found, returning original: {response[:200]}...")
                return response.strip()
            
            # Fix common JSON issues
            json_str = json_str.replace('\n', ' ')  # Remove newlines
            json_str = re.sub(r',\s*}', '}', json_str)  # Remove trailing commas
            json_str = re.sub(r',\s*]', ']', json_str)  # Remove trailing commas in arrays
            
            logger.info(f"🔍 Final cleaned JSON: {json_str[:200]}...")
            return json_str.strip()
            
        except Exception as e:
            logger.warning(f"Error cleaning JSON response: {e}")
            return response

    def _apply_updates_to_docx(self, doc: Document, updates: List[Dict[str, Any]]):
        """Apply LLM-generated updates directly to DOCX paragraphs."""
        try:
            for update in updates:
                section_text = update.get('section', '')
                action = update.get('action', 'add')
                content = update.get('content', '')
                position = update.get('position', 'after')
                
                if not section_text or not content:
                    continue
                
                # Find the paragraph that matches the section text using smart detection
                target_paragraph = self._find_best_matching_section(doc, section_text)
                
                if target_paragraph:
                    # Log the original content before making changes
                    original_text = target_paragraph.text.strip()
                    
                    if action == 'add' and position == 'after':
                        # Add new paragraph after the target
                        new_paragraph = target_paragraph._element.getparent().add_paragraph(content)
                        # Copy formatting from target paragraph
                        if target_paragraph.style:
                            new_paragraph.style = target_paragraph.style
                        logger.info(f"✅ ADDED new paragraph after '{original_text[:50]}...'")
                        logger.info(f"   📝 New content: {content[:100]}...")

                    elif action == 'modify' and position == 'replace':
                        # Replace content while preserving formatting
                        original_formats = capture_detailed_formatting(target_paragraph)
                        replace_paragraph_text_with_formatting(target_paragraph, content, original_formats)
                        logger.info(f"✅ MODIFIED paragraph: '{original_text[:50]}...'")
                        logger.info(f"   📝 New content: {content[:100]}...")

                    elif action == 'insert_after':
                        # Insert content after the target paragraph
                        new_paragraph = target_paragraph._element.getparent().add_paragraph(content)
                        if target_paragraph.style:
                            new_paragraph.style = target_paragraph.style
                        logger.info(f"✅ INSERTED new paragraph after '{original_text[:50]}...'")
                        logger.info(f"   📝 New content: {content[:100]}...")

                    logger.info(f"Applied update: {action} {position} to section containing '{section_text[:50]}...'")
                else:
                    logger.warning(f"Could not find section containing: {section_text[:50]}...")
                    
        except Exception as e:
            logger.error(f"Error applying updates to DOCX: {e}")

    def _find_best_matching_section(self, doc: Document, section_text: str):
        """
        Smart section detection using multiple strategies and scoring.
        Returns the best matching paragraph or None if no good match found.
        """
        try:
            # Clean the section text to remove markdown formatting
            clean_section_text = section_text
            if section_text.startswith('#'):
                clean_section_text = re.sub(r'^#+\s*', '', section_text).strip()
            
            logger.info(f"🔍 Looking for section: '{clean_section_text}'")
            
            # Collect all potential matches with scores
            candidates = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                if not para_text:
                    continue
                
                # Calculate similarity score for this paragraph
                score = self._calculate_section_similarity(clean_section_text, para_text)
                
                if score > 0.3:  # Only consider reasonable matches
                    candidates.append({
                        'paragraph': paragraph,
                        'score': score,
                        'text': para_text,
                        'index': i
                    })
            
            # Sort by score (highest first)
            candidates.sort(key=lambda x: x['score'], reverse=True)
            
            if candidates:
                best_match = candidates[0]
                logger.info(f"✅ BEST MATCH found: '{best_match['text'][:60]}...' (score: {best_match['score']:.2f})")
                return best_match['paragraph']
            else:
                logger.warning(f"❌ No suitable match found for: '{clean_section_text}'")
                return None
                
        except Exception as e:
            logger.error(f"Error in smart section detection: {e}")
            return None

    def _calculate_section_similarity(self, target_text: str, paragraph_text: str) -> float:
        """
        Calculate similarity score between target text and paragraph text.
        Returns a score between 0.0 and 1.0, where 1.0 is perfect match.
        """
        try:
            target_lower = target_text.lower()
            para_lower = paragraph_text.lower()
            
            # Strategy 1: Exact match (score: 1.0)
            if target_lower == para_lower:
                return 1.0
            
            # Strategy 2: Target contained in paragraph (score: 0.9)
            if target_lower in para_lower:
                return 0.9
            
            # Strategy 3: Paragraph contained in target (score: 0.8)
            if para_lower in target_lower:
                return 0.8
            
            # Strategy 4: Remove list numbers and try again
            target_clean = re.sub(r'^\d+\.\s*', '', target_lower)
            para_clean = re.sub(r'^\d+\.\s*', '', para_lower)
            
            if target_clean in para_clean:
                return 0.7
            if para_clean in target_clean:
                return 0.6
            
            # Strategy 5: Fuzzy matching (remove common suffixes)
            target_fuzzy = re.sub(r'(s|es|ing|ed)$', '', target_clean)
            para_fuzzy = re.sub(r'(s|es|ing|ed)$', '', para_clean)
            
            if target_fuzzy in para_fuzzy:
                return 0.5
            if para_fuzzy in target_fuzzy:
                return 0.4
            
            # Strategy 6: Word overlap scoring
            target_words = set(target_fuzzy.split())
            para_words = set(para_fuzzy.split())
            
            if target_words and para_words:
                overlap = len(target_words.intersection(para_words))
                total_words = len(target_words.union(para_words))
                word_score = overlap / total_words if total_words > 0 else 0
                
                if word_score > 0.3:  # At least 30% word overlap
                    return word_score * 0.3  # Scale down to 0.3 max
            
            return 0.0
            
        except Exception as e:
            logger.error(f"Error calculating similarity: {e}")
            return 0.0

    def _parse_and_apply_instruction_to_docx(self, doc: Document, instruction: str, commit_message: str, files_changed: List[str], diff_summary: str, rag_context: str) -> int:
        """
        Parse the instruction and apply updates directly to the DOCX document.
        Returns the number of updates applied.
        """
        try:
            updates_applied = 0
            
            # Parse the instruction to understand what needs to be updated
            parsed_actions = self._parse_instruction(instruction)
            
            logger.info(f"🔍 Parsed {len(parsed_actions)} actions from instruction")
            
            for action in parsed_actions:
                logger.info(f"🎯 Processing action: {action}")
                
                if action['type'] == 'add_bullet_point':
                    success = self._add_bullet_point_to_section(doc, action['section'], action['content'])
                    if success:
                        updates_applied += 1
                        logger.info(f"✅ Added bullet point to '{action['section']}'")
                    else:
                        logger.warning(f"❌ Failed to add bullet point to '{action['section']}'")
                
                elif action['type'] == 'update_table':
                    success = self._add_row_to_table(doc, action['table_name'], action['row_data'])
                    if success:
                        updates_applied += 1
                        logger.info(f"✅ Added row to table '{action['table_name']}'")
                    else:
                        logger.warning(f"❌ Failed to add row to table '{action['table_name']}'")
                
                elif action['type'] == 'add_paragraph':
                    success = self._add_paragraph_after_section(doc, action['section'], action['content'])
                    if success:
                        updates_applied += 1
                        logger.info(f"✅ Added paragraph after '{action['section']}'")
                    else:
                        logger.warning(f"❌ Failed to add paragraph after '{action['section']}'")
                
                elif action['type'] == 'insert_content':
                    success = self._insert_content_in_section(doc, action['section'], action['content'], action.get('position', 'after'))
                    if success:
                        updates_applied += 1
                        logger.info(f"✅ Inserted content in '{action['section']}'")
                    else:
                        logger.warning(f"❌ Failed to insert content in '{action['section']}'")
            
            return updates_applied
            
        except Exception as e:
            logger.error(f"Error parsing and applying instruction to DOCX: {e}")
            return 0

    def _parse_instruction(self, instruction: str) -> List[Dict[str, Any]]:
        """
        Parse the instruction string to extract actionable items.
        """
        try:
            actions = []
            
            # Use LLM to parse the instruction into structured actions
            prompt = f"""
You are a documentation parser. Parse this instruction into specific actionable items.

INSTRUCTION:
{instruction}

TASK:
Convert this instruction into a JSON array of specific actions that can be performed on a DOCX document.

ACTION TYPES:
- add_bullet_point: Add a bullet point to a specific section
- update_table: Add a row to a specific table
- add_paragraph: Add a paragraph after a specific section
- insert_content: Insert content in a specific section

RESPONSE FORMAT (JSON only):
[
  {{
    "type": "add_bullet_point",
    "section": "Feature Specifications",
    "content": "Factorial Operation: Calculates the factorial of a number"
  }},
  {{
    "type": "update_table",
    "table_name": "Calculator Operations",
    "row_data": {{
      "Operation": "Factorial",
      "Method Name": "factorial",
      "Input Parameters": "n (integer)",
      "Output": "factorial of n",
      "Error Handling": "Raises ValueError for negative numbers"
    }}
  }}
]

IMPORTANT: 
- Respond with ONLY valid JSON
- Use exact section names from the instruction
- Be specific about what content to add
- No explanations, no markdown code blocks
- CRITICAL: row_data MUST be an object with key-value pairs, NOT an array
- Each key in row_data should match a table column header exactly
"""

            response = self.provider.generate_response(prompt)
            
            if response and len(response.strip()) > 10:
                try:
                    cleaned_response = self._clean_json_response(response)
                    actions = json.loads(cleaned_response)
                    logger.info(f"Successfully parsed {len(actions)} actions from instruction")
                except json.JSONDecodeError as e:
                    logger.warning(f"Failed to parse JSON from instruction: {e}")
                    logger.warning(f"Raw response: {response[:500]}...")
                
                # Save debug files
                import os
                debug_dir = "debug_json"
                os.makedirs(debug_dir, exist_ok=True)
                
                # Save raw response
                with open(f"{debug_dir}/raw_instruction_response.txt", "w", encoding="utf-8") as f:
                    f.write(response)
                
                # Save cleaned response
                with open(f"{debug_dir}/cleaned_instruction_response.txt", "w", encoding="utf-8") as f:
                    f.write(cleaned_response)
                
                logger.info(f"💾 Saved debug files to {debug_dir}/")
            
            return actions
            
        except Exception as e:
            logger.error(f"Error parsing instruction: {e}")
            return []

    def _add_bullet_point_to_section(self, doc: Document, section_name: str, content: str) -> bool:
        """Add a bullet point to a specific section."""
        try:
            # Find the section
            target_paragraph = self._find_section_paragraph(doc, section_name)
            if not target_paragraph:
                logger.warning(f"Could not find section '{section_name}' for bullet point")
                return False
            
            # Simply add a new paragraph with bullet point content
            # This is much simpler and more reliable than XML manipulation
            new_paragraph = doc.add_paragraph(f"• {content}")
            
            logger.info(f"✅ Added bullet point after section '{section_name}': {content[:50]}...")
            return True
            
        except Exception as e:
            logger.error(f"Error adding bullet point to section '{section_name}': {e}")
            return False

    def _add_row_to_table(self, doc: Document, table_name: str, row_data: Dict[str, str]) -> bool:
        """Add a row to a specific table."""
        try:
            # Find the table by name (look for nearby text)
            target_table = self._find_table_by_name(doc, table_name)
            if not target_table:
                return False
            
            # Add new row
            new_row = target_table.add_row()
            
            # Fill the row with data
            for i, (key, value) in enumerate(row_data.items()):
                if i < len(new_row.cells):
                    new_row.cells[i].text = value
            
            return True
            
        except Exception as e:
            logger.error(f"Error adding row to table '{table_name}': {e}")
            return False

    def _add_paragraph_after_section(self, doc: Document, section_name: str, content: str) -> bool:
        """Add a paragraph after a specific section."""
        try:
            # Find the section
            target_paragraph = self._find_section_paragraph(doc, section_name)
            if not target_paragraph:
                logger.warning(f"Could not find section '{section_name}' for paragraph")
                return False
            
            # Simply add a new paragraph with content
            new_paragraph = doc.add_paragraph(content)
            
            logger.info(f"✅ Added paragraph after section '{section_name}': {content[:50]}...")
            return True
            
        except Exception as e:
            logger.error(f"Error adding paragraph after section '{section_name}': {e}")
            return False

    def _insert_content_in_section(self, doc: Document, section_name: str, content: str, position: str = 'after') -> bool:
        """Insert content in a specific section."""
        try:
            # Find the section
            target_paragraph = self._find_section_paragraph(doc, section_name)
            if not target_paragraph:
                return False
            
            if position == 'after':
                return self._add_paragraph_after_section(doc, section_name, content)
            elif position == 'before':
                # Insert before the section
                new_paragraph = target_paragraph._element.getparent().add_paragraph(content)
                target_paragraph._element.getparent().insert_before(new_paragraph._element, target_paragraph._element)
                return True
            else:
                # Append to existing paragraph
                target_paragraph.text += f"\n{content}"
                return True
            
        except Exception as e:
            logger.error(f"Error inserting content in section '{section_name}': {e}")
            return False

    def _find_section_paragraph(self, doc: Document, section_name: str) -> Optional[Any]:
        """Find a paragraph that contains the section name."""
        try:
            clean_section_name = section_name.lower().strip()
            
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip().lower()
                if not para_text:
                    continue
                
                # Try multiple matching strategies
                if clean_section_name in para_text:
                    return paragraph
                
                # Remove list numbers and try again
                clean_para = re.sub(r'^\d+\.\s*', '', para_text)
                if clean_section_name in clean_para:
                    return paragraph
                
                # Try fuzzy matching
                if self._fuzzy_match_section(clean_section_name, para_text):
                    return paragraph
            
            return None
            
        except Exception as e:
            logger.error(f"Error finding section paragraph '{section_name}': {e}")
            return None

    def _find_table_by_name(self, doc: Document, table_name: str) -> Optional[Any]:
        """Find a table by looking for nearby text that mentions the table name."""
        try:
            clean_table_name = table_name.lower().strip()
            
            for table in doc.tables:
                # Check paragraphs before and after the table
                table_element = table._element
                parent = table_element.getparent()
                
                # Look at surrounding paragraphs
                for paragraph in doc.paragraphs:
                    para_element = paragraph._element
                    if para_element.getparent() == parent:
                        para_text = paragraph.text.strip().lower()
                        if clean_table_name in para_text:
                            return table
            
            # Fallback: return first table if no specific match
            if doc.tables:
                return doc.tables[0]
            
            return None
            
        except Exception as e:
            logger.error(f"Error finding table '{table_name}': {e}")
            return None

    def _fuzzy_match_section(self, target: str, candidate: str) -> bool:
        """Perform fuzzy matching between section names."""
        try:
            # Remove common words and try matching
            target_clean = re.sub(r'\b(the|a|an|and|or|but|in|on|at|to|for|of|with|by)\b', '', target)
            candidate_clean = re.sub(r'\b(the|a|an|and|or|but|in|on|at|to|for|of|with|by)\b', '', candidate)
            
            # Check if one contains the other
            return target_clean in candidate_clean or candidate_clean in target_clean
            
        except Exception as e:
            logger.error(f"Error in fuzzy matching: {e}")
            return False

    def _generate_update_instruction(self, doc_content: str, commit_message: str, files_changed: List[str], diff_summary: str, rag_context: str) -> str:
        """
        Generate a specific update instruction for the document.
        Returns a clear instruction like: "MODIFY_APPEND: 'Feature Specifications' -> 'Factorial Operation: Calculates...'"
        """
        try:
            prompt = f"""
You are a technical documentation editor. Analyze the document and commit changes to generate a specific update instruction.

CURRENT DOCUMENT CONTENT:
{doc_content}

COMMIT CHANGES:
- Message: {commit_message}
- Files Changed: {', '.join(files_changed)}
- Summary: {diff_summary}

ADDITIONAL CONTEXT:
{rag_context}

TASK:
Generate a clear, specific instruction that tells exactly what needs to be changed in the document.

INSTRUCTION REQUIREMENTS:
- Be clear and unambiguous about what action to take
- Specify exactly which section/content to modify
- Describe what new content to add or how to modify existing content
- Use natural language that another system can understand

EXAMPLES OF GOOD INSTRUCTIONS:
- "Add a new bullet point about factorial operation after the 'Feature Specifications' section"
- "Insert a new paragraph describing the factorial function before the 'Technical Implementation' section"
- "Update the 'Calculator Operations' table to include factorial as a new row"
- "Add factorial operation documentation to the features list"

IMPORTANT: 
- Use the EXACT section names from the document
- Be specific about what content to add
- Respond with ONLY the instruction, no explanations

INSTRUCTION:
"""

            response = self.provider.generate_response(prompt)
            
            if response and len(response.strip()) > 10:
                # Clean the response to extract just the instruction
                instruction = response.strip()
                if '\n' in instruction:
                    instruction = instruction.split('\n')[0].strip()
                
                logger.info(f"Generated instruction: {instruction}")
                return instruction
            
            return ""

        except Exception as e:
            logger.error(f"Error generating update instruction: {e}")
            return ""

    def _llm_update_document_with_instruction(self, doc_content: str, instruction: str, commit_message: str, files_changed: List[str], diff_summary: str, rag_context: str) -> str:
        """
        Pass the instruction and full document to LLM for direct update.
        Returns the complete updated document content.
        """
        try:
            prompt = f"""
You are a technical documentation editor. You have been given a specific instruction to update a document.

ORIGINAL DOCUMENT CONTENT:
{doc_content}

UPDATE INSTRUCTION:
{instruction}

COMMIT CONTEXT:
- Message: {commit_message}
- Files Changed: {', '.join(files_changed)}
- Summary: {diff_summary}

ADDITIONAL CONTEXT:
{rag_context}

TASK:
Apply the update instruction to the document and return the COMPLETE updated document content.

INSTRUCTIONS:
1. Carefully read and understand the update instruction
2. Find the relevant section or content mentioned in the instruction
3. Apply the requested changes while maintaining document structure and formatting
4. Return the ENTIRE updated document content
5. Do not add explanations or comments - just the updated document

IMPORTANT:
- Preserve the original document structure and formatting
- Only make the changes specified in the instruction
- Ensure the updated content flows naturally with the existing text

UPDATED DOCUMENT:
"""

            response = self.provider.generate_response(prompt)
            
            if response and len(response.strip()) > len(doc_content.strip()):
                logger.info(f"LLM generated updated content: {len(response)} characters")
                return response.strip()
            
            logger.warning("LLM response too short or empty")
            return ""

        except Exception as e:
            logger.error(f"Error updating document with LLM: {e}")
            return ""

    def _convert_text_to_docx(self, text_content: str, original_doc: Document) -> Document:
        """
        Convert the updated text content back to a DOCX Document object.
        This is a simplified approach - in production you'd want more sophisticated conversion.
        """
        try:
            # For now, create a new document and add the content as paragraphs
            new_doc = Document()
            
            # Split content into paragraphs and add them
            paragraphs = text_content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    new_doc.add_paragraph(para_text.strip())
            
            logger.info(f"Converted text to DOCX: {len(new_doc.paragraphs)} paragraphs")
            return new_doc
            
        except Exception as e:
            logger.error(f"Error converting text to DOCX: {e}")
            return original_doc

    def _log_document_structure(self, doc: Document, stage: str):
        """Log the structure of the document for debugging purposes."""
        try:
            logger.info(f"🔍 {stage} - Document Analysis:")
            logger.info(f"   📊 Total paragraphs: {len(doc.paragraphs)}")
            logger.info(f"   📊 Total tables: {len(doc.tables)}")
            
            # Log first few paragraphs with their styles
            logger.info(f"   📝 First 10 paragraphs:")
            for i, para in enumerate(doc.paragraphs[:10]):
                if para.text.strip():
                    style_name = para.style.name if para.style else "Normal"
                    text_preview = para.text.strip()[:80]
                    logger.info(f"      {i+1:2d}. [{style_name}] {text_preview}...")
            
            # Log table information
            if doc.tables:
                logger.info(f"   📊 Tables:")
                for i, table in enumerate(doc.tables):
                    rows = len(table.rows)
                    cols = len(table.columns) if table.rows else 0
                    logger.info(f"      Table {i+1}: {rows} rows × {cols} columns")
                    
                    # Show first row (headers) if available
                    if table.rows:
                        headers = [cell.text.strip() for cell in table.rows[0].cells if cell.text.strip()]
                        if headers:
                            logger.info(f"         Headers: {' | '.join(headers[:5])}{'...' if len(headers) > 5 else ''}")
            
        except Exception as e:
            logger.error(f"Error logging document structure: {e}")

    def _convert_markdown_to_docx(self, markdown_content: str, original_doc: Document) -> Document:
        """Convert markdown content back to DOCX format, preserving original formatting."""
        # For now, return the original document with updated content
        # This is a simplified approach - in production, you'd want more sophisticated conversion
        return original_doc

    def _generate_content_for_heading(self, heading: str, commit_message: str, files_changed: List[str], diff_summary: str, heading_content: str = "") -> str:
        """
        Generate content for a heading based on the actual commit changes, heading context, and diff summary.
        
        Args:
            heading: The heading/section name
            commit_message: The commit message
            files_changed: List of files that were changed
            diff_summary: Summary of the actual changes made
            heading_content: Content under this heading for context
            
        Returns:
            Generated content relevant to the heading and actual changes
        """
        try:
            # Use LLM to analyze the actual changes and generate relevant content
            prompt = f"""
You are a technical documentation writer. Analyze the existing content structure and add relevant information about the commit changes while maintaining the same style and flow.

COMMIT INFORMATION:
- Message: {commit_message}
- Files Changed: {', '.join(files_changed)}
- Changes Summary: {diff_summary}

DOCUMENTATION SECTION:
- Heading: {heading}
- Existing Content: {heading_content[:800] if heading_content else "No existing content"}

TASK:
Analyze the existing content structure, style, and flow. Then add information about the commit changes that:
1. Follows the same writing style and tone as existing content
2. Maintains the same structure and formatting patterns
3. Integrates naturally with the existing flow
4. Adds relevant information about the specific changes made
5. Doesn't disrupt the existing content organization

IMPORTANT:
- Study the existing content structure carefully
- Match the existing writing style (formal/informal, technical level, sentence structure)
- Follow the same formatting patterns (bullet points, paragraphs, lists, etc.)
- Integrate the new information seamlessly into the existing flow
- Don't add redundant information already present
"""

            response = self.provider.generate_response(prompt)
            
            if response and len(response.strip()) > 20:
                return response.strip()
            else:
                # Fallback to basic analysis if LLM fails
                return self._generate_fallback_content(heading, commit_message, files_changed, diff_summary)
                
        except Exception as e:
            logger.warning(f"Failed to generate content with LLM: {e}")
            return self._generate_fallback_content(heading, commit_message, files_changed, diff_summary)
    
    def _generate_fallback_content(self, heading: str, commit_message: str, files_changed: List[str], diff_summary: str) -> str:
        """Fallback content generation when LLM fails."""
        heading_lower = heading.lower()
        commit_lower = commit_message.lower()
        
        # Extract key information from commit dynamically
        feature_desc = "new feature"
        
        # Analyze commit message for key terms
        if "add" in commit_lower:
            if "operation" in commit_lower:
                feature_desc = "new operation"
            elif "function" in commit_lower:
                feature_desc = "new function"
            elif "feature" in commit_lower:
                feature_desc = "new feature"
            else:
                feature_desc = "new functionality"
        
        # Analyze diff summary for more specific information
        if diff_summary:
            diff_lower = diff_summary.lower()
            # Extract key terms from diff summary dynamically
            if "operation" in diff_lower:
                feature_desc = "mathematical operation"
            elif "function" in diff_lower:
                feature_desc = "new function"
            elif "menu" in diff_lower:
                feature_desc = "menu option"
            elif "interface" in diff_lower:
                feature_desc = "interface update"
            elif "api" in diff_lower:
                feature_desc = "API enhancement"
            elif "config" in diff_lower:
                feature_desc = "configuration update"
        
        # Generate structured content based on heading type and existing content
        if any(word in heading_lower for word in ['feature', 'specification']):
            return f"The application now includes a {feature_desc} that enhances its capabilities."
        
        elif any(word in heading_lower for word in ['mathematical', 'operation', 'arithmetic']):
            return f"A new {feature_desc} has been added to expand the mathematical capabilities."
        
        elif any(word in heading_lower for word in ['implementation', 'technical']):
            return f"The {feature_desc} has been implemented with proper error handling and validation."
        
        elif any(word in heading_lower for word in ['testing', 'quality']):
            return f"The new {feature_desc} includes comprehensive validation and error handling."
        
        elif any(word in heading_lower for word in ['user', 'interface', 'menu']):
            return f"The {feature_desc} has been integrated into the user interface for easy access."
        
        else:
            return f"The application has been enhanced with a {feature_desc}."
    
    def _generate_minimal_addition(self, heading: str, commit_message: str, files_changed: List[str], diff_summary: str) -> str:
        """Generate minimal content when existing content already mentions the change."""
        heading_lower = heading.lower()
        commit_lower = commit_message.lower()
        
        # Extract key information from commit dynamically
        feature_desc = "new feature"
        
        # Analyze commit message for key terms
        if "add" in commit_lower:
            if "operation" in commit_lower:
                feature_desc = "operation"
            elif "function" in commit_lower:
                feature_desc = "function"
            elif "feature" in commit_lower:
                feature_desc = "feature"
            else:
                feature_desc = "functionality"
        
        # Analyze diff summary for more specific information
        if diff_summary:
            diff_lower = diff_summary.lower()
            if "operation" in diff_lower:
                feature_desc = "operation"
            elif "function" in diff_lower:
                feature_desc = "function"
            elif "menu" in diff_lower:
                feature_desc = "menu option"
            elif "interface" in diff_lower:
                feature_desc = "interface update"
            elif "api" in diff_lower:
                feature_desc = "API enhancement"
            elif "config" in diff_lower:
                feature_desc = "configuration update"
        
        # Generate minimal content based on heading type
        if any(word in heading_lower for word in ['feature', 'specification']):
            return f"Added {feature_desc}."
        elif any(word in heading_lower for word in ['mathematical', 'operation', 'arithmetic']):
            return f"New {feature_desc}."
        elif any(word in heading_lower for word in ['implementation', 'technical']):
            return f"Implemented {feature_desc}."
        elif any(word in heading_lower for word in ['testing', 'quality']):
            return f"Added {feature_desc}."
        elif any(word in heading_lower for word in ['user', 'interface', 'menu']):
            return f"Added {feature_desc}."
        else:
            return f"Added {feature_desc}."
    
    def _select_best_paragraph_with_llm(self, paragraphs: List[Dict], heading: str, commit_message: str, files_changed: List[str], diff_summary: str) -> Optional[Dict]:
        """Use LLM to select the best paragraph for content insertion."""
        try:
            commit_lower = commit_message.lower()
            heading_lower = heading.lower()
            
            # Score each paragraph based on relevance
            scored_paragraphs = []
            
            for para_info in paragraphs:
                text = para_info['text']
                text_lower = text.lower()
                
                score = 0
                reasons = []
                
                # Check for content type matches - analyze commit keywords dynamically
                commit_keywords = [word for word in commit_lower.split() if len(word) > 3]
                text_keywords = [word for word in text_lower.split() if len(word) > 3]
                
                # Score based on keyword overlap
                keyword_overlap = len(set(commit_keywords) & set(text_keywords))
                if keyword_overlap > 0:
                    score += min(keyword_overlap, 3)
                    reasons.append(f"Content matches commit keywords ({keyword_overlap} matches)")
                
                # Check for similar context based on heading
                if any(word in heading_lower for word in ['recommendation', 'usage', 'example', 'guide']):
                    if any(word in text_lower for word in ['recommend', 'suggest', 'use', 'develop', 'example', 'guide']):
                        score += 2
                        reasons.append("Content aligns with heading theme")
                
                # Check for feature-related content
                if any(word in commit_lower for word in ['feature', 'new', 'add', 'implement']):
                    if any(word in text_lower for word in ['feature', 'capability', 'functionality', 'enhancement']):
                        score += 2
                        reasons.append("Content relates to feature development")
                
                # Prefer paragraphs that are descriptive but not too long
                length_score = min(len(text) / 100, 3)  # Cap at 3 points
                score += length_score
                reasons.append(f"Length score: {len(text)} characters")
                
                # Avoid paragraphs that are too generic or too specific
                if any(word in text_lower for word in ['document', 'analysis', 'report', 'generated']):
                    score -= 1
                    reasons.append("Penalty for generic document language")
                
                # Prefer paragraphs that are more specific and actionable
                if any(word in text_lower for word in ['recommend', 'suggest', 'use', 'should', 'can', 'will']):
                    score += 1
                    reasons.append("Content contains actionable language")
                
                scored_paragraphs.append({
                    **para_info,
                    'score': score,
                    'reasons': reasons
                })
            
            # Select the paragraph with the highest score
            if not scored_paragraphs:
                return None
                
            best_para = max(scored_paragraphs, key=lambda x: x['score'])
            
            # Add reasoning to the result
            best_para['reasoning'] = '; '.join(best_para['reasons'])
            
            return best_para
                
        except Exception as e:
            logger.error(f"Error in LLM paragraph selection: {e}")
            # Fallback to longest paragraph
            return max(paragraphs, key=lambda x: len(x['text']))
    
    def _find_best_insertion_point(self, paragraphs: List[Dict], best_para_info: Dict, heading: str, commit_message: str) -> int:
        """Find the best insertion point for the new content within the section."""
        try:
            best_para = best_para_info['paragraph_obj']
            parent = best_para._element.getparent()
            best_para_index = parent.index(best_para._element)
            
            # Default: insert right after the selected paragraph
            insertion_index = best_para_index + 1
            
            # If this is a recommendation/usage section, try to insert at a more logical position
            heading_lower = heading.lower()
            commit_lower = commit_message.lower()
            
            if any(word in heading_lower for word in ['recommendation', 'usage', 'example']):
                # For recommendation sections, try to insert after introductory content
                # Look for paragraphs that seem to be introductory vs specific recommendations
                for i, para_info in enumerate(paragraphs):
                    text_lower = para_info['text'].lower()
                    
                    # If we find a paragraph that seems to be setting up recommendations
                    if any(word in text_lower for word in ['recommend', 'suggest', 'should', 'can', 'will']):
                        # Insert after this paragraph instead
                        para_obj = para_info['paragraph_obj']
                        para_index = parent.index(para_obj._element)
                        insertion_index = para_index + 1
                        logger.info(f"Found better insertion point after recommendation paragraph at index {para_index}")
                        break
            
            logger.info(f"Using insertion index: {insertion_index}")
            return insertion_index
            
        except Exception as e:
            logger.error(f"Error finding insertion point: {e}")
            # Fallback to inserting after the selected paragraph
            best_para = best_para_info['paragraph_obj']
            parent = best_para._element.getparent()
            return parent.index(best_para._element) + 1
    
    def _parse_json_safely(self, response: str) -> Optional[Any]:
        """
        Safely parse JSON from LLM response with multiple fallback strategies.
        
        Args:
            response: Raw LLM response
            
        Returns:
            Parsed JSON object or None if parsing fails
        """
        if not response:
            return None
        
        # Strategy 1: Try direct parsing
        try:
            return json.loads(response.strip())
        except json.JSONDecodeError:
            pass
        
        # Strategy 2: Extract JSON from markdown code blocks
        try:
            # Look for ```json blocks
            json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', response, re.DOTALL)
            if json_match:
                return json.loads(json_match.group(1))
        except json.JSONDecodeError:
            pass
        
        # Strategy 3: Find JSON object boundaries
        try:
            start_idx = response.find('{')
            end_idx = response.rfind('}')
            if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
                json_str = response[start_idx:end_idx+1]
                return json.loads(json_str)
        except json.JSONDecodeError:
            pass
        
        # Strategy 4: Clean and try again
        try:
            cleaned = self._clean_json_string(response)
            return json.loads(cleaned)
        except json.JSONDecodeError:
            pass
        
        logger.warning(f"Failed to parse JSON from response: {response[:200]}...")
        return None
    
    def _clean_json_string(self, json_str: str) -> str:
        """
        Clean JSON string by fixing common issues.
        
        Args:
            json_str: Raw JSON string
            
        Returns:
            Cleaned JSON string
        """
        # Remove markdown formatting
        cleaned = re.sub(r'```(?:json)?\s*', '', json_str)
        cleaned = re.sub(r'```\s*', '', cleaned)
        
        # Remove trailing commas
        cleaned = re.sub(r',(\s*[}\]])', r'\1', cleaned)
        
        # Normalize boolean values
        cleaned = re.sub(r'\bTrue\b', 'true', cleaned)
        cleaned = re.sub(r'\bFalse\b', 'false', cleaned)
        cleaned = re.sub(r'\bNull\b', 'null', cleaned)
        
        # Remove control characters
        cleaned = ''.join(char for char in cleaned if ord(char) >= 32 or char in '\n\t\r')
        
        return cleaned.strip()
    
    def _parse_document_sections(self, doc_content: str) -> List[Dict[str, Any]]:
        """
        Parse document content into sections.
        
        Args:
            doc_content: Document content
            
        Returns:
            List of section dictionaries
        """
        try:
            sections = []
            lines = doc_content.split('\n')
            current_section = None
            
            for i, line in enumerate(lines):
                line = line.strip()
                
                # Check for headings (simple heuristic)
                if line and (line.isupper() or line.startswith('#') or 
                           (len(line) < 100 and not line.startswith('-') and not line.startswith('*'))):
                    # Save previous section
                    if current_section:
                        sections.append(current_section)
                    
                    # Start new section
                    current_section = {
                        'heading': line,
                        'line_start': i + 1,
                        'line_end': i + 1,
                        'content_lines': [line],
                        'word_count': len(line.split())
                    }
                elif current_section:
                    current_section['content_lines'].append(line)
                    current_section['line_end'] = i + 1
                    current_section['word_count'] += len(line.split())
            
            # Add final section
            if current_section:
                sections.append(current_section)
            
            return sections
            
        except Exception as e:
            logger.error(f"Error parsing document sections: {e}")
            return []
    
    def _format_sections_for_analysis(self, sections: List[Dict[str, Any]]) -> str:
        """
        Format sections for LLM analysis.
        
        Args:
            sections: List of section dictionaries
            
        Returns:
            Formatted string
        """
        try:
            formatted = []
            for section in sections:
                formatted.append(f"Section: {section['heading']}")
                formatted.append(f"Lines: {section['line_start']}-{section['line_end']}")
                formatted.append(f"Content: {' '.join(section['content_lines'][:3])}...")
                formatted.append("")
            
            return '\n'.join(formatted)
                
        except Exception as e:
            logger.error(f"Error formatting sections: {e}")
            return "Error formatting sections"
    
    def _create_fallback_commit_analysis(self, commit_message: str, files_changed: List[str]) -> Dict[str, Any]:
        """Create fallback commit analysis."""
        commit_lower = commit_message.lower()
        
        if any(keyword in commit_lower for keyword in ['feat', 'feature', 'new', 'add']):
            impact_level = "high"
        elif any(keyword in commit_lower for keyword in ['fix', 'bug', 'update']):
            impact_level = "medium"
        else:
            impact_level = "low"
        
        return {
            "impact_level": impact_level,
            "key_changes": [f"Changes to {', '.join(files_changed[:3])}"],
            "reasoning": f"Fallback analysis based on commit message: {commit_message}",
            "recommended_docs": [
                {
                    "doc_type": "user_guide",
                    "priority": impact_level,
                    "reason": "General documentation update needed"
                }
            ]
        }
    
    def _create_fallback_document_analysis(self, commit_message: str, sections: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Create fallback document analysis."""
        commit_lower = commit_message.lower()
        
        # Find best section based on keywords
        best_section = "Document Content"
        for section in sections:
            heading_lower = section['heading'].lower()
            if any(keyword in heading_lower for keyword in ['feature', 'function', 'api', 'usage']):
                best_section = section['heading']
                break
        
        return {
            "section_analysis": [
                {
                    "heading": best_section,
                    "relevance_score": 7,
                    "should_update": True,
                    "update_reason": f"Fallback analysis: Commit appears relevant to {best_section}",
                    "content_type": "paragraph"
                }
            ],
            "placement_decision": {
                "target_section": best_section,
                "placement_type": "append",
                "content_type": "description",
                "reasoning": "Fallback placement decision"
            }
        }
    
    def _extract_table_content(self, table) -> List[List[str]]:
        """
        Extract all content from a table as a list of rows.
        
        Args:
            table: Table object from python-docx
            
        Returns:
            List of lists representing table rows and cells
        """
        table_data = []
        
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell_idx, cell in enumerate(row.cells):
                # Get all text from cell paragraphs
                cell_text = ' '.join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                row_data.append(cell_text)
            table_data.append(row_data)
        
        logger.info(f"Extracted table with {len(table_data)} rows and {len(table_data[0]) if table_data else 0} columns")
        return table_data
    
    def _convert_table_to_json(self, table_data: List[List[str]]) -> Dict[str, Any]:
        """
        Convert table data to structured JSON format for better analysis.
        
        Args:
            table_data: Table content as list of rows
            
        Returns:
            Structured JSON representation of the table
        """
        if not table_data:
            return {"headers": [], "rows": [], "column_count": 0, "row_count": 0}
        
        # First row is typically headers
        headers = table_data[0] if table_data else []
        data_rows = table_data[1:] if len(table_data) > 1 else []
        
        # Convert rows to objects with header keys
        structured_rows = []
        for row in data_rows:
            row_dict = {}
            for i, cell_value in enumerate(row):
                header = headers[i] if i < len(headers) else f"Column_{i+1}"
                row_dict[header] = cell_value
            structured_rows.append(row_dict)
        
        return {
            "headers": headers,
            "rows": structured_rows,
            "column_count": len(headers),
            "row_count": len(data_rows),
            "raw_data": table_data
        }
    
    def _analyze_table_structure(self, table_json: Dict[str, Any]) -> Dict[str, Any]:
        """
        Analyze table structure to understand column purposes and data patterns.
        
        Args:
            table_json: Structured JSON representation of the table
            
        Returns:
            Analysis of table structure and patterns
        """
        headers = table_json.get("headers", [])
        rows = table_json.get("rows", [])
        
        # Analyze column types and purposes
        column_analysis = {}
        for i, header in enumerate(headers):
            column_values = [row.get(header, "") for row in rows if header in row]
            
            # Determine column type based on content patterns
            column_type = "text"
            if any(val.isdigit() for val in column_values if val):
                column_type = "numeric"
            elif any(val.lower() in ["true", "false", "yes", "no"] for val in column_values if val):
                column_type = "boolean"
            elif any(val.lower() in ["high", "medium", "low", "critical", "minor"] for val in column_values if val):
                column_type = "priority"
            
            column_analysis[header] = {
                "type": column_type,
                "sample_values": column_values[:3],  # First 3 values as samples
                "unique_values": len(set(column_values)),
                "empty_count": sum(1 for val in column_values if not val.strip())
            }
        
        # Analyze row patterns
        row_patterns = {
            "total_rows": len(rows),
            "average_cell_length": sum(
                sum(len(str(val)) for val in row.values()) 
                for row in rows
            ) / max(len(rows), 1) if rows else 0,
            "common_patterns": []
        }
        
        return {
            "column_analysis": column_analysis,
            "row_patterns": row_patterns,
            "table_type": self._determine_table_type(headers, rows)
        }
    
    def _determine_table_type(self, headers: List[str], rows: List[Dict]) -> str:
        """
        Determine the type/purpose of the table based on headers and content.
        
        Args:
            headers: Table column headers
            rows: Table data rows
            
        Returns:
            String describing the table type
        """
        header_text = " ".join(headers).lower()
        
        # Common table type patterns
        if any(word in header_text for word in ["function", "method", "api", "command"]):
            return "function_reference"
        elif any(word in header_text for word in ["package", "library", "dependency", "module"]):
            return "package_list"
        elif any(word in header_text for word in ["test", "coverage", "status", "quality"]):
            return "test_coverage"
        elif any(word in header_text for word in ["example", "demo", "usage", "sample"]):
            return "examples"
        elif any(word in header_text for word in ["performance", "metric", "benchmark", "speed"]):
            return "performance_metrics"
        elif any(word in header_text for word in ["feature", "enhancement", "planned", "roadmap"]):
            return "feature_roadmap"
        elif any(word in header_text for word in ["download", "star", "popularity", "usage"]):
            return "popularity_stats"
        else:
            return "general_data"
    
    def _fix_json_syntax(self, json_str: str) -> str:
        """
        Fix common JSON syntax issues that LLMs sometimes generate.
        
        Args:
            json_str: Potentially malformed JSON string
            
        Returns:
            Fixed JSON string
        """
        import re
        
        # Fix common issues
        fixed = json_str
        
        # Remove comments (// style)
        fixed = re.sub(r'//.*$', '', fixed, flags=re.MULTILINE)
        
        # Remove trailing commas before closing braces/brackets
        fixed = re.sub(r',(\s*[}\]])', r'\1', fixed)
        
        # Fix single quotes to double quotes (but be careful with apostrophes in strings)
        # This is tricky, so we'll be conservative
        fixed = re.sub(r"'([^']*)':", r'"\1":', fixed)  # Keys
        fixed = re.sub(r":\s*'([^']*)'", r': "\1"', fixed)  # String values
        
        # Fix unescaped quotes in strings (basic fix)
        # This is complex, so we'll just log if we detect potential issues
        if "'" in fixed and '"' in fixed:
            logger.warning("Mixed quotes detected in JSON, may need manual review")
        
        return fixed
    
    def _analyze_table_with_llm(self, table_data: List[List[str]], heading: str, heading_context: str, commit_context: Dict[str, Any]) -> Dict[str, Any]:
        """
        Simplified table analysis - understand headers, check row fit, and conditionally add rows.
        
        Args:
            table_data: Table content as list of rows
            heading: The heading that contains this table
            heading_context: Context paragraph under the heading
            commit_context: Commit information and changes
            
        Returns:
            Dictionary with analysis results and update recommendations
        """
        try:
            if not table_data or len(table_data) < 2:
                logger.info(f"Table too small or empty - skipping analysis")
                return {
                    "table_purpose": f"Table with {len(table_data)} rows (too small for analysis)",
                    "needs_update": False,
                    "update_reason": "Insufficient data for meaningful analysis",
                    "recommended_updates": [],
                    "confidence": 0.0
                }
            
            # Step 1: Understand the table headers
            headers = table_data[0] if table_data else []
            data_rows = table_data[1:] if len(table_data) > 1 else []
            
            logger.info(f"📊 Table Analysis - Headers: {headers}")
            logger.info(f"📊 Table Analysis - Data rows: {len(data_rows)}")
            
            # Step 2: Analyze existing row patterns
            row_patterns = []
            for i, row in enumerate(data_rows[:5]):  # Analyze first 5 rows
                if len(row) == len(headers):
                    row_dict = {headers[j]: row[j] for j in range(len(headers))}
                    row_patterns.append(row_dict)
                    logger.info(f"📊 Row {i+1} pattern: {row_dict}")
                else:
                    logger.warning(f"⚠️ Row {i+1} has {len(row)} columns but table has {len(headers)} headers")
            
            commit_message = commit_context.get('message', '')
            diff_summary = commit_context.get('diff_summary', '')
            
            # Step 3: Simple LLM analysis for table purpose and relevance
            prompt = f"""
You are analyzing a table to determine if it needs updates based on a commit. Keep it simple and focused.

TABLE INFORMATION:
- Section Heading: {heading}
- Section Context: {heading_context}
- Table Headers: {headers}
- Number of Columns: {len(headers)}
- Number of Data Rows: {len(data_rows)}

COMMIT INFORMATION:
- Commit Message: {commit_message}
- Diff Summary: {diff_summary}

ANALYSIS REQUIRED:
1. What is the purpose of this table based on the headers?
2. Is this table relevant to the commit changes?
3. If relevant, can we create a meaningful new row with data from the commit?

CRITICAL: Use EXACT column names from the headers: {headers}
Do NOT modify, misspell, or change the column names in any way.

RESPOND WITH SIMPLE JSON (no complex structures):
{{
    "table_purpose": "Brief description of what this table represents",
    "is_relevant": true,
    "relevance_reason": "Why this table is relevant to the commit",
    "can_add_row": true,
    "new_row_data": {{
        "{headers[0] if headers else 'Column1'}": "value1",
        "{headers[1] if len(headers) > 1 else 'Column2'}": "value2"
    }},
    "confidence": 0.8
}}

IMPORTANT: 
- Use EXACT column names: {headers}
- Only respond with valid JSON
- Use double quotes, no comments, no trailing commas
- Do NOT modify column names
"""
            
            response = self.provider.generate_response(
                prompt=prompt,
                max_tokens=600,
                temperature=0.2
            )
            
            # Step 4: Simple JSON parsing with fallback
            try:
                # Basic cleaning
                cleaned_response = response.strip()
                if cleaned_response.startswith('```json'):
                    cleaned_response = cleaned_response[7:]
                elif cleaned_response.startswith('```'):
                    cleaned_response = cleaned_response[3:]
                if cleaned_response.endswith('```'):
                    cleaned_response = cleaned_response[:-3]
                
                cleaned_response = cleaned_response.strip()
                
                # Remove trailing commas
                import re
                cleaned_response = re.sub(r',(\s*[}\]])', r'\1', cleaned_response)
                
                analysis_result = json.loads(cleaned_response)
                
                # Step 5: Validate that we can fill all columns
                if analysis_result.get('can_add_row', False):
                    new_row_data = analysis_result.get('new_row_data', {})
                    
                    # Check if we have data for all headers
                    missing_columns = []
                    for header in headers:
                        if header not in new_row_data or not new_row_data[header]:
                            missing_columns.append(header)
                    
                    if missing_columns:
                        logger.warning(f"⚠️ Cannot add row - missing data for columns: {missing_columns}")
                        analysis_result['can_add_row'] = False
                        analysis_result['reason'] = f"Missing data for columns: {missing_columns}"
                    else:
                        logger.info(f"✅ Can add row with data for all {len(headers)} columns")
                        analysis_result['reason'] = "All columns have data"
                
                logger.info(f"📊 Table analysis completed - Purpose: {analysis_result.get('table_purpose', 'Unknown')}")
                logger.info(f"📊 Relevance: {analysis_result.get('is_relevant', False)}")
                logger.info(f"📊 Can add row: {analysis_result.get('can_add_row', False)}")
                
                return analysis_result
                
            except json.JSONDecodeError as e:
                logger.error(f"❌ JSON parsing failed: {e}")
                logger.error(f"Raw response: {response}")
                
                # Fallback: Simple analysis without LLM
                return {
                    "table_purpose": f"Table with {len(headers)} columns",
                    "is_relevant": False,
                    "relevance_reason": "Unable to analyze table structure",
                    "can_add_row": False,
                    "new_row_data": {},
                    "confidence": 0.0,
                    "reason": "JSON parsing failed"
                }
                
        except Exception as e:
            logger.error(f"❌ Error in simplified table analysis: {e}")
            return {
                "table_purpose": f"Analysis error: {str(e)[:50]}",
                "is_relevant": False,
                "relevance_reason": f"Error during analysis: {str(e)[:100]}",
                "can_add_row": False,
                "new_row_data": {},
                "confidence": 0.0,
                "reason": f"Analysis error: {str(e)}"
            }
    
    def _update_table_with_new_data(self, table, table_data: List[List[str]], analysis_result: Dict[str, Any]) -> List[List[str]]:
        """
        Update table data with simplified approach - only add rows if all columns can be filled.
        
        Args:
            table: Original table object
            table_data: Current table data
            analysis_result: Analysis result from simplified LLM analysis
            
        Returns:
            Updated table data
        """
        try:
            updated_data = [row[:] for row in table_data]  # Deep copy
            
            # Only proceed if we can add a row and have data for all columns
            if analysis_result.get('can_add_row', False) and analysis_result.get('is_relevant', False):
                new_row_data = analysis_result.get('new_row_data', {})
                
                if new_row_data:
                    # Get headers from first row
                    headers = table_data[0] if table_data else []
                    
                    # Create ordered row based on headers
                    ordered_row = []
                    for header in headers:
                        ordered_row.append(new_row_data.get(header, ""))
                    
                    # Add the new row
                    updated_data.append(ordered_row)
                    
                    logger.info(f"✅ Added new row to table: {ordered_row}")
                    logger.info(f"📝 Reason: {analysis_result.get('relevance_reason', 'No reason provided')}")
                else:
                    logger.warning("⚠️ Cannot add row - no row data provided")
            else:
                reason = analysis_result.get('reason', 'Not relevant or cannot add row')
                logger.info(f"ℹ️ Skipping table update: {reason}")
            
            logger.info(f"Table updated: {len(table_data)} -> {len(updated_data)} rows")
            return updated_data
            
        except Exception as e:
            logger.error(f"❌ Error updating table: {e}")
            return table_data  # Return original data on error
    
    def _replace_table_in_document(self, doc, old_table, new_table_data: List[List[str]], table_index: int) -> bool:
        """
        Replace an existing table in the document with updated data.
        Uses the table_editor.py approach for formatting.
        
        Args:
            doc: Document object
            old_table: Original table to replace
            new_table_data: New table data
            table_index: Index for logging
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Import table editor functions
            from docx.shared import Pt, RGBColor
            from docx.oxml.shared import OxmlElement, qn
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            
            # Find the table element
            table_element = old_table._element
            parent = table_element.getparent()
            
            # Create new table with updated data
            if not new_table_data:
                logger.warning(f"Table {table_index} has no data")
                return False
            
            rows_count = len(new_table_data)
            cols_count = len(new_table_data[0]) if new_table_data else 0
            
            new_table = doc.add_table(rows=rows_count, cols=cols_count)
            
            # Apply formatting to new table
            for row_idx, row_data in enumerate(new_table_data):
                row = new_table.rows[row_idx]
                
                for cell_idx, cell_text in enumerate(row_data):
                    cell = row.cells[cell_idx]
                    cell.text = cell_text
                    
                    # Get the paragraph for formatting
                    paragraph = cell.paragraphs[0]
                    
                    # Left align the text
                    paragraph.alignment = 0  # Left alignment
                    
                    # Set font size to 10.5
                    for run in paragraph.runs:
                        run.font.size = Pt(10.5)
                    
                    # Format header row (first row) with custom background and white text
                    if row_idx == 0:
                        # Make text bold and white
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                        
                        # Set custom background color RGB(52, 73, 94)
                        shading_elm = parse_xml(r'<w:shd {} w:fill="34495E"/>'.format(nsdecls('w')))
                        cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # Add borders to the table
            tbl = new_table._tbl
            tblPr = tbl.tblPr
            
            # Add table borders
            tblBorders = OxmlElement('w:tblBorders')
            
            # Define border styles
            border_elements = [
                ('w:top', 'single', '000000', '4'),
                ('w:left', 'single', '000000', '4'),
                ('w:bottom', 'single', '000000', '4'),
                ('w:right', 'single', '000000', '4'),
                ('w:insideH', 'single', '000000', '4'),
                ('w:insideV', 'single', '000000', '4')
            ]
            
            for border_name, border_style, border_color, border_size in border_elements:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), border_style)
                border.set(qn('w:color'), border_color)
                border.set(qn('w:sz'), border_size)
                tblBorders.append(border)
            
            tblPr.append(tblBorders)
            
            # Add cell padding
            tblCellMar = OxmlElement('w:tblCellMar')
            
            padding_elements = [
                ('w:top', '120'),    # 6pt padding
                ('w:left', '240'),   # 12pt padding
                ('w:bottom', '120'), # 6pt padding
                ('w:right', '120')   # 6pt padding
            ]
            
            for padding_name, padding_value in padding_elements:
                padding = OxmlElement(padding_name)
                padding.set(qn('w:w'), padding_value)
                padding.set(qn('w:type'), 'dxa')
                tblCellMar.append(padding)
            
            tblPr.append(tblCellMar)
            
            # Set minimum row height
            for row_idx, row in enumerate(new_table.rows):
                tr = row._tr
                trPr = tr.trPr
                if trPr is None:
                    trPr = OxmlElement('w:trPr')
                    tr.insert(0, trPr)
                
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), '600')  # 30pt minimum height
                trHeight.set(qn('w:hRule'), 'atLeast')
                trPr.append(trHeight)
                
                # Add vertical centering to all cells
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.tcPr
                    if tcPr is None:
                        tcPr = OxmlElement('w:tcPr')
                        tc.insert(0, tcPr)
                    
                    vAlign = OxmlElement('w:vAlign')
                    vAlign.set(qn('w:val'), 'center')
                    tcPr.append(vAlign)
            
            # Replace the old table with the new one
            new_table_element = new_table._element
            parent.replace(table_element, new_table_element)
            
            logger.info(f"SUCCESS: Replaced table {table_index} in document")
            return True
            
        except Exception as e:
            logger.error(f"Failed to replace table {table_index}: {e}")
            return False