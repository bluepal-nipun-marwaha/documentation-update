#!/usr/bin/env python3
"""
Simple test to fix paragraph duplication for "Usage Recommendations" heading.
"""

import sys
from docx import Document
from datetime import datetime
import os

def test_single_heading_update(docx_path, output_path):
    """Test updating just the 'Usage Recommendations' heading."""
    try:
        doc = Document(docx_path)
        
        # Find the "Usage Recommendations" heading
        target_heading = "Usage Recommendations"
        heading_para = None
        heading_index = -1
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip() == target_heading:
                heading_para = paragraph
                heading_index = i
                break
        
        if not heading_para:
            print(f"Could not find heading: {target_heading}")
            return False
        
        print(f"Found heading '{target_heading}' at index {heading_index}")
        
        # Find paragraphs under this heading
        paragraphs_under_heading = []
        for i in range(heading_index + 1, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            # Stop if we hit another heading
            if (para.style.name.startswith('Heading') or 
                (para.runs and any(run.bold for run in para.runs) and len(text) < 100 and not text.endswith('.'))):
                break
            
            if text:  # Only non-empty paragraphs
                paragraphs_under_heading.append({
                    'paragraph': para,
                    'text': text,
                    'index': i
                })
        
        print(f"Found {len(paragraphs_under_heading)} paragraphs under the heading")
        
        if not paragraphs_under_heading:
            print("No paragraphs to duplicate")
            return False
        
        # Show the paragraphs
        for i, para_info in enumerate(paragraphs_under_heading):
            print(f"  {i+1}. {para_info['text'][:100]}...")
        
        # Find the best paragraph to duplicate (longest one)
        best_para_info = max(paragraphs_under_heading, key=lambda x: len(x['text']))
        best_para = best_para_info['paragraph']
        
        print(f"\nBest paragraph to duplicate: {best_para_info['text'][:100]}...")
        
        # Try a simpler duplication approach
        # Get the parent element (document body)
        parent = best_para._element.getparent()
        
        # Create a copy of the paragraph element
        new_element = best_para._element.__copy__()
        
        # Insert the new element after the original paragraph
        parent.insert(parent.index(best_para._element) + 1, new_element)
        
        # Find the new paragraph object
        new_paragraph = None
        for para in doc.paragraphs:
            if para._element == new_element:
                new_paragraph = para
                break
        
        if new_paragraph:
            # Update the duplicated paragraph with new content
            new_content = "Interactive Builder Usage: The new interactive builder feature is recommended for developers who want to quickly prototype CLI applications or prefer a guided approach to CLI creation. This feature is particularly useful for complex command structures with multiple options and arguments, making CLI development more accessible to beginners while maintaining Click's powerful capabilities."
            
            new_paragraph.text = new_content
            print(f"Successfully duplicated paragraph and added: {new_content[:100]}...")
            
            # Save the document
            doc.save(output_path)
            print(f"Document saved to: {output_path}")
            return True
        else:
            print("Could not find the duplicated paragraph")
            return False
        
    except Exception as e:
        print(f"Error: {e}")
        return False

def main():
    """Test the single heading update."""
    
    input_docx = "test_step_by_step/input/Click_Professional_Documentation.docx"
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_docx = f"test_step_by_step/output/test_single_update_{timestamp}.docx"
    
    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_docx), exist_ok=True)
    
    print("Testing Single Heading Update")
    print("=" * 80)
    
    success = test_single_heading_update(input_docx, output_docx)
    
    if success:
        print("\nTest completed successfully!")
    else:
        print("\nTest failed!")

if __name__ == "__main__":
    main()
