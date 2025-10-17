#!/usr/bin/env python3
"""
Test script to verify paragraph duplication logic
"""

from docx import Document
import os

def test_paragraph_duplication():
    """Test the paragraph duplication logic"""
    
    # Load the original document
    original_file = 'test_step_by_step/input/Click_Professional_Documentation.docx'
    if not os.path.exists(original_file):
        print(f"Original file not found: {original_file}")
        return
    
    doc = Document(original_file)
    print(f"Original document has {len(doc.paragraphs)} paragraphs")
    
    # Find the Usage Recommendations section
    usage_rec_index = None
    for i, para in enumerate(doc.paragraphs):
        if 'Usage Recommendations' in para.text:
            usage_rec_index = i
            print(f"Found Usage Recommendations at paragraph {i}")
            break
    
    if usage_rec_index is None:
        print("Usage Recommendations section not found")
        return
    
    # Get the next paragraph (the one we want to duplicate)
    if usage_rec_index + 1 < len(doc.paragraphs):
        next_para = doc.paragraphs[usage_rec_index + 1]
        print(f"Next paragraph text: {next_para.text[:100]}...")
        
        # Test duplication
        parent = next_para._element.getparent()
        new_element = next_para._element.__copy__()
        parent.insert(parent.index(next_para._element) + 1, new_element)
        
        # Find the new paragraph
        new_paragraph = None
        for para in doc.paragraphs:
            if para._element == new_element:
                new_paragraph = para
                break
        
        if new_paragraph:
            print(f"Successfully duplicated paragraph")
            print(f"New paragraph text: {new_paragraph.text[:100]}...")
            
            # Update the new paragraph with test content
            new_paragraph.text = "TEST CONTENT: This is a test of the paragraph duplication logic."
            
            print(f"Updated new paragraph text: {new_paragraph.text}")
            
            # Save the test document
            test_file = "test_paragraph_duplication_output.docx"
            doc.save(test_file)
            print(f"Saved test document: {test_file}")
            
            # Verify the changes
            verify_doc = Document(test_file)
            print(f"Verification document has {len(verify_doc.paragraphs)} paragraphs")
            
            # Find our test content
            for i, para in enumerate(verify_doc.paragraphs):
                if "TEST CONTENT" in para.text:
                    print(f"Found test content at paragraph {i}: {para.text}")
                    break
        else:
            print("Failed to find duplicated paragraph")
    else:
        print("No paragraph after Usage Recommendations")

if __name__ == "__main__":
    test_paragraph_duplication()
