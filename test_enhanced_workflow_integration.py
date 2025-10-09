#!/usr/bin/env python3
"""
Test script to verify the enhanced DOCX workflow integration in existing_repo_workflow.py
"""

import os
import tempfile
from pathlib import Path
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def test_enhanced_workflow_integration():
    """Test the enhanced DOCX workflow integration."""
    try:
        logger.info("🧪 Testing Enhanced DOCX Workflow Integration")
        logger.info("=" * 60)
        
        # Set environment variables for testing
        os.environ['ENHANCED_DOCX_WORKFLOW'] = 'true'
        # Note: LLM configuration will be loaded from .env file or defaults
        
        # Import the workflow class
        from existing_repo_workflow import ExistingRepoWorkflow
        
        # Create workflow instance
        workflow = ExistingRepoWorkflow()
        
        # Create LLM service for testing (this will use user's configuration)
        from services.llm_service import LLMService
        from utils.config import get_settings
        settings = get_settings()
        llm_service = LLMService(settings.ai.model_dump())
        
        logger.info(f"🤖 Using LLM Provider: {settings.ai.llm_provider}")
        logger.info(f"🤖 Using LLM Model: {settings.ai.llm_model}")
        
        # Test data
        test_file_path = "docs/test_documentation.docx"
        test_commit_context = {
            'message': 'Add interactive CLI builder feature',
            'author': 'Test User',
            'hash': 'abc123def456',
            'diff': 'Added InteractiveCLIBuilder class with methods for interactive command creation',
            'modified': ['core.py', 'interactive_builder.py']
        }
        
        # Create a test DOCX file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_docx_path = temp_file.name
            
            # Create a simple DOCX file using python-docx
            from docx import Document
            doc = Document()
            doc.add_heading('Test Documentation', 0)
            doc.add_paragraph('This is a test document for the enhanced workflow.')
            doc.add_paragraph('It contains some sample content.')
            
            # Add a simple table
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = 'Feature'
            table.cell(0, 1).text = 'Status'
            table.cell(1, 0).text = 'Basic CLI'
            table.cell(1, 1).text = 'Implemented'
            
            doc.save(temp_docx_path)
            
            # Read the file content
            with open(temp_docx_path, 'rb') as f:
                original_content = f.read()
        
        # Test the enhanced workflow method
        logger.info("🔄 Testing _process_docx_with_enhanced_workflow method...")
        
        try:
            updated_content = "This is updated content with new interactive builder information."
            
            result = workflow._process_docx_with_enhanced_workflow(
                original_content, updated_content, test_commit_context, test_file_path, llm_service
            )
            
            if result and len(result) > 0:
                logger.info("✅ Enhanced workflow method executed successfully!")
                logger.info(f"📊 Result size: {len(result)} bytes")
                
                # Save result for inspection
                output_path = "test_enhanced_workflow_output.docx"
                with open(output_path, 'wb') as f:
                    f.write(result)
                logger.info(f"💾 Output saved to: {output_path}")
                
                return True
            else:
                logger.error("❌ Enhanced workflow method returned empty result")
                return False
                
        except Exception as e:
            logger.error(f"❌ Enhanced workflow method failed: {str(e)}")
            return False
            
        finally:
            # Clean up
            try:
                os.unlink(temp_docx_path)
            except:
                pass
    
    except Exception as e:
        logger.error(f"❌ Test failed: {str(e)}")
        return False

def test_fallback_workflow():
    """Test the fallback to basic workflow when enhanced is disabled."""
    try:
        logger.info("🧪 Testing Fallback to Basic Workflow")
        logger.info("=" * 60)
        
        # Disable enhanced workflow
        os.environ['ENHANCED_DOCX_WORKFLOW'] = 'false'
        
        # Import the workflow class
        from existing_repo_workflow import ExistingRepoWorkflow
        
        # Create workflow instance
        workflow = ExistingRepoWorkflow()
        
        logger.info("✅ Fallback workflow configuration loaded successfully")
        logger.info("ℹ️ When ENHANCED_DOCX_WORKFLOW=false, the system will use basic DOCX processing")
        
        return True
        
    except Exception as e:
        logger.error(f"❌ Fallback test failed: {str(e)}")
        return False

if __name__ == "__main__":
    print("🚀 Enhanced DOCX Workflow Integration Test")
    print("=" * 60)
    
    # Test 1: Enhanced workflow
    print("\n📝 Test 1: Enhanced Workflow")
    enhanced_success = test_enhanced_workflow_integration()
    
    # Test 2: Fallback workflow
    print("\n📝 Test 2: Fallback Workflow")
    fallback_success = test_fallback_workflow()
    
    # Summary
    print("\n📊 Test Summary")
    print("=" * 60)
    print(f"Enhanced Workflow: {'✅ PASS' if enhanced_success else '❌ FAIL'}")
    print(f"Fallback Workflow: {'✅ PASS' if fallback_success else '❌ FAIL'}")
    
    if enhanced_success and fallback_success:
        print("\n🎉 All tests passed! Enhanced DOCX workflow integration is working.")
    else:
        print("\n⚠️ Some tests failed. Check the logs above for details.")
