"""
Word Document Editor POC - Table Formatter
Extracts table content and recreates tables with custom formatting.
"""

import os
import random
import logging
from docx import Document
from docx.shared import Pt, RGBColor
from openai import OpenAI
from dotenv import load_dotenv

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def read_document(file_path):
    """
    Loads a Word document.
    
    Args:
        file_path: Path to the .docx file
        
    Returns:
        Document object
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Document not found: {file_path}")
    
    logger.info(f"Reading document: {file_path}")
    return Document(file_path)


def get_all_content_items(doc, min_length=50):
    """
    Extracts all editable content from document: paragraphs, bullets, table cells, headers, and footers.
    
    Args:
        doc: Document object
        min_length: Minimum character length to prefer
        
    Returns:
        List of tuples: (content_type, content_object, text, location_info)
        content_type: 'paragraph', 'list', 'table_cell', 'header', or 'footer'
    """
    content_items = []
    
    # 1. Get all paragraphs (includes body text and list items)
    for idx, para in enumerate(doc.paragraphs):
        if not para.text.strip() or len(para.runs) == 0:
            continue
            
        # Check if it's a list item (bullet or numbered)
        if para.style.name.startswith('List'):
            content_type = 'list'
            location = f"List item {idx + 1}"
        else:
            content_type = 'paragraph'
            location = f"Paragraph {idx + 1}"
        
        content_items.append({
            'type': content_type,
            'object': para,
            'text': para.text.strip(),
            'location': location,
            'length': len(para.text.strip())
        })
    
    # 2. Get all table cells
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                # Get cell text from all paragraphs in cell
                cell_text = ' '.join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                
                if cell_text and len(cell.paragraphs) > 0 and len(cell.paragraphs[0].runs) > 0:
                    content_items.append({
                        'type': 'table_cell',
                        'object': cell.paragraphs[0],  # First paragraph in cell
                        'text': cell_text,
                        'location': f"Table {table_idx + 1}, Row {row_idx + 1}, Cell {cell_idx + 1}",
                        'length': len(cell_text)
                    })
    
    # 3. Get headers from all sections
    for section_idx, section in enumerate(doc.sections):
        # Primary header
        for para_idx, para in enumerate(section.header.paragraphs):
            if para.text.strip() and len(para.runs) > 0:
                content_items.append({
                    'type': 'header',
                    'object': para,
                    'text': para.text.strip(),
                    'location': f"Header (Section {section_idx + 1}, Para {para_idx + 1})",
                    'length': len(para.text.strip())
                })
        
        # First page header (if different)
        if section.first_page_header:
            for para_idx, para in enumerate(section.first_page_header.paragraphs):
                if para.text.strip() and len(para.runs) > 0:
                    content_items.append({
                        'type': 'header',
                        'object': para,
                        'text': para.text.strip(),
                        'location': f"First Page Header (Section {section_idx + 1})",
                        'length': len(para.text.strip())
                    })
    
    # 4. Get footers from all sections
    for section_idx, section in enumerate(doc.sections):
        # Primary footer
        for para_idx, para in enumerate(section.footer.paragraphs):
            if para.text.strip() and len(para.runs) > 0:
                content_items.append({
                    'type': 'footer',
                    'object': para,
                    'text': para.text.strip(),
                    'location': f"Footer (Section {section_idx + 1}, Para {para_idx + 1})",
                    'length': len(para.text.strip())
                })
        
        # First page footer (if different)
        if section.first_page_footer:
            for para_idx, para in enumerate(section.first_page_footer.paragraphs):
                if para.text.strip() and len(para.runs) > 0:
                    content_items.append({
                        'type': 'footer',
                        'object': para,
                        'text': para.text.strip(),
                        'location': f"First Page Footer (Section {section_idx + 1})",
                        'length': len(para.text.strip())
                    })
    
    # Filter by minimum length if specified
    longer_items = [item for item in content_items if item['length'] >= min_length]
    
    if not longer_items:
        logger.info(f"No content >= {min_length} chars found, using any content")
        longer_items = content_items
    
    return longer_items


def select_random_content(doc, min_length=50):
    """
    Selects random content from document (paragraph, list item, or table cell).
    
    Args:
        doc: Document object
        min_length: Minimum character length to prefer (default 50)
    
    Returns:
        Tuple of (content_type, content_object, text, location)
    """
    content_items = get_all_content_items(doc, min_length)
    
    if not content_items:
        raise ValueError("No suitable content found in document")
    
    # Select random item
    selected = random.choice(content_items)
    
    logger.info(f"Selected {selected['type']} (length: {selected['length']} chars)")
    logger.info(f"  Location: {selected['location']}")
    logger.info(f"  Text: '{selected['text'][:80]}...'")
    logger.info(f"  Has {len(selected['object'].runs)} runs")
    
    return selected['type'], selected['object'], selected['text'], selected['location']


def create_formatting_instruction(original_formats):
    """
    Creates a human-readable description of formatting structure for LLM.
    
    Args:
        original_formats: List of run formatting info
        
    Returns:
        String describing the formatting structure
    """
    if not original_formats:
        return "Plain text with no special formatting."
    
    if len(original_formats) == 1:
        fmt = original_formats[0]
        styles = []
        if fmt['bold']:
            styles.append("bold")
        if fmt['italic']:
            styles.append("italic")
        if fmt['underline']:
            styles.append("underlined")
        if fmt['highlight']:
            styles.append("highlighted")
        
        if styles:
            return f"Entire text is {', '.join(styles)}."
        else:
            return "Plain text with no special formatting."
    
    # Multiple runs - describe the pattern
    parts = []
    for idx, fmt in enumerate(original_formats):
        styles = []
        if fmt['bold']:
            styles.append("BOLD")
        if fmt['italic']:
            styles.append("ITALIC")
        if fmt['underline']:
            styles.append("UNDERLINED")
        
        style_str = "+".join(styles) if styles else "NORMAL"
        parts.append(f'Part {idx+1}: "{fmt["text"]}" [{style_str}]')
    
    return "Text has multiple formatting styles:\n" + "\n".join(parts)


def get_most_used_font(doc):
    """
    Extracts the most commonly used font from the document.
    
    Args:
        doc: Document object
        
    Returns:
        String with the most used font name, or 'Calibri' as fallback
    """
    font_counts = {}
    
    # Count fonts from paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:
                font_name = run.font.name
                font_counts[font_name] = font_counts.get(font_name, 0) + 1
    
    # Count fonts from table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.name:
                            font_name = run.font.name
                            font_counts[font_name] = font_counts.get(font_name, 0) + 1
    
    if font_counts:
        most_used_font = max(font_counts, key=font_counts.get)
        logger.info(f"Most used font: {most_used_font} (used {font_counts[most_used_font]} times)")
        return most_used_font
    else:
        logger.info("No fonts found, using Calibri as fallback")
        return 'Calibri'


def extract_table_content(table):
    """
    Extracts all content from a table as a list of rows.
    
    Args:
        table: Table object from python-docx
        
    Returns:
        List of lists representing table rows and cells
    """
    table_data = []
    
    for row_idx, row in enumerate(table.rows):
        row_data = []
        for cell_idx, cell in enumerate(row.cells):
            # Get all text from cell paragraphs
            cell_text = ' '.join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
            row_data.append(cell_text)
        table_data.append(row_data)
    
    logger.info(f"Extracted table with {len(table_data)} rows and {len(table_data[0]) if table_data else 0} columns")
    return table_data


def create_formatted_table(doc, table_data, table_index):
    """
    Creates a new table with custom formatting.
    
    Args:
        doc: Document object
        table_data: List of lists containing table content
        table_index: Index of the table for logging
        
    Returns:
        New table object with custom formatting
    """
    if not table_data:
        logger.warning(f"Table {table_index} has no data")
        return None
    
    # Create new table
    rows_count = len(table_data)
    cols_count = len(table_data[0]) if table_data else 0
    
    new_table = doc.add_table(rows=rows_count, cols=cols_count)
    new_table.style = 'Table Grid'  # Basic table style
    
    # Apply custom formatting
    for row_idx, row_data in enumerate(table_data):
        row = new_table.rows[row_idx]
        
        for cell_idx, cell_text in enumerate(row_data):
            cell = row.cells[cell_idx]
            
            # Clear existing paragraphs and add new one
            cell.text = cell_text
            
            # Get the paragraph for formatting
            paragraph = cell.paragraphs[0]
            
            # Set font size to 12
            for run in paragraph.runs:
                run.font.size = Pt(12)
            
            # Format header row (first row) with light blue background and bold
            if row_idx == 0:
                # Make text bold
                for run in paragraph.runs:
                    run.font.bold = True
                
                # Set light blue background (this is a simplified approach)
                # Note: python-docx has limited background color support
                # We'll use shading instead
                from docx.oxml.shared import OxmlElement, qn
                from docx.oxml.ns import nsdecls
                from docx.oxml import parse_xml
                
                # Create shading element
                shading_elm = parse_xml(r'<w:shd {} w:fill="ADD8E6"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
    
    logger.info(f"Created formatted table {table_index} with {rows_count} rows and {cols_count} columns")
    return new_table


def parse_formatted_response(llm_response, original_formats):
    """
    Parses LLM response with formatting tags into segments.
    
    Args:
        llm_response: String from LLM with <BOLD>, <ITALIC>, <UNDERLINE> tags
        original_formats: Original formatting to use as fallback
        
    Returns:
        List of dicts with 'text' and formatting attributes
    """
    import re
    
    segments = []
    
    # If no tags found, use proportional distribution as fallback
    if '<BOLD>' not in llm_response and '<ITALIC>' not in llm_response and '<UNDERLINE>' not in llm_response:
        logger.info("No formatting tags found, using original structure")
        # Return segments matching original structure proportionally
        return None  # Signals to use proportional method
    
    # Parse tags
    position = 0
    text = llm_response
    
    # Pattern to match <TAG>content</TAG>
    tag_pattern = r'<(BOLD|ITALIC|UNDERLINE)>(.*?)</\1>'
    
    last_end = 0
    for match in re.finditer(tag_pattern, text):
        # Add any plain text before this tag
        if match.start() > last_end:
            plain_text = text[last_end:match.start()]
            if plain_text.strip():
                segments.append({
                    'text': plain_text,
                    'bold': False,
                    'italic': False,
                    'underline': False
                })
        
        # Add tagged content
        tag = match.group(1)
        content = match.group(2)
        
        segments.append({
            'text': content,
            'bold': tag == 'BOLD',
            'italic': tag == 'ITALIC',
            'underline': tag == 'UNDERLINE'
        })
        
        last_end = match.end()
    
    # Add any remaining plain text
    if last_end < len(text):
        remaining = text[last_end:]
        if remaining.strip():
            segments.append({
                'text': remaining,
                'bold': False,
                'italic': False,
                'underline': False
            })
    
    logger.info(f"Parsed {len(segments)} formatted segments from LLM response")
    for idx, seg in enumerate(segments):
        styles = []
        if seg['bold']:
            styles.append('BOLD')
        if seg['italic']:
            styles.append('ITALIC')
        if seg['underline']:
            styles.append('UNDERLINE')
        style_str = '+'.join(styles) if styles else 'NORMAL'
        logger.info(f"  Segment {idx+1}: '{seg['text']}' [{style_str}]")
    
    return segments


def has_hyperlinks(paragraph):
    """
    Checks if a paragraph contains hyperlinks.
    
    Args:
        paragraph: Paragraph object to check
        
    Returns:
        Boolean indicating if hyperlinks are present
    """
    from docx.oxml.ns import qn
    for element in paragraph._element.iterchildren():
        if element.tag == qn('w:hyperlink'):
            return True
    return False


def get_hyperlink_info(paragraph):
    """
    Extracts hyperlink information from a paragraph.
    
    Args:
        paragraph: Paragraph object
        
    Returns:
        List of dicts with hyperlink info: {'text': str, 'url': str, 'run_index': int}
    """
    from docx.oxml.ns import qn
    hyperlinks = []
    
    for element in paragraph._element.iterchildren():
        if element.tag == qn('w:hyperlink'):
            # Get relationship ID
            r_id = element.get(qn('r:id'))
            if r_id:
                try:
                    rel = paragraph.part.rels[r_id]
                    url = rel.target_ref
                    
                    # Get text from runs within hyperlink
                    text_parts = []
                    for run_elem in element.iterchildren():
                        if run_elem.tag == qn('w:r'):
                            for t in run_elem.iterchildren():
                                if t.tag == qn('w:t'):
                                    text_parts.append(t.text)
                    
                    hyperlinks.append({
                        'text': ''.join(text_parts),
                        'url': url
                    })
                except KeyError:
                    pass
    
    return hyperlinks


def capture_detailed_formatting(paragraph):
    """
    Captures formatting for each run in detail, including word-level information.
    
    Args:
        paragraph: Paragraph to analyze
        
    Returns:
        List of formatting info for each run with detailed properties
    """
    run_formats = []
    
    for idx, run in enumerate(paragraph.runs):
        if not run.text:
            continue
            
        format_info = {
            'run_index': idx,
            'text': run.text,
            'length': len(run.text),
            'font_name': run.font.name,
            'font_size': run.font.size,
            'bold': run.font.bold,
            'italic': run.font.italic,
            'underline': run.font.underline,
            'color': run.font.color.rgb if run.font.color.rgb else None,
            'highlight': run.font.highlight_color,
            'strikethrough': getattr(run.font, 'strike', None)
        }
        run_formats.append(format_info)
    
    return run_formats


def apply_proportional_formatting_safe(paragraph, new_text, original_formats):
    """
    Applies formatting proportionally while reusing existing runs (SAFE for tables/lists).
    
    Strategy: Reuses cleared runs where possible, only adds new runs if needed.
    
    Args:
        paragraph: Paragraph to modify (with cleared runs)
        new_text: New text to apply
        original_formats: List of format info from capture_detailed_formatting
    """
    if not original_formats:
        logger.warning("No original formats, using plain text")
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)
        return
    
    existing_runs = list(paragraph.runs)
    
    # Calculate total original length
    total_original_length = sum(fmt['length'] for fmt in original_formats)
    new_text_length = len(new_text)
    
    # Check if formatting is uniform
    is_uniform = True
    if len(original_formats) > 1:
        first_fmt = original_formats[0]
        
        def normalize_bool(val):
            return bool(val) if val is not None else False
        
        for idx, fmt in enumerate(original_formats[1:], 1):
            bold_same = normalize_bool(fmt['bold']) == normalize_bool(first_fmt['bold'])
            italic_same = normalize_bool(fmt['italic']) == normalize_bool(first_fmt['italic'])
            underline_same = normalize_bool(fmt['underline']) == normalize_bool(first_fmt['underline'])
            color_same = fmt['color'] == first_fmt['color']
            highlight_same = fmt['highlight'] == first_fmt['highlight']
            
            if not (bold_same and italic_same and underline_same and color_same and highlight_same):
                is_uniform = False
                break
    
    if is_uniform:
        # Uniform formatting - reuse first run
        logger.info("Applying uniform formatting (reusing first run)")
        if existing_runs:
            run = existing_runs[0]
            run.text = new_text
        else:
            run = paragraph.add_run(new_text)
        
        fmt = original_formats[0]
        if fmt['font_name']:
            run.font.name = fmt['font_name']
        if fmt['font_size']:
            run.font.size = fmt['font_size']
        if fmt['bold'] is not None:
            run.font.bold = fmt['bold']
        if fmt['italic'] is not None:
            run.font.italic = fmt['italic']
        if fmt['underline'] is not None:
            run.font.underline = fmt['underline']
        if fmt['color']:
            run.font.color.rgb = fmt['color']
        if fmt['highlight']:
            run.font.highlight_color = fmt['highlight']
    else:
        # Mixed formatting - distribute proportionally, reuse runs
        logger.info(f"Applying mixed formatting across {len(original_formats)} runs (reusing existing)")
        
        position = 0
        for idx, fmt in enumerate(original_formats):
            # Calculate segment
            if idx == len(original_formats) - 1:
                segment_text = new_text[position:]
            else:
                proportion = fmt['length'] / total_original_length
                segment_length = int(new_text_length * proportion)
                
                # Try to split at word boundary
                if segment_length > 0 and position + segment_length < new_text_length:
                    space_pos = new_text.find(' ', position + segment_length)
                    if space_pos != -1 and space_pos - position < segment_length + 20:
                        segment_length = space_pos - position + 1
                
                segment_text = new_text[position:position + segment_length]
                position += segment_length
            
            if not segment_text:
                continue
            
            # Reuse existing run or create new
            if idx < len(existing_runs):
                run = existing_runs[idx]
                run.text = segment_text
            else:
                run = paragraph.add_run(segment_text)
            
            # Apply formatting
            if fmt['font_name']:
                run.font.name = fmt['font_name']
            if fmt['font_size']:
                run.font.size = fmt['font_size']
            if fmt['bold'] is not None:
                run.font.bold = fmt['bold']
            if fmt['italic'] is not None:
                run.font.italic = fmt['italic']
            if fmt['underline'] is not None:
                run.font.underline = fmt['underline']
            if fmt['color']:
                run.font.color.rgb = fmt['color']
            if fmt['highlight']:
                run.font.highlight_color = fmt['highlight']
            
            logger.info(f"  Run {idx+1}: '{segment_text[:30]}...' with bold={fmt['bold']}, "
                       f"italic={fmt['italic']}, color={fmt['color']}")


def apply_proportional_formatting(paragraph, new_text, original_formats):
    """
    Applies formatting proportionally based on original run distribution.
    
    Strategy: If original had multiple runs with different formatting,
    distribute the new text across runs proportionally to maintain visual pattern.
    
    Args:
        paragraph: Paragraph to modify
        new_text: New text to apply
        original_formats: List of format info from capture_detailed_formatting
    """
    if not original_formats:
        logger.warning("No original formats, using plain text")
        paragraph.add_run(new_text)
        return
    
    # Calculate total original length
    total_original_length = sum(fmt['length'] for fmt in original_formats)
    new_text_length = len(new_text)
    
    # Check if formatting is uniform (all runs have same bold/italic/underline)
    # Need to handle None values properly - treat None as False for comparison
    is_uniform = True
    if len(original_formats) > 1:
        first_fmt = original_formats[0]
        
        def normalize_bool(val):
            """Treat None as False for formatting comparison"""
            return bool(val) if val is not None else False
        
        for idx, fmt in enumerate(original_formats[1:], 1):
            # Compare each formatting attribute
            bold_same = normalize_bool(fmt['bold']) == normalize_bool(first_fmt['bold'])
            italic_same = normalize_bool(fmt['italic']) == normalize_bool(first_fmt['italic'])
            underline_same = normalize_bool(fmt['underline']) == normalize_bool(first_fmt['underline'])
            color_same = fmt['color'] == first_fmt['color']
            highlight_same = fmt['highlight'] == first_fmt['highlight']
            
            if not (bold_same and italic_same and underline_same and color_same and highlight_same):
                is_uniform = False
                logger.info(f"Detected MIXED formatting at run {idx+1}:")
                logger.info(f"  Bold: {normalize_bool(first_fmt['bold'])} vs {normalize_bool(fmt['bold'])} (same={bold_same})")
                logger.info(f"  Italic: {normalize_bool(first_fmt['italic'])} vs {normalize_bool(fmt['italic'])} (same={italic_same})")
                logger.info(f"  Underline: {normalize_bool(first_fmt['underline'])} vs {normalize_bool(fmt['underline'])} (same={underline_same})")
                logger.info(f"  Color: {first_fmt['color']} vs {fmt['color']} (same={color_same})")
                logger.info(f"  Highlight: {first_fmt['highlight']} vs {fmt['highlight']} (same={highlight_same})")
                break
    
    if is_uniform:
        logger.info("All runs have UNIFORM formatting")
    
    if is_uniform:
        # Uniform formatting - apply single format to all text
        logger.info("Applying uniform formatting (all runs have same style)")
        new_run = paragraph.add_run(new_text)
        
        fmt = original_formats[0]
        if fmt['font_name']:
            new_run.font.name = fmt['font_name']
        if fmt['font_size']:
            new_run.font.size = fmt['font_size']
        if fmt['bold'] is not None:
            new_run.font.bold = fmt['bold']
        if fmt['italic'] is not None:
            new_run.font.italic = fmt['italic']
        if fmt['underline'] is not None:
            new_run.font.underline = fmt['underline']
        if fmt['color']:
            new_run.font.color.rgb = fmt['color']
        if fmt['highlight']:
            new_run.font.highlight_color = fmt['highlight']
            
    else:
        # Mixed formatting - distribute proportionally
        logger.info(f"Applying mixed formatting across {len(original_formats)} runs")
        
        position = 0
        for idx, fmt in enumerate(original_formats):
            # Calculate how much of new text this run should get
            if idx == len(original_formats) - 1:
                # Last run gets remaining text
                segment_text = new_text[position:]
            else:
                # Proportional distribution
                proportion = fmt['length'] / total_original_length
                segment_length = int(new_text_length * proportion)
                
                # Try to split at word boundary
                if segment_length > 0 and position + segment_length < new_text_length:
                    # Look for nearest space after calculated position
                    space_pos = new_text.find(' ', position + segment_length)
                    if space_pos != -1 and space_pos - position < segment_length + 20:
                        segment_length = space_pos - position + 1
                
                segment_text = new_text[position:position + segment_length]
                position += segment_length
            
            if not segment_text:
                continue
            
            # Create run with this segment
            new_run = paragraph.add_run(segment_text)
            
            # Apply formatting from original run
            if fmt['font_name']:
                new_run.font.name = fmt['font_name']
            if fmt['font_size']:
                new_run.font.size = fmt['font_size']
            if fmt['bold'] is not None:
                new_run.font.bold = fmt['bold']
            if fmt['italic'] is not None:
                new_run.font.italic = fmt['italic']
            if fmt['underline'] is not None:
                new_run.font.underline = fmt['underline']
            if fmt['color']:
                new_run.font.color.rgb = fmt['color']
            if fmt['highlight']:
                new_run.font.highlight_color = fmt['highlight']
            
            logger.info(f"  Run {idx+1}: '{segment_text[:30]}...' with bold={fmt['bold']}, "
                       f"italic={fmt['italic']}, color={fmt['color']}")


def replace_paragraph_text(paragraph, new_text_or_segments, original_formats):
    """
    Replaces paragraph text while preserving word-level formatting.
    
    Strategy: 
    1. If new_text_or_segments is a string: Use proportional distribution
    2. If new_text_or_segments is a list: Use LLM-generated segments with formatting
    
    Args:
        paragraph: Paragraph to modify
        new_text_or_segments: Either string (proportional) or list of dicts (LLM segments)
        original_formats: Original formatting captured
    """
    if not paragraph.runs:
        logger.warning("No runs found, adding plain text")
        if isinstance(new_text_or_segments, str):
            paragraph.add_run(new_text_or_segments)
        else:
            for seg in new_text_or_segments:
                paragraph.add_run(seg['text'])
        return
    
    # SAFE REMOVAL: Clear text from existing runs instead of removing XML elements
    # This preserves document structure for tables, lists, headers, footers
    logger.info(f"Clearing text from {len(paragraph.runs)} existing runs")
    for run in paragraph.runs:
        run.text = ""
    
    logger.info(f"After clearing, paragraph has {len(paragraph.runs)} runs")
    
    # Check if we have LLM-generated segments or just text
    if isinstance(new_text_or_segments, list):
        # LLM-generated segments with formatting
        logger.info(f"Applying {len(new_text_or_segments)} LLM-generated formatted segments")
        
        segments_to_apply = [seg for seg in new_text_or_segments if seg.get('text', '')]
        existing_runs = list(paragraph.runs)
        
        for idx, seg in enumerate(segments_to_apply):
            # Reuse existing run if available, otherwise create new
            if idx < len(existing_runs):
                run = existing_runs[idx]
                run.text = seg['text']
            else:
                run = paragraph.add_run(seg['text'])
            
            # Apply LLM-specified formatting
            run.font.bold = seg.get('bold', False)
            run.font.italic = seg.get('italic', False)
            run.font.underline = seg.get('underline', False)
            
            # Apply original visual properties (font, size, color) from first original format
            if original_formats:
                fmt = original_formats[0]
                if fmt['font_name']:
                    run.font.name = fmt['font_name']
                if fmt['font_size']:
                    run.font.size = fmt['font_size']
                if fmt['color']:
                    run.font.color.rgb = fmt['color']
            
            logger.info(f"  Segment {idx+1}: '{seg['text']}' [bold={seg.get('bold')}, italic={seg.get('italic')}]")
    else:
        # Plain text - use proportional distribution
        logger.info("Applying proportional formatting distribution")
        apply_proportional_formatting_safe(paragraph, new_text_or_segments, original_formats)
    
    logger.info(f"Paragraph now has {len(paragraph.runs)} runs with text: '{paragraph.text}'")


def save_document(doc, output_path):
    """
    Saves the document.
    """
    doc.save(output_path)
    logger.info(f"Document saved: {output_path}")


def replace_table_in_document(doc, old_table, new_table_data, table_index, font_name):
    """
    Replaces an existing table in the document with a new formatted table.
    
    Args:
        doc: Document object
        old_table: Original table to replace
        new_table_data: New table data with formatting
        table_index: Index for logging
        font_name: Font name to use for the table
    """
    try:
        # Find the paragraph that contains the table
        table_element = old_table._element
        
        # Get the parent element (usually a paragraph)
        parent = table_element.getparent()
        
        # Create new table element with one extra row for 1s
        new_table = doc.add_table(rows=len(new_table_data) + 1, cols=len(new_table_data[0]) if new_table_data else 0)
        # Don't set style - use default formatting
        
        # Apply formatting to new table
        for row_idx, row_data in enumerate(new_table_data):
            row = new_table.rows[row_idx]
            
            for cell_idx, cell_text in enumerate(row_data):
                cell = row.cells[cell_idx]
                cell.text = cell_text
                
                # Get the paragraph for formatting
                paragraph = cell.paragraphs[0]
                
                # Left align the text (no center alignment)
                paragraph.alignment = 0  # Left alignment
                
                # Set font size to 10.5 and font name
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = font_name
                
                # Format header row (first row) with custom background and white text
                if row_idx == 0:
                    # Make text bold and white
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                    
                    # Set custom background color RGB(31, 56, 100)
                    from docx.oxml.shared import OxmlElement, qn
                    from docx.oxml.ns import nsdecls
                    from docx.oxml import parse_xml
                    
                    # Create shading element with custom RGB color
                    # RGB(52, 73, 94) = #34495E - custom blue-gray for headers
                    shading_elm = parse_xml(r'<w:shd {} w:fill="34495E"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Add row of 1s at the end
        if new_table_data:  # Only if there's data
            last_row_idx = len(new_table_data)  # This is the extra row we added
            last_row = new_table.rows[last_row_idx]
            
            for cell_idx in range(len(new_table_data[0])):  # Use number of columns from original data
                cell = last_row.cells[cell_idx]
                cell.text = "1"
                
                # Get the paragraph for formatting
                paragraph = cell.paragraphs[0]
                
                # Left align the text
                paragraph.alignment = 0  # Left alignment
                
                # Set font size to 10.5 and font name
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = font_name
        
        # Add borders to the table
        from docx.oxml.shared import OxmlElement, qn
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        # Create table borders
        tbl = new_table._tbl
        tblPr = tbl.tblPr
        
        # Add table borders
        tblBorders = OxmlElement('w:tblBorders')
        
        # Define border styles
        border_elements = [
            ('w:top', 'single', '000000', '4'),
            ('w:left', 'single', '000000', '4'),
            ('w:bottom', 'single', '000000', '4'),
            ('w:right', 'single', '000000', '4'),
            ('w:insideH', 'single', '000000', '4'),
            ('w:insideV', 'single', '000000', '4')
        ]
        
        for border_name, border_style, border_color, border_size in border_elements:
            border = OxmlElement(border_name)
            border.set(qn('w:val'), border_style)
            border.set(qn('w:color'), border_color)
            border.set(qn('w:sz'), border_size)
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
        
        # Add cell padding with more left padding and vertical centering
        tblCellMar = OxmlElement('w:tblCellMar')
        
        # Define padding for all sides (in twips - 1/20th of a point)
        # More left padding, less right padding for better readability
        padding_elements = [
            ('w:top', '120'),    # 6pt padding
            ('w:left', '240'),   # 12pt padding (more left padding)
            ('w:bottom', '120'), # 6pt padding
            ('w:right', '120')   # 6pt padding
        ]
        
        for padding_name, padding_value in padding_elements:
            padding = OxmlElement(padding_name)
            padding.set(qn('w:w'), padding_value)
            padding.set(qn('w:type'), 'dxa')  # dxa = 1/20th of a point
            tblCellMar.append(padding)
        
        tblPr.append(tblCellMar)
        
        # Set minimum row height for consistent GitHub docs appearance
        for row_idx, row in enumerate(new_table.rows):
            # Set minimum row height (in twips - 1/20th of a point)
            # 24pt = 480 twips for good GitHub docs spacing
            tr = row._tr
            trPr = tr.trPr
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                tr.insert(0, trPr)
            
            # Add row height
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), '600')  # 30pt minimum height (taller rows)
            trHeight.set(qn('w:hRule'), 'atLeast')  # At least this height
            trPr.append(trHeight)
            
            # Add vertical centering to all cells in the row
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.tcPr
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)
                
                # Add vertical alignment
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')  # Center vertically
                tcPr.append(vAlign)
        
        # Replace the old table with the new one
        new_table_element = new_table._element
        parent.replace(table_element, new_table_element)
        
        logger.info(f"SUCCESS: Replaced table {table_index} in document")
        return True
        
    except Exception as e:
        logger.error(f"Failed to replace table {table_index}: {e}")
        return False


def main():
    """
    Main execution.
    
    Replaces all tables in the original Word document with formatted versions:
    - Font size 12
    - Header rows: Light blue background and bold text
    """
    input_file = 'sample_document.docx'
    output_file = 'sample_document_formatted.docx'
    
    try:
        # Read document
        doc = read_document(input_file)
        logger.info(f"Document has {len(doc.paragraphs)} paragraphs, "
                   f"{len(doc.tables)} tables, {len(doc.sections)} sections")
        
        if len(doc.tables) == 0:
            logger.warning("No tables found in document")
            return
        
        logger.info(f"\nFound {len(doc.tables)} tables to process")
        
        # Extract the most used font from the document
        most_used_font = get_most_used_font(doc)
        
        # Process each table (work backwards to avoid index issues)
        processed_count = 0
        tables_to_process = list(doc.tables)  # Create a copy of the list
        
        for table_idx, table in enumerate(tables_to_process, 1):
            logger.info("\n" + "=" * 60)
            logger.info(f"PROCESSING TABLE {table_idx}/{len(tables_to_process)}")
            logger.info(f"Table has {len(table.rows)} rows and {len(table.columns)} columns")
            logger.info("=" * 60)
            
            try:
                # Extract table content
                table_data = extract_table_content(table)
                
                if not table_data:
                    logger.warning(f"Table {table_idx} has no content, skipping")
                    continue
                
                # Replace the table in the document
                success = replace_table_in_document(doc, table, table_data, table_idx, most_used_font)
                
                if success:
                    logger.info(f"SUCCESS: Table {table_idx} replaced with formatted version")
                    processed_count += 1
                else:
                    logger.error(f"ERROR: Failed to replace table {table_idx}")
                
            except Exception as e:
                logger.error(f"Failed to process table {table_idx}: {e}")
                import traceback
                traceback.print_exc()
                continue
        
        # Save the modified document
        doc.save(output_file)
        
        logger.info("\n" + "=" * 60)
        logger.info("TABLE REPLACEMENT COMPLETED")
        logger.info(f"Processed: {processed_count}/{len(tables_to_process)} tables")
        logger.info(f"Output: {output_file}")
        logger.info("=" * 60)
        
    except Exception as e:
        logger.error(f"Error: {e}")
        raise


if __name__ == "__main__":
    main()
