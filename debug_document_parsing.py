#!/usr/bin/env python3
"""
Debug script to compare document parsing between final_heading_update.py and main system.
"""

from docx import Document
import os

def debug_document_parsing():
    """Debug document parsing to see what's different."""
    
    # Use the same input file as final_heading_update.py
    input_docx = "test_step_by_step/input/Click_Professional_Documentation.docx"
    
    if not os.path.exists(input_docx):
        print(f"Input file not found: {input_docx}")
        return
    
    print("Loading document...")
    doc = Document(input_docx)
    
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print("=" * 80)
    
    # Parse document structure like final_heading_update.py
    headings_with_content = []
    current_section = None
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        
        if not text:
            continue
            
        # Check if this is a heading
        is_heading = False
        heading_level = 0
        
        if paragraph.style.name.startswith('Heading'):
            is_heading = True
            heading_level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
        elif paragraph.runs and any(run.bold for run in paragraph.runs):
            if len(text) < 100 and not text.endswith('.'):
                is_heading = True
                heading_level = 1
        
        if is_heading:
            # Save previous section if it has content
            if current_section and current_section['paragraphs']:
                headings_with_content.append(current_section)
            
            # Start new section
            current_section = {
                'heading': text,
                'level': heading_level,
                'paragraph_index': i,
                'paragraphs': []
            }
            print(f"FOUND HEADING [{heading_level}]: {text}")
        else:
            # Add paragraph to current section
            if current_section:
                current_section['paragraphs'].append({
                    'text': text,
                    'index': i,
                    'paragraph_obj': paragraph
                })
                print(f"  Added paragraph: {text[:50]}...")
    
    # Don't forget the last section
    if current_section and current_section['paragraphs']:
        headings_with_content.append(current_section)
    
    print("=" * 80)
    print(f"Found {len(headings_with_content)} headings with content")
    
    for i, section in enumerate(headings_with_content):
        print(f"\n{i+1}. [{section['level']}] {section['heading']}")
        print(f"   Paragraphs: {len(section['paragraphs'])}")
        for j, para in enumerate(section['paragraphs'][:3]):  # Show first 3 paragraphs
            print(f"     {j+1}. {para['text'][:60]}...")
        if len(section['paragraphs']) > 3:
            print(f"     ... and {len(section['paragraphs']) - 3} more paragraphs")

if __name__ == "__main__":
    debug_document_parsing()
