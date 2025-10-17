#!/usr/bin/env python3
"""
Test script to parse DOCX headings and their paragraphs.
This will help us understand the document structure before implementing
the heading-by-heading analysis logic.
"""

import sys
from docx import Document
from docx.shared import Inches
import re

def parse_docx_headings_and_paragraphs(docx_path):
    """
    Parse a DOCX file and extract headings with their associated paragraphs.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary with heading structure
    """
    try:
        doc = Document(docx_path)
        
        # Structure to hold our parsed content
        document_structure = {
            'headings': [],
            'total_paragraphs': len(doc.paragraphs)
        }
        
        current_heading = None
        current_section = None
        
        print(f"Parsing document: {docx_path}")
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        print("=" * 80)
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
                
            # Check if this is a heading
            # Look for common heading patterns
            is_heading = False
            heading_level = 0
            
            # Check paragraph style for headings
            if paragraph.style.name.startswith('Heading'):
                is_heading = True
                heading_level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
            # Check for bold text (common heading indicator)
            elif paragraph.runs and any(run.bold for run in paragraph.runs):
                # Additional check: is it short and looks like a heading?
                if len(text) < 100 and not text.endswith('.'):
                    is_heading = True
                    heading_level = 1
            
            if is_heading:
                # Save previous section if it exists
                if current_section:
                    document_structure['headings'].append(current_section)
                
                # Start new section
                current_section = {
                    'heading': text,
                    'level': heading_level,
                    'paragraph_index': i,
                    'paragraphs': []
                }
                current_heading = text
                
                print(f"\nHEADING {heading_level}: {text}")
                print(f"   Paragraph index: {i}")
                print("-" * 60)
                
            else:
                # This is a regular paragraph
                if current_section:
                    current_section['paragraphs'].append({
                        'text': text,
                        'index': i,
                        'length': len(text),
                        'word_count': len(text.split())
                    })
                    
                    print(f"   Paragraph {len(current_section['paragraphs'])}: {text[:100]}{'...' if len(text) > 100 else ''}")
                    print(f"      Length: {len(text)} chars, {len(text.split())} words")
        
        # Don't forget the last section
        if current_section:
            document_structure['headings'].append(current_section)
        
        return document_structure
        
    except Exception as e:
        print(f"Error parsing document: {e}")
        return None

def display_document_summary(structure):
    """Display a summary of the document structure."""
    if not structure:
        return
        
    print("\n" + "=" * 80)
    print("DOCUMENT STRUCTURE SUMMARY")
    print("=" * 80)
    
    print(f"Total headings found: {len(structure['headings'])}")
    print(f"Total paragraphs: {structure['total_paragraphs']}")
    
    print("\nHEADINGS OVERVIEW:")
    for i, section in enumerate(structure['headings']):
        print(f"  {i+1}. [{section['level']}] {section['heading']}")
        print(f"     Paragraphs: {len(section['paragraphs'])}")
        print(f"     Start index: {section['paragraph_index']}")
        
        # Show first few words of each paragraph
        if section['paragraphs']:
            print("     Content preview:")
            for j, para in enumerate(section['paragraphs'][:3]):  # Show first 3 paragraphs
                preview = para['text'][:80] + "..." if len(para['text']) > 80 else para['text']
                print(f"        {j+1}. {preview}")
            if len(section['paragraphs']) > 3:
                print(f"        ... and {len(section['paragraphs']) - 3} more paragraphs")

def main():
    """Main function to test the heading parser."""
    
    # Test with the Click documentation file
    docx_path = "test_step_by_step/input/Click_Professional_Documentation.docx"
    
    print("Testing DOCX Heading Parser")
    print("=" * 80)
    
    # Parse the document
    structure = parse_docx_headings_and_paragraphs(docx_path)
    
    if structure:
        # Display summary
        display_document_summary(structure)
        
        print("\nParsing completed successfully!")
        print("\nNext steps:")
        print("   1. Review the heading structure above")
        print("   2. Identify which headings would be relevant for commit content")
        print("   3. Implement logic to match commit changes to appropriate headings")
        
    else:
        print("Failed to parse document")

if __name__ == "__main__":
    main()
