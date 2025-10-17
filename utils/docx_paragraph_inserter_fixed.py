"""
DOCX Paragraph Inserter Module - Fixed Version

Handles adding new paragraphs to DOCX documents by duplicating existing paragraphs.
Based on the proven approach from double_para.py
"""

import io
import logging
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import RGBColor

logger = logging.getLogger(__name__)


class DocxParagraphInserter:
    """Handles adding new paragraphs to DOCX documents by duplicating existing ones."""
    
    def __init__(self):
        self.logger = logger
    
    def insert_new_paragraphs_with_analysis(self, docx_bytes: bytes, commit_context: Dict[str, Any], llm_service, analysis: Dict[str, Any]) -> bytes:
        """
        Enhanced method for inserting paragraphs using intelligent document analysis.
        
        Args:
            docx_bytes: Original DOCX file as bytes
            commit_context: Commit information and context
            llm_service: LLM service for generating content
            analysis: Document analysis results from LLM
            
        Returns:
            Updated DOCX file as bytes
        """
        try:
            logger.info("🧠 Starting intelligent paragraph insertion with document analysis")
            
            # Check if no changes should be made
            edit_type = analysis.get('edit_type', 'add_paragraph')
            if edit_type == 'no_change':
                logger.info("🚫 LLM determined no changes should be made - preserving document integrity")
                return docx_bytes
            
            # Load the document
            doc = Document(io.BytesIO(docx_bytes))
            logger.info(f"📄 Loaded DOCX document with {len(doc.paragraphs)} paragraphs")
            
            # Extract placement information from analysis
            placement = analysis.get('placement_decision', {})
            target_section = placement.get('target_section')
            
            # Check if target section is None (no appropriate section found)
            if not target_section:
                logger.info("🚫 No appropriate section found for content placement - preserving document integrity")
                return docx_bytes
            
            placement_type = placement.get('placement_type', 'append')
            
            logger.info(f"🎯 Target section: {target_section}")
            logger.info(f"📍 Placement type: {placement_type}")
            
            # Find the target section and determine insertion point
            insertion_point = self._find_intelligent_insertion_point(doc, target_section, placement_type)
            
            if insertion_point is None:
                logger.warning("⚠️ Could not find target section, preserving document integrity")
                return docx_bytes
            
            logger.info(f"📍 Insertion point determined: paragraph {insertion_point}")
            
            # Get target section content for style matching
            target_content = self._extract_target_section_content(doc, insertion_point)
            
            # Generate contextual content using LLM
            new_content = llm_service.generate_contextual_content(commit_context, analysis, target_content)
            
            logger.info(f"✨ Generated contextual content: {len(new_content)} characters")
            
            # Insert the content at the determined point
            self._insert_contextual_content(doc, insertion_point, new_content, placement_type)
            
            # Save the updated document
            output_buffer = io.BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            
            logger.info("✅ Intelligent paragraph insertion completed successfully")
            return output_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"❌ Intelligent paragraph insertion failed: {str(e)}")
            # Fallback to original method
            return self.insert_new_paragraphs(docx_bytes, commit_context, llm_service)
    
    def _find_intelligent_insertion_point(self, doc: Document, target_section: str, placement_type: str) -> Optional[int]:
        """
        Find the intelligent insertion point based on document analysis.
        
        Args:
            doc: Document object
            target_section: Name of the target section
            placement_type: Type of placement (insert_after, insert_before, replace, append)
            
        Returns:
            Index of the insertion point, or None if not found
        """
        try:
            # Look for section headings that match the target section
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip().lower()
                target_lower = target_section.lower()
                
                # Check if this paragraph is a heading that matches our target
                if (target_lower in text or 
                    text in target_lower or 
                    any(word in text for word in target_lower.split() if len(word) > 3)):
                    
                    logger.info(f"🎯 Found potential target section at paragraph {i}: '{paragraph.text}'")
                    
                    if placement_type == "insert_after":
                        return i + 1
                    elif placement_type == "insert_before":
                        return i
                    elif placement_type == "replace":
                        return i
                    elif placement_type == "append":
                        # Find the end of this section
                        return self._find_section_end(doc, i)
            
            # If no specific section found, use intelligent fallback
            return self._find_intelligent_fallback_point(doc, placement_type)
            
        except Exception as e:
            logger.error(f"Error finding insertion point: {str(e)}")
            return None
    
    def _find_section_end(self, doc: Document, start_index: int) -> int:
        """Find the end of a section starting from the given index."""
        try:
            for i in range(start_index + 1, len(doc.paragraphs)):
                paragraph = doc.paragraphs[i]
                text = paragraph.text.strip()
                
                # Look for next heading (usually shorter text or different style)
                if (len(text) < 50 and 
                    (text.isupper() or 
                     any(char.isupper() for char in text[:10]) or
                     text.startswith(('##', '###', '####')))):
                    return i
            
            # If no next section found, return near the end
            return max(start_index + 1, len(doc.paragraphs) - 2)
            
        except Exception as e:
            logger.error(f"Error finding section end: {str(e)}")
            return len(doc.paragraphs) - 1
    
    def _find_intelligent_fallback_point(self, doc: Document, placement_type: str) -> int:
        """Find an intelligent fallback insertion point."""
        try:
            # Look for common document sections
            common_sections = ['features', 'overview', 'introduction', 'getting started', 'examples', 'usage']
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip().lower()
                
                for section in common_sections:
                    if section in text and len(text) < 100:  # Likely a heading
                        logger.info(f"🎯 Found fallback section '{section}' at paragraph {i}")
                        return i + 1 if placement_type == "insert_after" else i
            
            # Default fallback
            return len(doc.paragraphs) - 1
            
        except Exception as e:
            logger.error(f"Error in fallback point finding: {str(e)}")
            return len(doc.paragraphs) - 1
    
    def _extract_target_section_content(self, doc: Document, insertion_point: int) -> str:
        """Extract content from the target section for style matching."""
        try:
            # Get content around the insertion point
            start_idx = max(0, insertion_point - 2)
            end_idx = min(len(doc.paragraphs), insertion_point + 3)
            
            content_parts = []
            for i in range(start_idx, end_idx):
                if i < len(doc.paragraphs):
                    text = doc.paragraphs[i].text.strip()
                    if text:
                        content_parts.append(text)
            
            return "\n".join(content_parts)
            
        except Exception as e:
            logger.error(f"Error extracting target content: {str(e)}")
            return ""
    
    def _insert_contextual_content(self, doc: Document, insertion_point: int, content: str, placement_type: str):
        """Insert contextual content at the specified point."""
        try:
            # Split content into paragraphs
            content_paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
            
            if not content_paragraphs:
                logger.warning("⚠️ No content to insert")
                return
            
            logger.info(f"📝 Inserting {len(content_paragraphs)} contextual paragraphs")
            
            # Find template paragraph for formatting
            template_paragraph = None
            if insertion_point < len(doc.paragraphs):
                template_paragraph = doc.paragraphs[insertion_point]
            elif len(doc.paragraphs) > 0:
                template_paragraph = doc.paragraphs[-1]
            
            # Insert each paragraph
            for i, content_text in enumerate(content_paragraphs):
                if template_paragraph:
                    # Duplicate the template paragraph with new content
                    new_para = self._duplicate_paragraph(doc, template_paragraph, content_text)
                else:
                    # Create a new paragraph
                    new_para = doc.add_paragraph(content_text)
                
                # Insert at the correct position
                if placement_type == "insert_after":
                    insert_idx = insertion_point + i + 1
                elif placement_type == "insert_before":
                    insert_idx = insertion_point + i
                else:  # append or replace
                    insert_idx = insertion_point + i
                
                if insert_idx < len(doc.paragraphs):
                    # Insert before existing paragraph
                    para_element = new_para._element
                    ref_para = doc.paragraphs[insert_idx]
                    ref_para._element.addprevious(para_element)
                
                logger.info(f"✅ Inserted contextual paragraph {i+1}: '{content_text[:50]}...'")
            
            logger.info(f"🎯 Contextual content insertion completed")
            
        except Exception as e:
            logger.error(f"❌ Error inserting contextual content: {str(e)}")

    def insert_new_paragraphs(self, docx_bytes: bytes, commit_context: Dict[str, Any], llm_service) -> bytes:
        """
        Main entry point for inserting new paragraphs into a DOCX document.
        
        Args:
            docx_bytes: Original DOCX file as bytes
            commit_context: Commit information and context
            llm_service: LLM service for generating content
            
        Returns:
            Updated DOCX file as bytes
        """
        try:
            # Load document from bytes
            doc = Document(io.BytesIO(docx_bytes))
            self.logger.info(f"Loaded DOCX document with {len(doc.paragraphs)} paragraphs")
            
            # Find insertion point
            insertion_point = self._find_insertion_point(doc)
            if insertion_point is None:
                self.logger.warning("No suitable insertion point found")
                return docx_bytes
            
            # Generate new content
            new_content = self._generate_new_content(commit_context, llm_service)
            if not new_content:
                self.logger.warning("No new content generated")
                return docx_bytes
            
            # Insert new paragraphs
            self._insert_paragraphs_at_point(doc, insertion_point, new_content)
            
            self.logger.info(f"Inserted new content at position {insertion_point}")
            
            # Save document to bytes
            output_stream = io.BytesIO()
            doc.save(output_stream)
            return output_stream.getvalue()
            
        except Exception as e:
            self.logger.error(f"Error inserting paragraphs into DOCX document: {e}")
            return docx_bytes
    
    def _find_insertion_point(self, doc: Document) -> Optional[int]:
        """
        Find the best insertion point for new content.
        
        Args:
            doc: Document object
            
        Returns:
            Index where to insert new content, or None if no suitable point found
        """
        # Look for a good insertion point (e.g., before the last paragraph)
        if len(doc.paragraphs) > 0:
            # Insert before the last paragraph
            return len(doc.paragraphs) - 1
        
        return None
    
    def _generate_new_content(self, commit_context: Dict[str, Any], llm_service) -> List[str]:
        """
        Generate new content based on commit context.
        
        Args:
            commit_context: Commit information
            llm_service: LLM service
            
        Returns:
            List of paragraph texts to insert
        """
        try:
            commit_message = commit_context.get('message', '')
            files_changed = commit_context.get('files_changed', [])
            
            prompt = f"""
Based on the commit information, generate new documentation content:

COMMIT INFORMATION:
- Message: {commit_message}
- Files Changed: {', '.join(files_changed)}

INSTRUCTIONS:
- Generate 2-3 paragraphs of documentation content that explains the changes
- Write in a professional, technical documentation style
- Focus on what was added or changed and why it's important
- Each paragraph should be on a separate line
- Do not include markdown formatting - this is for a Word document

NEW DOCUMENTATION CONTENT:
"""
            
            response = llm_service.generate_response(prompt, temperature=0.3, max_tokens=800)
            
            # Split into paragraphs
            paragraphs = [p.strip() for p in response.split('\n') if p.strip()]
            
            # Filter out any markdown-like content
            clean_paragraphs = []
            for para in paragraphs:
                # Remove markdown formatting
                clean_para = para.replace('**', '').replace('*', '').replace('#', '').replace('`', '')
                if clean_para and len(clean_para) > 20:  # Only keep substantial paragraphs
                    clean_paragraphs.append(clean_para)
            
            self.logger.info(f"Generated {len(clean_paragraphs)} new paragraphs")
            return clean_paragraphs
            
        except Exception as e:
            self.logger.error(f"Error generating new content: {e}")
            return []
    
    def _insert_paragraphs_at_point(self, doc: Document, insertion_point: int, new_content: List[str]) -> None:
        """
        Insert new paragraphs at the specified point.
        
        Args:
            doc: Document object
            insertion_point: Index where to insert
            new_content: List of paragraph texts to insert
        """
        try:
            # Find a template paragraph to copy formatting from
            template_paragraph = None
            if insertion_point < len(doc.paragraphs):
                template_paragraph = doc.paragraphs[insertion_point]
            elif len(doc.paragraphs) > 0:
                template_paragraph = doc.paragraphs[-1]
            
            # Insert each new paragraph
            for i, content in enumerate(new_content):
                if template_paragraph:
                    # Duplicate the template paragraph
                    new_para = self._duplicate_paragraph(doc, template_paragraph, content)
                else:
                    # Create a new paragraph
                    new_para = doc.add_paragraph(content)
                
                # Insert at the correct position
                if insertion_point + i < len(doc.paragraphs):
                    # Insert before existing paragraph
                    para_element = new_para._element
                    ref_para = doc.paragraphs[insertion_point + i]
                    ref_para._element.addprevious(para_element)
                else:
                    # Add to end
                    pass  # Already added by add_paragraph
                
                self.logger.info(f"Inserted paragraph {i+1}: '{content[:50]}...'")
                
        except Exception as e:
            self.logger.error(f"Error inserting paragraphs: {e}")
    
    def _duplicate_paragraph(self, doc, template_paragraph, new_text: str):
        """
        Duplicate a paragraph with new text while preserving formatting.
        
        Args:
            doc: Document object
            template_paragraph: Paragraph to duplicate
            new_text: New text for the duplicated paragraph
            
        Returns:
            New paragraph with duplicated formatting
        """
        try:
            # Create new paragraph using the document
            new_para = doc.add_paragraph()
            
            # Copy paragraph-level formatting
            if template_paragraph.style:
                new_para.style = template_paragraph.style
            
            # Copy alignment
            if template_paragraph.alignment:
                new_para.alignment = template_paragraph.alignment
            
            # Copy formatting from template runs
            if template_paragraph.runs:
                # Use the first run's formatting as template
                template_run = template_paragraph.runs[0]
                
                # Add new text with template formatting
                new_run = new_para.add_run(new_text)
                new_run.font.bold = template_run.font.bold
                new_run.font.italic = template_run.font.italic
                new_run.font.underline = template_run.font.underline
                new_run.font.size = template_run.font.size
                
                if template_run.font.color and template_run.font.color.rgb:
                    new_run.font.color.rgb = template_run.font.color.rgb
            else:
                # No runs in template, add plain text
                new_para.add_run(new_text)
            
            return new_para
            
        except Exception as e:
            self.logger.error(f"Error duplicating paragraph: {e}")
            # Fallback: create simple paragraph
            return doc.add_paragraph(new_text)
