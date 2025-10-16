"""
DOCX Table Editor Module

Handles editing tables in DOCX documents by extracting table content,
updating it based on commit context, and applying GitHub docs style formatting.
Based on the workflow from table_editor.py but integrated with the LLM service.
"""

import io
import logging
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from collections import Counter

logger = logging.getLogger(__name__)


class DocxTableEditor:
    """Handles table editing in DOCX documents with GitHub docs style formatting."""
    
    def __init__(self):
        self.logger = logger
    
    def edit_tables(self, docx_bytes: bytes, commit_context: Dict[str, Any], llm_service) -> bytes:
        """
        Main entry point for editing tables in a DOCX document.
        
        Args:
            docx_bytes: Original DOCX file as bytes
            commit_context: Commit information and context
            llm_service: LLM service instance for text generation
            
        Returns:
            Updated DOCX file as bytes
        """
        try:
            # Load document
            doc = Document(io.BytesIO(docx_bytes))
            
            # Find tables in the document
            tables = doc.tables
            
            if not tables:
                self.logger.warning("No tables found in document")
                return docx_bytes
            
            # Get the most used font in the document
            most_used_font = self._get_most_used_font(doc)
            
            # Process each table
            for table in tables:
                try:
                    # Extract table content
                    table_data = self._extract_table_content(table)
                    
                    # Update table content based on commit context
                    updated_data = self._update_table_content(table_data, commit_context, llm_service)
                    
                    # Replace table with updated content and formatting
                    self._replace_table_in_document(table, updated_data, most_used_font)
                    
                except Exception as e:
                    self.logger.error(f"Error processing table: {e}")
                    continue
            
            # Save updated document
            output = io.BytesIO()
            doc.save(output)
            return output.getvalue()
            
        except Exception as e:
            self.logger.error(f"Error editing tables: {e}")
            return docx_bytes
    
    def _get_most_used_font(self, doc: Document) -> str:
        """
        Determine the most commonly used font in the document.
        Returns the font name or 'Calibri' as default.
        """
        try:
            font_counter = Counter()
            
            # Count fonts in paragraphs
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.font.name:
                        font_counter[run.font.name] += 1
            
            # Count fonts in table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.font.name:
                                    font_counter[run.font.name] += 1
            
            # Return most common font
            if font_counter:
                return font_counter.most_common(1)[0][0]
            else:
                return 'Calibri'  # Default font
                
        except Exception as e:
            self.logger.error(f"Error determining most used font: {e}")
            return 'Calibri'
    
    def _extract_table_content(self, table) -> List[List[str]]:
        """
        Extract table content as a list of lists (rows and columns).
        """
        try:
            table_data = []
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Get text from all paragraphs in the cell
                    cell_text = []
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            cell_text.append(paragraph.text.strip())
                    
                    # Join paragraphs with newlines
                    cell_content = '\n'.join(cell_text) if cell_text else ''
                    row_data.append(cell_content)
                
                table_data.append(row_data)
            
            return table_data
            
        except Exception as e:
            self.logger.error(f"Error extracting table content: {e}")
            return []
    
    def _update_table_content(self, table_data: List[List[str]], commit_context: Dict[str, Any], llm_service) -> List[List[str]]:
        """
        Update table content based on commit context using LLM.
        """
        try:
            if not table_data:
                return table_data
            
            # Convert table data to text format for LLM
            table_text = self._table_data_to_text(table_data)
            
            prompt = f"""
            Based on this commit information, update the following table content:
            
            Commit Message: {commit_context.get('message', '')}
            Files Changed: {', '.join(commit_context.get('files', []))}
            
            Current table content:
            {table_text}
            
            Update the table content to reflect the changes described in the commit.
            Return the updated table in the same format (pipe-separated values).
            Maintain the same number of rows and columns.
            """
            
            response = llm_service.generate_response(prompt)
            
            # Parse the response back to table data
            updated_data = self._text_to_table_data(response)
            
            # Ensure we have the same dimensions as original
            if len(updated_data) != len(table_data) or (table_data and len(updated_data[0]) != len(table_data[0])):
                self.logger.warning("LLM response doesn't match table dimensions, using original data")
                return table_data
            
            return updated_data
            
        except Exception as e:
            self.logger.error(f"Error updating table content: {e}")
            return table_data
    
    def _table_data_to_text(self, table_data: List[List[str]]) -> str:
        """Convert table data to text format for LLM processing."""
        try:
            lines = []
            for row in table_data:
                # Escape pipes in cell content
                escaped_row = [cell.replace('|', '\\|') for cell in row]
                lines.append('| ' + ' | '.join(escaped_row) + ' |')
            
            return '\n'.join(lines)
            
        except Exception as e:
            self.logger.error(f"Error converting table data to text: {e}")
            return str(table_data)
    
    def _text_to_table_data(self, text: str) -> List[List[str]]:
        """Convert LLM response text back to table data format."""
        try:
            lines = text.strip().split('\n')
            table_data = []
            
            for line in lines:
                if '|' in line:
                    # Remove leading/trailing pipes and split
                    cells = line.strip('|').split('|')
                    # Clean up cell content
                    cleaned_cells = [cell.strip().replace('\\|', '|') for cell in cells]
                    table_data.append(cleaned_cells)
            
            return table_data
            
        except Exception as e:
            self.logger.error(f"Error converting text to table data: {e}")
            return []
    
    def _replace_table_in_document(self, old_table, new_data: List[List[str]], font_name: str):
        """
        Replace the old table with new data and apply GitHub docs style formatting.
        """
        try:
            if not new_data:
                return
            
            # Get table dimensions
            rows = len(new_data)
            cols = len(new_data[0]) if new_data else 0
            
            if rows == 0 or cols == 0:
                return
            
            # Clear existing table content
            for row in old_table.rows:
                for cell in row.cells:
                    # Clear all paragraphs in the cell
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.text = ""
                        paragraph.text = ""
            
            # Ensure we have enough rows
            while len(old_table.rows) < rows:
                old_table.add_row()
            
            # Ensure we have enough columns
            while len(old_table.columns) < cols:
                old_table.add_column(Inches(1.5))
            
            # Fill table with new data
            for i, row_data in enumerate(new_data):
                if i < len(old_table.rows):
                    row = old_table.rows[i]
                    for j, cell_data in enumerate(row_data):
                        if j < len(row.cells):
                            cell = row.cells[j]
                            
                            # Clear existing content
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = ""
                                paragraph.text = ""
                            
                            # Set new content
                            if cell.paragraphs:
                                cell.paragraphs[0].text = cell_data
                            else:
                                cell.add_paragraph(cell_data)
            
            # Apply GitHub docs style formatting
            self._apply_github_docs_formatting(old_table, font_name)
            
        except Exception as e:
            self.logger.error(f"Error replacing table: {e}")
    
    def _apply_github_docs_formatting(self, table, font_name: str):
        """
        Apply GitHub docs style formatting to the table.
        """
        try:
            # Set table alignment
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Apply formatting to each cell
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    # Set cell properties
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    # Apply formatting to paragraphs in the cell
                    for paragraph in cell.paragraphs:
                        # Set paragraph alignment
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Apply formatting to runs
                        for run in paragraph.runs:
                            # Set font
                            run.font.name = font_name
                            run.font.size = Pt(10.5)
                            
                            # Header row formatting (first row)
                            if i == 0:
                                run.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                
                                # Set cell background color
                                self._set_cell_background_color(cell, RGBColor(52, 73, 94))  # Dark blue-gray
                            else:
                                run.bold = False
                                run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
                                
                                # Set cell background color
                                self._set_cell_background_color(cell, RGBColor(255, 255, 255))  # White
            
            # Set table borders
            self._set_table_borders(table)
            
            # Set row heights
            for row in table.rows:
                row.height = Inches(0.4)  # Minimum 30pt height
            
            # Add a row of "1"s at the bottom (as per requirements)
            if table.rows:
                # Add new row
                new_row = table.add_row()
                for cell in new_row.cells:
                    cell.text = "1"
                    # Apply same formatting as other rows
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.name = font_name
                            run.font.size = Pt(10.5)
                            run.bold = False
                            run.font.color.rgb = RGBColor(0, 0, 0)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    self._set_cell_background_color(cell, RGBColor(255, 255, 255))
            
        except Exception as e:
            self.logger.error(f"Error applying GitHub docs formatting: {e}")
    
    def _set_cell_background_color(self, cell, color: RGBColor):
        """Set the background color of a table cell."""
        try:
            # Get the cell's XML element
            cell_xml = cell._tc
            
            # Create shading element
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), color.hex[1:])  # Remove # from hex color
            
            # Add shading to cell properties
            cell_props = cell_xml.get_or_add_tcPr()
            cell_props.append(shading)
            
        except Exception as e:
            self.logger.error(f"Error setting cell background color: {e}")
    
    def _set_table_borders(self, table):
        """Set table borders to black, single line, 4pt."""
        try:
            # Get table XML element
            table_xml = table._tbl
            
            # Create table borders element
            borders = OxmlElement('w:tblBorders')
            
            # Define border types
            border_types = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
            
            for border_type in border_types:
                border = OxmlElement(f'w:{border_type}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')  # Black
                borders.append(border)
            
            # Add borders to table properties
            table_props = table_xml.get_or_add_tblPr()
            table_props.append(borders)
            
        except Exception as e:
            self.logger.error(f"Error setting table borders: {e}")
    
    def _set_cell_padding(self, cell, left: int = 12, top: int = 6, right: int = 6, bottom: int = 6):
        """Set cell padding (left, top, right, bottom in points)."""
        try:
            # Get cell XML element
            cell_xml = cell._tc
            
            # Create cell margins element
            margins = OxmlElement('w:tcMar')
            
            # Set margins
            margin_types = [
                ('left', left),
                ('top', top),
                ('right', right),
                ('bottom', bottom)
            ]
            
            for margin_type, value in margin_types:
                margin = OxmlElement(f'w:{margin_type}')
                margin.set(qn('w:w'), str(value))
                margin.set(qn('w:type'), 'dxa')  # Points
                margins.append(margin)
            
            # Add margins to cell properties
            cell_props = cell_xml.get_or_add_tcPr()
            cell_props.append(margins)
            
        except Exception as e:
            self.logger.error(f"Error setting cell padding: {e}")

