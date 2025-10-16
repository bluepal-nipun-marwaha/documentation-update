"""
DOCX Line Editor Module - Fixed Version

Handles editing existing lines/paragraphs in DOCX documents while preserving formatting.
Based on the proven approach from document_editor (fixed).py
"""

import io
import logging
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import RGBColor

logger = logging.getLogger(__name__)


class DocxLineEditor:
    """Handles line/paragraph editing in DOCX documents with formatting preservation."""
    
    def __init__(self):
        self.logger = logger
    
    def edit_document_lines(self, docx_bytes: bytes, commit_context: Dict[str, Any], llm_service) -> bytes:
        """
        Main entry point for editing lines in a DOCX document.
        
        Args:
            docx_bytes: Original DOCX file as bytes
            commit_context: Commit information and context
            llm_service: LLM service for generating updates
            
        Returns:
            Updated DOCX file as bytes
        """
        try:
            # Load document from bytes
            doc = Document(io.BytesIO(docx_bytes))
            self.logger.info(f"Loaded DOCX document with {len(doc.paragraphs)} paragraphs")
            
            # Get all content items that can be edited
            content_items = self._get_all_content_items(doc)
            self.logger.info(f"Found {len(content_items)} editable content items")
            
            if not content_items:
                self.logger.warning("No content items found to edit")
                return docx_bytes
            
            # Select items to edit based on commit context
            items_to_edit = self._select_items_to_edit(content_items, commit_context)
            self.logger.info(f"Selected {len(items_to_edit)} items to edit")
            
            # Edit each selected item
            edited_count = 0
            for item in items_to_edit:
                try:
                    self._edit_content_item(item, commit_context, llm_service)
                    edited_count += 1
                except Exception as e:
                    self.logger.error(f"Failed to edit item: {e}")
            
            self.logger.info(f"Successfully edited {edited_count} content items")
            
            # Save document to bytes
            output_stream = io.BytesIO()
            doc.save(output_stream)
            return output_stream.getvalue()
            
        except Exception as e:
            self.logger.error(f"Error editing DOCX document: {e}")
            return docx_bytes
    
    def _get_all_content_items(self, doc: Document) -> List[Dict[str, Any]]:
        """
        Extract all editable content from document: paragraphs, table cells, headers, footers.
        
        Args:
            doc: Document object
            
        Returns:
            List of content items with metadata
        """
        content_items = []
        
        # 1. Get paragraphs from main document
        for para_idx, para in enumerate(doc.paragraphs):
            if para.text.strip() and len(para.runs) > 0:
                content_items.append({
                    'type': 'paragraph',
                    'object': para,
                    'text': para.text.strip(),
                    'location': f"Paragraph {para_idx + 1}",
                    'length': len(para.text.strip())
                })
        
        # 2. Get table cells
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        if para.text.strip() and len(para.runs) > 0:
                            content_items.append({
                                'type': 'table_cell',
                                'object': para,
                                'text': para.text.strip(),
                                'location': f"Table {table_idx + 1}, Row {row_idx + 1}, Cell {cell_idx + 1}",
                                'length': len(para.text.strip())
                            })
        
        # 3. Get headers from all sections
        for section_idx, section in enumerate(doc.sections):
            for para_idx, para in enumerate(section.header.paragraphs):
                if para.text.strip() and len(para.runs) > 0:
                    content_items.append({
                        'type': 'header',
                        'object': para,
                        'text': para.text.strip(),
                        'location': f"Header (Section {section_idx + 1})",
                        'length': len(para.text.strip())
                    })
        
        # 4. Get footers from all sections
        for section_idx, section in enumerate(doc.sections):
            for para_idx, para in enumerate(section.footer.paragraphs):
                if para.text.strip() and len(para.runs) > 0:
                    content_items.append({
                        'type': 'footer',
                        'object': para,
                        'text': para.text.strip(),
                        'location': f"Footer (Section {section_idx + 1})",
                        'length': len(para.text.strip())
                    })
        
        # Filter by minimum length
        min_length = 20
        longer_items = [item for item in content_items if item['length'] >= min_length]
        
        if not longer_items:
            self.logger.info(f"No content >= {min_length} chars found, using any content")
            longer_items = content_items
        
        return longer_items
    
    def _select_items_to_edit(self, content_items: List[Dict[str, Any]], commit_context: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Select which content items to edit based on commit context.
        
        Args:
            content_items: All available content items
            commit_context: Commit information
            
        Returns:
            Selected items to edit
        """
        # For now, select a few random items for editing
        # In a real implementation, this would use LLM to select relevant content
        import random
        
        # Select up to 3 items to edit
        max_items = min(3, len(content_items))
        selected_items = random.sample(content_items, max_items)
        
        self.logger.info(f"Selected {len(selected_items)} items for editing")
        for item in selected_items:
            self.logger.info(f"  - {item['type']}: {item['text'][:50]}...")
        
        return selected_items
    
    def _edit_content_item(self, item: Dict[str, Any], commit_context: Dict[str, Any], llm_service) -> None:
        """
        Edit a single content item using LLM.
        
        Args:
            item: Content item to edit
            commit_context: Commit information
            llm_service: LLM service for generating updates
        """
        try:
            paragraph = item['object']
            original_text = item['text']
            
            # Capture original formatting
            original_formats = self._capture_detailed_formatting(paragraph)
            
            # Generate updated text using LLM
            updated_text = self._generate_updated_text(original_text, commit_context, llm_service)
            
            if updated_text and updated_text != original_text:
                # Replace text while preserving formatting
                self._replace_paragraph_text(paragraph, updated_text, original_formats)
                self.logger.info(f"Updated {item['type']}: '{original_text[:30]}...' -> '{updated_text[:30]}...'")
            else:
                self.logger.info(f"No changes needed for {item['type']}")
                
        except Exception as e:
            self.logger.error(f"Error editing content item: {e}")
    
    def _capture_detailed_formatting(self, paragraph) -> List[Dict[str, Any]]:
        """
        Capture detailed formatting information from a paragraph.
        
        Args:
            paragraph: Paragraph object
            
        Returns:
            List of formatting information for each run
        """
        formats = []
        
        for run in paragraph.runs:
            if run.text.strip():
                formats.append({
                    'text': run.text,
                    'bold': run.font.bold,
                    'italic': run.font.italic,
                    'underline': run.font.underline,
                    'color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None
                })
        
        return formats
    
    def _generate_updated_text(self, original_text: str, commit_context: Dict[str, Any], llm_service) -> str:
        """
        Generate updated text using LLM based on commit context.
        
        Args:
            original_text: Original text to update
            commit_context: Commit information
            llm_service: LLM service
            
        Returns:
            Updated text
        """
        try:
            commit_message = commit_context.get('message', '')
            files_changed = commit_context.get('files_changed', [])
            
            prompt = f"""
Update the following text based on the commit information:

COMMIT INFORMATION:
- Message: {commit_message}
- Files Changed: {', '.join(files_changed)}

ORIGINAL TEXT:
{original_text}

INSTRUCTIONS:
- Update the text to reflect the changes described in the commit
- Keep the same tone and style as the original
- Make the text more accurate or up-to-date based on the commit
- Return only the updated text, no explanations

UPDATED TEXT:
"""
            
            updated_text = llm_service.generate_response(prompt, temperature=0.3, max_tokens=500)
            
            # Clean up the response
            updated_text = updated_text.strip()
            if updated_text.startswith('"') and updated_text.endswith('"'):
                updated_text = updated_text[1:-1]
            
            return updated_text
            
        except Exception as e:
            self.logger.error(f"Error generating updated text: {e}")
            return original_text
    
    def _replace_paragraph_text(self, paragraph, new_text: str, original_formats: List[Dict[str, Any]]) -> None:
        """
        Replace paragraph text while preserving formatting.
        
        Args:
            paragraph: Paragraph to modify
            new_text: New text to insert
            original_formats: Original formatting information
        """
        if not paragraph.runs:
            self.logger.warning("No runs found, adding plain text")
            paragraph.add_run(new_text)
            return
        
        # Clear text from existing runs (preserve structure)
        self.logger.info(f"Clearing text from {len(paragraph.runs)} existing runs")
        for run in paragraph.runs:
            run.text = ""
        
        # Apply new text with proportional formatting
        self._apply_proportional_formatting(paragraph, new_text, original_formats)
    
    def _apply_proportional_formatting(self, paragraph, new_text: str, original_formats: List[Dict[str, Any]]) -> None:
        """
        Apply formatting proportionally to new text based on original formatting.
        
        Args:
            paragraph: Paragraph object
            new_text: New text to format
            original_formats: Original formatting information
        """
        if not original_formats:
            # No original formatting, add plain text
            paragraph.add_run(new_text)
            return
        
        # Calculate proportional distribution
        total_original_length = sum(len(fmt['text']) for fmt in original_formats)
        if total_original_length == 0:
            paragraph.add_run(new_text)
            return
        
        # Distribute new text proportionally across runs
        existing_runs = list(paragraph.runs)
        new_text_chars = list(new_text)
        char_index = 0
        
        for fmt_idx, fmt in enumerate(original_formats):
            if char_index >= len(new_text_chars):
                break
                
            # Calculate how many characters this run should get
            proportion = len(fmt['text']) / total_original_length
            chars_for_this_run = max(1, int(len(new_text) * proportion))
            
            # Get characters for this run
            run_text = ''.join(new_text_chars[char_index:char_index + chars_for_this_run])
            char_index += chars_for_this_run
            
            # Apply to existing run or create new one
            if fmt_idx < len(existing_runs):
                run = existing_runs[fmt_idx]
                run.text = run_text
            else:
                run = paragraph.add_run(run_text)
            
            # Apply formatting
            run.font.bold = fmt['bold']
            run.font.italic = fmt['italic']
            run.font.underline = fmt['underline']
            
            if fmt['color']:
                try:
                    run.font.color.rgb = RGBColor.from_string(fmt['color'])
                except:
                    pass  # Ignore color errors
        
        # Add any remaining characters
        if char_index < len(new_text_chars):
            remaining_text = ''.join(new_text_chars[char_index:])
            paragraph.add_run(remaining_text)
