"""
DOCX Table Editor Module - Fixed Version

Handles updating tables in DOCX documents with proper formatting.
Based on the proven approach from table_editor.py
"""

import io
import logging
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import RGBColor, Pt
from docx.table import Table
from docx.oxml.shared import OxmlElement, qn

logger = logging.getLogger(__name__)


class DocxTableEditor:
    """Handles table editing in DOCX documents with proper formatting."""
    
    def __init__(self):
        self.logger = logger
    
    def edit_document_tables(self, docx_bytes: bytes, commit_context: Dict[str, Any], llm_service) -> bytes:
        """
        Main entry point for editing tables in a DOCX document.
        
        Args:
            docx_bytes: Original DOCX file as bytes
            commit_context: Commit information and context
            llm_service: LLM service for generating updates
            
        Returns:
            Updated DOCX file as bytes
        """
        try:
            # Load document from bytes
            doc = Document(io.BytesIO(docx_bytes))
            self.logger.info(f"Loaded DOCX document with {len(doc.tables)} tables")
            
            if not doc.tables:
                self.logger.info("No tables found in document")
                return docx_bytes
            
            # Process each table
            updated_tables = 0
            for table_idx, table in enumerate(doc.tables):
                try:
                    if self._update_table_content(table, commit_context, llm_service):
                        updated_tables += 1
                        self.logger.info(f"Updated table {table_idx + 1}")
                except Exception as e:
                    self.logger.error(f"Failed to update table {table_idx + 1}: {e}")
            
            self.logger.info(f"Successfully updated {updated_tables} tables")
            
            # Save document to bytes
            output_stream = io.BytesIO()
            doc.save(output_stream)
            return output_stream.getvalue()
            
        except Exception as e:
            self.logger.error(f"Error editing DOCX tables: {e}")
            return docx_bytes
    
    def _update_table_content(self, table: Table, commit_context: Dict[str, Any], llm_service) -> bool:
        """
        Update content in a table based on commit context.
        
        Args:
            table: Table object to update
            commit_context: Commit information
            llm_service: LLM service
            
        Returns:
            True if table was updated, False otherwise
        """
        try:
            # Extract current table content
            current_content = self._extract_table_content(table)
            if not current_content:
                self.logger.warning("No content found in table")
                return False
            
            # Generate updated content
            updated_content = self._generate_updated_table_content(current_content, commit_context, llm_service)
            if not updated_content:
                self.logger.warning("No updated content generated")
                return False
            
            # Apply updates to table
            self._apply_table_updates(table, updated_content)
            
            return True
            
        except Exception as e:
            self.logger.error(f"Error updating table content: {e}")
            return False
    
    def _extract_table_content(self, table: Table) -> List[List[str]]:
        """
        Extract content from a table.
        
        Args:
            table: Table object
            
        Returns:
            List of rows, each containing list of cell texts
        """
        content = []
        
        for row_idx, row in enumerate(table.rows):
            row_content = []
            for cell_idx, cell in enumerate(row.cells):
                # Get text from all paragraphs in the cell
                cell_text = ""
                for para in cell.paragraphs:
                    if para.text.strip():
                        cell_text += para.text.strip() + " "
                
                row_content.append(cell_text.strip())
            content.append(row_content)
        
        return content
    
    def _generate_updated_table_content(self, current_content: List[List[str]], commit_context: Dict[str, Any], llm_service) -> List[List[str]]:
        """
        Generate updated table content using LLM.
        
        Args:
            current_content: Current table content
            commit_context: Commit information
            llm_service: LLM service
            
        Returns:
            Updated table content
        """
        try:
            commit_message = commit_context.get('message', '')
            files_changed = commit_context.get('files_changed', [])
            
            # Convert current content to text format
            current_text = self._table_content_to_text(current_content)
            
            prompt = f"""
Update the following table content based on the commit information:

COMMIT INFORMATION:
- Message: {commit_message}
- Files Changed: {', '.join(files_changed)}

CURRENT TABLE CONTENT:
{current_text}

INSTRUCTIONS:
- Update the table content to reflect the changes described in the commit
- Keep the same table structure (same number of rows and columns)
- Update relevant cells with new information
- Return the updated table in the same format
- Do not add markdown formatting

UPDATED TABLE CONTENT:
"""
            
            response = llm_service.generate_response(prompt, temperature=0.3, max_tokens=1000)
            
            # Parse the response back to table format
            updated_content = self._text_to_table_content(response, len(current_content[0]) if current_content else 0)
            
            return updated_content
            
        except Exception as e:
            self.logger.error(f"Error generating updated table content: {e}")
            return current_content
    
    def _table_content_to_text(self, content: List[List[str]]) -> str:
        """
        Convert table content to text format for LLM processing.
        
        Args:
            content: Table content as list of rows
            
        Returns:
            Text representation of the table
        """
        if not content:
            return "Empty table"
        
        lines = []
        for row_idx, row in enumerate(content):
            line = " | ".join(cell for cell in row)
            lines.append(line)
        
        return "\n".join(lines)
    
    def _text_to_table_content(self, text: str, expected_columns: int) -> List[List[str]]:
        """
        Convert text response back to table content format.
        
        Args:
            text: Text response from LLM
            expected_columns: Expected number of columns
            
        Returns:
            Table content as list of rows
        """
        try:
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            content = []
            
            for line in lines:
                # Split by | separator
                if '|' in line:
                    cells = [cell.strip() for cell in line.split('|')]
                else:
                    # If no separator, treat as single cell
                    cells = [line]
                
                # Ensure we have the expected number of columns
                while len(cells) < expected_columns:
                    cells.append("")
                
                # Truncate if too many columns
                if len(cells) > expected_columns:
                    cells = cells[:expected_columns]
                
                content.append(cells)
            
            return content
            
        except Exception as e:
            self.logger.error(f"Error parsing table content: {e}")
            return []
    
    def _apply_table_updates(self, table: Table, updated_content: List[List[str]]) -> None:
        """
        Apply updated content to the table.
        
        Args:
            table: Table object to update
            updated_content: Updated content to apply
        """
        try:
            # Ensure table has enough rows
            while len(table.rows) < len(updated_content):
                table.add_row()
            
            # Update each cell
            for row_idx, row_content in enumerate(updated_content):
                if row_idx < len(table.rows):
                    row = table.rows[row_idx]
                    
                    # Ensure row has enough cells
                    while len(row.cells) < len(row_content):
                        row.add_cell()
                    
                    # Update each cell
                    for cell_idx, cell_text in enumerate(row_content):
                        if cell_idx < len(row.cells):
                            cell = row.cells[cell_idx]
                            
                            # Clear existing content
                            for para in cell.paragraphs:
                                para.clear()
                            
                            # Add new content
                            if cell_text:
                                cell.paragraphs[0].add_run(cell_text)
            
            # Apply GitHub docs style formatting
            self._apply_github_docs_formatting(table)
            
        except Exception as e:
            self.logger.error(f"Error applying table updates: {e}")
    
    def _apply_github_docs_formatting(self, table: Table) -> None:
        """
        Apply GitHub docs style formatting to the table.
        
        Args:
            table: Table object to format
        """
        try:
            # Set table borders
            self._set_table_borders(table)
            
            # Set header row formatting
            if len(table.rows) > 0:
                self._format_header_row(table.rows[0])
            
            # Set cell padding
            self._set_cell_padding(table)
            
        except Exception as e:
            self.logger.error(f"Error applying GitHub docs formatting: {e}")
    
    def _set_table_borders(self, table: Table) -> None:
        """
        Set table borders to GitHub docs style.
        
        Args:
            table: Table object
        """
        try:
            tbl = table._tbl
            tblBorders = OxmlElement('w:tblBorders')
            
            # Define border style
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)
            
            tbl.tblPr.append(tblBorders)
            
        except Exception as e:
            self.logger.error(f"Error setting table borders: {e}")
    
    def _format_header_row(self, header_row) -> None:
        """
        Format the header row with bold text and background color.
        
        Args:
            header_row: Header row object
        """
        try:
            for cell in header_row.cells:
                # Set background color
                self._set_cell_background_color(cell, 'F6F8FA')
                
                # Set text to bold
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        
        except Exception as e:
            self.logger.error(f"Error formatting header row: {e}")
    
    def _set_cell_background_color(self, cell, color_hex: str) -> None:
        """
        Set background color for a cell.
        
        Args:
            cell: Cell object
            color_hex: Hex color code (without #)
        """
        try:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:color'), 'auto')
            shading_elm.set(qn('w:fill'), color_hex)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
        except Exception as e:
            self.logger.error(f"Error setting cell background color: {e}")
    
    def _set_cell_padding(self, table: Table) -> None:
        """
        Set cell padding for the table.
        
        Args:
            table: Table object
        """
        try:
            tbl = table._tbl
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tbl.tblPr.append(tblLayout)
            
        except Exception as e:
            self.logger.error(f"Error setting cell padding: {e}")
