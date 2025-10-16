"""
DOCX Paragraph Inserter Module

Handles adding new paragraphs to DOCX documents by duplicating existing paragraphs
and then editing the duplicated content. Based on the workflow from double_para.py
but integrated with the LLM service.
"""

import io
import logging
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)


class DocxParagraphInserter:
    """Handles paragraph insertion in DOCX documents via duplication and editing."""
    
    def __init__(self):
        self.logger = logger
    
    def insert_paragraph(self, docx_bytes: bytes, commit_context: Dict[str, Any], llm_service) -> bytes:
        """
        Main entry point for inserting paragraphs in a DOCX document.
        
        Args:
            docx_bytes: Original DOCX file as bytes
            commit_context: Commit information and context
            llm_service: LLM service instance for text generation
            
        Returns:
            Updated DOCX file as bytes
        """
        try:
            # Load document
            doc = Document(io.BytesIO(docx_bytes))
            
            # Get all content items
            content_items = self._get_all_content_items(doc)
            
            if not content_items:
                self.logger.warning("No content items found in document")
                return docx_bytes
            
            # Find insertion point
            insertion_point = self._find_insertion_point(doc, commit_context, llm_service)
            
            if not insertion_point:
                self.logger.warning("No suitable insertion point found")
                return docx_bytes
            
            # Find nearby paragraph to duplicate
            source_paragraph = self._find_source_paragraph(content_items, insertion_point)
            
            if not source_paragraph:
                self.logger.warning("No suitable source paragraph found for duplication")
                return docx_bytes
            
            # Duplicate the paragraph
            duplicated_paragraph = self._duplicate_paragraph(source_paragraph, insertion_point)
            
            # Generate new content for the duplicated paragraph
            new_content = self._generate_new_content(
                source_paragraph.text, 
                commit_context, 
                llm_service
            )
            
            # Replace the duplicated paragraph's content
            self._replace_paragraph_content(duplicated_paragraph, new_content)
            
            # Save updated document
            output = io.BytesIO()
            doc.save(output)
            return output.getvalue()
            
        except Exception as e:
            self.logger.error(f"Error inserting paragraph: {e}")
            return docx_bytes
    
    def _get_all_content_items(self, doc: Document) -> List[Any]:
        """Get all content items from the document."""
        content_items = []
        
        # Add paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only non-empty paragraphs
                content_items.append(paragraph)
        
        # Add table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            content_items.append(paragraph)
        
        return content_items
    
    def _find_insertion_point(self, doc: Document, commit_context: Dict[str, Any], llm_service) -> Optional[Any]:
        """
        Find the best insertion point for new content based on commit context.
        Uses LLM to determine where new content should be added.
        """
        try:
            # Get document structure
            paragraphs = [p for p in doc.paragraphs if p.text.strip()]
            
            if not paragraphs:
                return None
            
            # Create context for LLM
            doc_structure = []
            for i, para in enumerate(paragraphs):
                doc_structure.append(f"{i}: {para.text[:100]}...")  # First 100 chars
            
            prompt = f"""
            Based on this commit information, determine where new content should be inserted in this document:
            
            Commit Message: {commit_context.get('message', '')}
            Files Changed: {', '.join(commit_context.get('files', []))}
            
            Document structure:
            {chr(10).join(doc_structure)}
            
            Return the paragraph number (0-based index) after which the new content should be inserted.
            If content should be added at the beginning, return -1.
            If content should be added at the end, return {len(paragraphs)}.
            """
            
            response = llm_service.generate_response(prompt)
            
            # Parse response
            try:
                insertion_index = int(response.strip())
                
                if insertion_index == -1:
                    # Insert at beginning
                    return paragraphs[0] if paragraphs else None
                elif insertion_index >= len(paragraphs):
                    # Insert at end
                    return paragraphs[-1] if paragraphs else None
                else:
                    # Insert after specified paragraph
                    return paragraphs[insertion_index]
                    
            except ValueError:
                # Fallback: insert at end
                return paragraphs[-1] if paragraphs else None
                
        except Exception as e:
            self.logger.error(f"Error finding insertion point: {e}")
            # Fallback: insert at end
            paragraphs = [p for p in doc.paragraphs if p.text.strip()]
            return paragraphs[-1] if paragraphs else None
    
    def _find_source_paragraph(self, content_items: List[Any], insertion_point: Any) -> Optional[Any]:
        """
        Find the best source paragraph to duplicate based on insertion point.
        Prefers paragraphs near the insertion point with similar formatting.
        """
        try:
            if not content_items:
                return None
            
            # If insertion point is in content_items, use it
            if insertion_point in content_items:
                return insertion_point
            
            # Find the closest paragraph to insertion point
            # For simplicity, use the last paragraph as source
            return content_items[-1]
            
        except Exception as e:
            self.logger.error(f"Error finding source paragraph: {e}")
            return content_items[0] if content_items else None
    
    def _duplicate_paragraph(self, source_paragraph: Any, insertion_point: Any) -> Any:
        """
        Duplicate a paragraph with formatting preservation.
        Creates a new paragraph with duplicated content and formatting.
        """
        try:
            # Get the parent element (document or table cell)
            parent = source_paragraph._element.getparent()
            
            # Create new paragraph element
            new_para_elem = source_paragraph._element.__copy__()
            
            # Insert the new paragraph after the insertion point
            if insertion_point._element.getnext() is not None:
                # Insert after insertion point
                insertion_point._element.addnext(new_para_elem)
            else:
                # Insert as last child
                parent.append(new_para_elem)
            
            # Get the new paragraph object
            new_paragraph = None
            for para in parent.xpath('.//w:p'):
                if para == new_para_elem:
                    # Find the paragraph object that corresponds to this element
                    for doc_para in source_paragraph._parent.paragraphs:
                        if doc_para._element == new_para_elem:
                            new_paragraph = doc_para
                            break
                    break
            
            if not new_paragraph:
                # Fallback: create a simple new paragraph
                new_paragraph = source_paragraph._parent.add_paragraph()
                # Copy formatting
                new_paragraph.style = source_paragraph.style
                for run in new_paragraph.runs:
                    run.text = source_paragraph.text + "\n"  # Add newline as separator
            
            return new_paragraph
            
        except Exception as e:
            self.logger.error(f"Error duplicating paragraph: {e}")
            # Fallback: create simple paragraph
            try:
                new_para = source_paragraph._parent.add_paragraph()
                new_para.style = source_paragraph.style
                new_para.text = source_paragraph.text + "\n"
                return new_para
            except:
                return None
    
    def _generate_new_content(self, source_text: str, commit_context: Dict[str, Any], llm_service) -> str:
        """
        Generate new content for the duplicated paragraph based on commit context.
        """
        try:
            prompt = f"""
            Based on this commit information, generate new content that should be added to the documentation.
            The content should be relevant to the changes described in the commit.
            
            Commit Message: {commit_context.get('message', '')}
            Files Changed: {', '.join(commit_context.get('files', []))}
            
            Original text context: {source_text[:200]}...
            
            Generate appropriate new content that should be added to the documentation.
            Return only the new content text.
            """
            
            response = llm_service.generate_response(prompt)
            return response.strip()
            
        except Exception as e:
            self.logger.error(f"Error generating new content: {e}")
            return f"New content related to: {commit_context.get('message', 'Update')}"
    
    def _replace_paragraph_content(self, paragraph: Any, new_content: str):
        """
        Replace paragraph content with new text while preserving basic formatting.
        """
        try:
            # Clear existing runs
            for run in paragraph.runs:
                run.text = ""
            
            # Add new content
            if paragraph.runs:
                paragraph.runs[0].text = new_content
            else:
                paragraph.add_run(new_content)
                
        except Exception as e:
            self.logger.error(f"Error replacing paragraph content: {e}")
            # Fallback: just set the text
            paragraph.text = new_content
    
    def _capture_detailed_formatting(self, paragraph) -> List[Dict[str, Any]]:
        """
        Capture detailed formatting information for each run in the paragraph.
        Returns a list of formatting dictionaries, one per run.
        """
        formats = []
        
        for run in paragraph.runs:
            format_info = {
                'text': run.text,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'color': None,
                'highlight': None
            }
            
            # Capture font color
            if run.font.color.rgb:
                format_info['color'] = run.font.color.rgb
            
            # Capture highlight color
            if hasattr(run.font, 'highlight_color') and run.font.highlight_color:
                format_info['highlight'] = run.font.highlight_color
            
            formats.append(format_info)
        
        return formats
    
    def _apply_proportional_formatting_safe(self, paragraph, new_text: str, original_formats: List[Dict[str, Any]]):
        """
        Apply text with proportional formatting distribution.
        Used as fallback when direct formatting application fails.
        """
        try:
            # Clear existing runs
            for run in paragraph.runs:
                run.text = ""
            
            # Distribute text proportionally across runs
            if not original_formats:
                # No original formatting, just set text
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)
                return
            
            # Calculate proportional distribution
            total_original_length = sum(len(fmt['text']) for fmt in original_formats)
            if total_original_length == 0:
                total_original_length = 1  # Avoid division by zero
            
            text_position = 0
            for i, orig_format in enumerate(original_formats):
                if i >= len(paragraph.runs):
                    run = paragraph.add_run()
                else:
                    run = paragraph.runs[i]
                
                # Calculate proportional length
                proportion = len(orig_format['text']) / total_original_length
                segment_length = int(len(new_text) * proportion)
                
                # Get text segment
                if i == len(original_formats) - 1:
                    # Last segment gets remaining text
                    segment_text = new_text[text_position:]
                else:
                    segment_text = new_text[text_position:text_position + segment_length]
                
                # Set text and formatting
                run.text = segment_text
                run.bold = orig_format.get('bold', False)
                run.italic = orig_format.get('italic', False)
                run.underline = orig_format.get('underline', False)
                
                if orig_format.get('color'):
                    run.font.color.rgb = orig_format['color']
                if orig_format.get('highlight'):
                    run.font.highlight_color = orig_format['highlight']
                
                text_position += segment_length
                
        except Exception as e:
            self.logger.error(f"Error applying proportional formatting: {e}")
            # Ultimate fallback: just set the text
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)

