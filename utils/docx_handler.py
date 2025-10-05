"""
DOCX file handling utilities for document processing.
Enhanced with Markdown conversion for better LLM processing.
"""

import io
import re
from typing import Optional, Dict, Any, Tuple, List
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import structlog

logger = structlog.get_logger(__name__)

class DOCXHandler:
    """Handler for DOCX file operations with enhanced Markdown conversion."""
    
    @staticmethod
    def docx_to_markdown_with_metadata(docx_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        Convert DOCX to Markdown while preserving formatting metadata for later restoration.
        
        Args:
            docx_content: Raw DOCX file content as bytes
            
        Returns:
            Tuple of (markdown_content, formatting_metadata)
        """
        try:
            if not docx_content or len(docx_content) < 100:
                logger.warning("[WARNING] Invalid or empty DOCX content provided")
                return "", {}
                
            doc = Document(io.BytesIO(docx_content))
            
            markdown_parts = []
            formatting_metadata = {
                'styles': {},
                'paragraph_formats': {},
                'run_formats': {},
                'tables': {},
                'document_properties': {},
                'numbering': {},
                'original_structure': []
            }
            
            # Extract document properties
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                formatting_metadata['document_properties']['title'] = doc.core_properties.title
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                formatting_metadata['document_properties']['author'] = doc.core_properties.author
            
            # Process paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():  # Skip empty paragraphs
                    md_text, para_format = DOCXHandler._paragraph_to_markdown(paragraph, i)
                    if md_text.strip():
                        markdown_parts.append(md_text)
                        formatting_metadata['paragraph_formats'][i] = para_format
                        formatting_metadata['original_structure'].append({'type': 'paragraph', 'index': i})
            
            # Process tables
            for i, table in enumerate(doc.tables):
                md_table, table_format = DOCXHandler._table_to_markdown(table, i)
                if md_table.strip():
                    markdown_parts.append(md_table)
                    formatting_metadata['tables'][i] = table_format
                    formatting_metadata['original_structure'].append({'type': 'table', 'index': i})
            
            markdown_content = '\n\n'.join(markdown_parts)
            logger.info(f"[SUCCESS] Converted DOCX to Markdown: {len(markdown_content)} characters")
            
            return markdown_content, formatting_metadata
            
        except Exception as e:
            logger.error(f"[ERROR] Failed to convert DOCX to Markdown: {str(e)}")
            # Fallback to simple text extraction
            return DOCXHandler.extract_text_from_docx(docx_content), {}
    
    @staticmethod
    def _paragraph_to_markdown(paragraph, index: int) -> Tuple[str, Dict[str, Any]]:
        """Convert a DOCX paragraph to Markdown with formatting metadata."""
        try:
            text = paragraph.text.strip()
            if not text:
                return "", {}
            
            # Extract paragraph formatting
            para_format = {
                'alignment': paragraph.alignment,
                'style': paragraph.style.name if paragraph.style else None,
                'space_before': paragraph.paragraph_format.space_before,
                'space_after': paragraph.paragraph_format.space_after,
                'line_spacing': paragraph.paragraph_format.line_spacing,
                'runs': []
            }
            
            # Check if it's a heading based on style
            if paragraph.style and 'heading' in paragraph.style.name.lower():
                level = 1
                if 'heading 1' in paragraph.style.name.lower():
                    level = 1
                elif 'heading 2' in paragraph.style.name.lower():
                    level = 2
                elif 'heading 3' in paragraph.style.name.lower():
                    level = 3
                elif 'heading 4' in paragraph.style.name.lower():
                    level = 4
                elif 'heading 5' in paragraph.style.name.lower():
                    level = 5
                elif 'heading 6' in paragraph.style.name.lower():
                    level = 6
                
                markdown_text = '#' * level + ' ' + text
                para_format['is_heading'] = True
                para_format['heading_level'] = level
                
            else:
                # Process runs for inline formatting
                markdown_text = ""
                for run in paragraph.runs:
                    run_text = run.text
                    run_format = {
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name if run.font.name else None,
                        'font_size': run.font.size.pt if run.font.size else None,
                        'font_color': None
                    }
                    
                    # Extract font color if available
                    if run.font.color and run.font.color.rgb:
                        run_format['font_color'] = str(run.font.color.rgb)
                    
                    para_format['runs'].append(run_format)
                    
                    # Apply markdown formatting
                    if run.bold and run.italic:
                        run_text = f"***{run_text}***"
                    elif run.bold:
                        run_text = f"**{run_text}**"
                    elif run.italic:
                        run_text = f"*{run_text}*"
                    
                    markdown_text += run_text
                
                # Handle list items
                if paragraph.style and ('list' in paragraph.style.name.lower() or 'bullet' in paragraph.style.name.lower()):
                    markdown_text = f"- {markdown_text}"
                    para_format['is_list'] = True
                elif paragraph.style and 'number' in paragraph.style.name.lower():
                    markdown_text = f"1. {markdown_text}"
                    para_format['is_numbered_list'] = True
            
            return markdown_text, para_format
            
        except Exception as e:
            logger.warning(f"[WARNING] Error converting paragraph to markdown: {str(e)}")
            return paragraph.text.strip(), {}
    
    @staticmethod
    def _table_to_markdown(table, index: int) -> Tuple[str, Dict[str, Any]]:
        """Convert a DOCX table to Markdown table format."""
        try:
            table_format = {
                'rows': len(table.rows),
                'columns': len(table.columns) if table.rows else 0,
                'style': table.style.name if table.style else None,
                'cell_formats': []
            }
            
            if not table.rows:
                return "", table_format
            
            markdown_rows = []
            
            # Process each row
            for row_idx, row in enumerate(table.rows):
                row_cells = []
                row_formats = []
                
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip().replace('\n', ' ')
                    cell_format = {
                        'alignment': None,
                        'bold': False,
                        'italic': False
                    }
                    
                    # Check for formatting in cell paragraphs
                    for para in cell.paragraphs:
                        if para.runs:
                            for run in para.runs:
                                if run.bold:
                                    cell_format['bold'] = True
                                if run.italic:
                                    cell_format['italic'] = True
                    
                    # Apply markdown formatting to cell content
                    if cell_format['bold'] and cell_format['italic']:
                        cell_text = f"***{cell_text}***"
                    elif cell_format['bold']:
                        cell_text = f"**{cell_text}**"
                    elif cell_format['italic']:
                        cell_text = f"*{cell_text}*"
                    
                    row_cells.append(cell_text)
                    row_formats.append(cell_format)
                
                markdown_rows.append('| ' + ' | '.join(row_cells) + ' |')
                table_format['cell_formats'].append(row_formats)
                
                # Add header separator after first row
                if row_idx == 0 and len(row_cells) > 0:
                    separator = '| ' + ' | '.join(['---'] * len(row_cells)) + ' |'
                    markdown_rows.append(separator)
            
            markdown_table = '\n'.join(markdown_rows)
            return markdown_table, table_format
            
        except Exception as e:
            logger.warning(f"[WARNING] Error converting table to markdown: {str(e)}")
            return "", {}
    
    @staticmethod
    def markdown_to_docx_with_formatting(markdown_content: str, 
                                       formatting_metadata: Dict[str, Any],
                                       original_docx: bytes = None) -> bytes:
        """
        Convert Markdown back to DOCX while restoring original formatting.
        
        Args:
            markdown_content: Updated Markdown content from LLM
            formatting_metadata: Original formatting metadata
            original_docx: Original DOCX file for style reference
            
        Returns:
            Updated DOCX file content as bytes
        """
        try:
            if not markdown_content or not markdown_content.strip():
                logger.warning("[WARNING] Empty markdown content provided")
                return original_docx if original_docx else DOCXHandler.create_docx_from_text("", "Empty Document")
            
            # Create new document or use original as template
            if original_docx:
                try:
                    # Use original document as template to preserve styles
                    doc = Document(io.BytesIO(original_docx))
                    # Clear existing content but keep styles and formatting
                    for paragraph in doc.paragraphs[:]:
                        p = paragraph._element
                        p.getparent().remove(p)
                        p._p = p._element = None
                except Exception as e:
                    logger.warning(f"[WARNING] Could not use original DOCX as template: {str(e)}")
                    doc = Document()
            else:
                doc = Document()
            
            # Add document title if available
            doc_props = formatting_metadata.get('document_properties', {})
            if doc_props.get('title'):
                try:
                    # Try different heading approaches
                    title_added = False
                    
                    # Try level 0 heading (Title style)
                    try:
                        title_para = doc.add_heading(doc_props['title'], 0)
                        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        title_added = True
                    except:
                        # Try level 1 heading
                        try:
                            title_para = doc.add_heading(doc_props['title'], 1)
                            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            title_added = True
                        except:
                            pass
                    
                    if not title_added:
                        # Fallback: add as regular paragraph with bold formatting
                        title_para = doc.add_paragraph()
                        title_run = title_para.add_run(doc_props['title'])
                        title_run.bold = True
                        title_run.font.size = Inches(18/72)  # Larger font for title
                        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                except Exception as e:
                    # Final fallback: simple paragraph
                    title_para = doc.add_paragraph()
                    title_run = title_para.add_run(doc_props['title'])
                    title_run.bold = True
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Get available styles to reduce warnings
            available_styles = DOCXHandler._get_available_styles(doc)
            
            # Parse markdown content
            lines = markdown_content.split('\n')
            i = 0
            
            while i < len(lines):
                line = lines[i].strip()
                
                if not line:
                    i += 1
                    continue
                
                # Handle headings
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    text = line.strip('#').strip()
                    try:
                        heading = doc.add_heading(text, level=min(level, 6))
                        # Apply original heading formatting to preserve document style
                        DOCXHandler._apply_original_formatting(heading, formatting_metadata, 'heading', level)
                    except Exception as e:
                        logger.warning(f"[WARNING] Could not create heading level {level}: {str(e)}")
                        # Fallback: create as bold paragraph
                        para = doc.add_paragraph()
                        run = para.add_run(text)
                        run.bold = True
                        run.font.name = 'Calibri'
                        if level == 1:
                            run.font.size = Inches(18/72)  # Larger font for main headings
                        elif level == 2:
                            run.font.size = Inches(16/72)  # Medium font for subheadings
                        elif level == 3:
                            run.font.size = Inches(14/72)  # Smaller font for sub-subheadings
                        else:
                            run.font.size = Inches(12/72)  # Default size
                    
                # Handle tables
                elif line.startswith('|') and '|' in line[1:]:
                    table_lines = [line]
                    i += 1
                    
                    # Collect all table lines
                    while i < len(lines) and lines[i].strip().startswith('|'):
                        table_lines.append(lines[i].strip())
                        i += 1
                    
                    # Create table
                    DOCXHandler._create_table_from_markdown(doc, table_lines, formatting_metadata)
                    continue  # i is already incremented
                
                # Handle list items
                elif line.startswith('- ') or line.startswith('* '):
                    list_items = [line[2:].strip()]
                    i += 1
                    
                    # Collect consecutive list items
                    while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                        list_items.append(lines[i].strip()[2:].strip())
                        i += 1
                    
                    # Create list
                    for item in list_items:
                        para = doc.add_paragraph()
                        # Try to find a suitable bullet list style
                        style_applied = False
                        bullet_styles = ['list bullet', 'listbullet', 'bullet list', 'list', 'list paragraph']
                        
                        for style_key in bullet_styles:
                            if style_key in available_styles:
                                try:
                                    para.style = available_styles[style_key]
                                    style_applied = True
                                    break
                                except:
                                    continue
                        
                        if not style_applied:
                            # Fallback: add bullet manually
                            bullet_run = para.add_run('• ')
                        
                        DOCXHandler._add_formatted_text_to_paragraph(para, item, formatting_metadata)
                        
                        # Apply original formatting to preserve document style
                        DOCXHandler._apply_original_formatting(para, formatting_metadata, 'paragraph')
                    
                    continue  # i is already incremented
                
                # Handle numbered lists
                elif re.match(r'^\d+\.\s', line):
                    list_items = [re.sub(r'^\d+\.\s', '', line)]
                    i += 1
                    
                    # Collect consecutive numbered list items
                    while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                        list_items.append(re.sub(r'^\d+\.\s', '', lines[i].strip()))
                        i += 1
                    
                    # Create numbered list
                    for idx, item in enumerate(list_items, 1):
                        para = doc.add_paragraph()
                        # Try to find a suitable numbered list style
                        style_applied = False
                        number_styles = ['list number', 'listnumber', 'numbered list', '1', 'list paragraph']
                        
                        for style_key in number_styles:
                            if style_key in available_styles:
                                try:
                                    para.style = available_styles[style_key]
                                    style_applied = True
                                    break
                                except:
                                    continue
                        
                        if not style_applied:
                            # Fallback: add number manually
                            number_run = para.add_run(f'{idx}. ')
                        
                        DOCXHandler._add_formatted_text_to_paragraph(para, item, formatting_metadata)
                        
                        # Apply original formatting to preserve document style
                        DOCXHandler._apply_original_formatting(para, formatting_metadata, 'paragraph')
                    
                    continue  # i is already incremented
                
                # Handle regular paragraphs
                else:
                    para = doc.add_paragraph()
                    DOCXHandler._add_formatted_text_to_paragraph(para, line, formatting_metadata)
                    # Apply original paragraph formatting to preserve document style
                    DOCXHandler._apply_original_formatting(para, formatting_metadata, 'paragraph')
                
                i += 1
            
            # Save to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            result_bytes = doc_bytes.getvalue()
            logger.info(f"[SUCCESS] Converted Markdown to DOCX: {len(result_bytes)} bytes")
            return result_bytes
            
        except Exception as e:
            logger.error(f"[ERROR] Failed to convert Markdown to DOCX: {str(e)}")
            # Fallback to simple text conversion
            return DOCXHandler.create_docx_from_text(markdown_content, "Updated Document")
    
    @staticmethod
    def _add_formatted_text_to_paragraph(paragraph, text: str, formatting_metadata: Dict[str, Any]):
        """Add formatted text to a paragraph, parsing markdown formatting."""
        try:
            # Parse markdown formatting in text
            parts = DOCXHandler._parse_markdown_formatting(text)
            
            for part in parts:
                run = paragraph.add_run(part['text'])
                
                # Apply markdown formatting
                if part.get('bold'):
                    run.bold = True
                if part.get('italic'):
                    run.italic = True
                if part.get('underline'):
                    run.underline = True
                
                # Note: Font properties will be applied later by _apply_original_formatting
                # to preserve the original document's font style
                        
        except Exception as e:
            logger.warning(f"[WARNING] Error adding formatted text: {str(e)}")
            paragraph.add_run(text)
    
    @staticmethod
    def _parse_markdown_formatting(text: str) -> List[Dict[str, Any]]:
        """Parse markdown formatting from text."""
        parts = []
        
        # Simple regex patterns for markdown formatting
        patterns = [
            (r'\*\*\*(.*?)\*\*\*', {'bold': True, 'italic': True}),  # Bold + Italic
            (r'\*\*(.*?)\*\*', {'bold': True}),                      # Bold
            (r'\*(.*?)\*', {'italic': True}),                        # Italic
            (r'__(.*?)__', {'underline': True}),                     # Underline
        ]
        
        remaining_text = text
        current_pos = 0
        
        while current_pos < len(remaining_text):
            earliest_match = None
            earliest_pos = len(remaining_text)
            
            # Find the earliest formatting pattern
            for pattern, formatting in patterns:
                match = re.search(pattern, remaining_text[current_pos:])
                if match and match.start() + current_pos < earliest_pos:
                    earliest_pos = match.start() + current_pos
                    earliest_match = (match, formatting, pattern)
            
            if earliest_match:
                match, formatting, pattern = earliest_match
                
                # Add text before the match
                if earliest_pos > current_pos:
                    parts.append({
                        'text': remaining_text[current_pos:earliest_pos],
                        'bold': False,
                        'italic': False,
                        'underline': False
                    })
                
                # Add the formatted text
                formatted_part = {'text': match.group(1)}
                formatted_part.update(formatting)
                parts.append(formatted_part)
                
                current_pos = earliest_pos + match.end()
            else:
                # No more formatting, add remaining text
                if current_pos < len(remaining_text):
                    parts.append({
                        'text': remaining_text[current_pos:],
                        'bold': False,
                        'italic': False,
                        'underline': False
                    })
                break
        
        return parts if parts else [{'text': text, 'bold': False, 'italic': False, 'underline': False}]
    
    @staticmethod
    def _create_table_from_markdown(doc, table_lines: List[str], formatting_metadata: Dict[str, Any]):
        """Create a DOCX table from markdown table lines."""
        try:
            # Filter out separator lines
            data_lines = [line for line in table_lines if not re.match(r'^\|\s*[-:]+\s*\|', line)]
            
            if not data_lines:
                return
            
            # Parse table data
            rows_data = []
            for line in data_lines:
                # Remove leading/trailing pipes and split
                cells = [cell.strip() for cell in line.strip('|').split('|')]
                if cells:
                    rows_data.append(cells)
            
            if not rows_data:
                return
            
            # Create table
            table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
            table.style = 'Table Grid'
            
            # Populate table
            for row_idx, row_data in enumerate(rows_data):
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < len(table.rows[row_idx].cells):
                        cell = table.rows[row_idx].cells[col_idx]
                        
                        # Clear existing content
                        cell.text = ""
                        
                        # Add formatted content
                        if cell.paragraphs:
                            para = cell.paragraphs[0]
                        else:
                            para = cell.add_paragraph()
                        
                        DOCXHandler._add_formatted_text_to_paragraph(para, cell_data, formatting_metadata)
                        
                        # Apply table formatting if available
                        if row_idx == 0:  # Header row
                            for run in para.runs:
                                run.bold = True
                        
                        # Apply original formatting to preserve document style
                        DOCXHandler._apply_original_formatting(para, formatting_metadata, 'paragraph')
            
        except Exception as e:
            logger.warning(f"[WARNING] Error creating table from markdown: {str(e)}")
    
    @staticmethod
    def _get_available_styles(doc) -> Dict[str, str]:
        """Get available styles in the document to avoid style warnings."""
        try:
            available_styles = {}
            for style in doc.styles:
                available_styles[style.name.lower()] = style.name
            return available_styles
        except Exception:
            return {}
    
    @staticmethod
    def _apply_original_formatting(element, formatting_metadata: Dict[str, Any], element_type: str, level: int = None):
        """Apply original formatting from metadata to preserve document style."""
        try:
            # Find matching original formatting
            original_format = None
            
            if element_type == 'heading' and level:
                # Look for heading formatting in paragraph_formats
                for para_format in formatting_metadata.get('paragraph_formats', {}).values():
                    if para_format.get('is_heading') and para_format.get('heading_level') == level:
                        original_format = para_format
                        break
            elif element_type == 'paragraph':
                # For regular paragraphs, use the first non-heading format as template
                for para_format in formatting_metadata.get('paragraph_formats', {}).values():
                    if not para_format.get('is_heading'):
                        original_format = para_format
                        break
            
            if original_format and original_format.get('runs'):
                # Apply original run formatting to preserve fonts and sizes
                for i, run in enumerate(element.runs):
                    if i < len(original_format['runs']):
                        run_format = original_format['runs'][i]
                        DOCXHandler._apply_original_run_formatting(run, run_format)
                    elif len(original_format['runs']) > 0:
                        # Use the last run format as template for additional runs
                        run_format = original_format['runs'][-1]
                        DOCXHandler._apply_original_run_formatting(run, run_format)
            
        except Exception as e:
            logger.warning(f"[WARNING] Could not apply original formatting: {str(e)}")
    
    @staticmethod
    def _apply_original_run_formatting(run, run_format_data):
        """Apply original run formatting from metadata."""
        try:
            if isinstance(run_format_data, dict):
                # Apply font properties only if they exist in original
                if run_format_data.get('font_name'):
                    run.font.name = run_format_data['font_name']
                if run_format_data.get('font_size'):
                    run.font.size = Inches(run_format_data['font_size']/72)
                if run_format_data.get('bold') is not None:
                    run.bold = run_format_data['bold']
                if run_format_data.get('italic') is not None:
                    run.italic = run_format_data['italic']
                if run_format_data.get('underline') is not None:
                    run.underline = run_format_data['underline']
        except Exception as e:
            logger.warning(f"[WARNING] Could not apply run formatting: {str(e)}")
    
    @staticmethod
    def _apply_paragraph_formatting(paragraph, formatting_metadata: Dict[str, Any], style_type: str):
        """Apply original paragraph formatting to a new paragraph."""
        try:
            # Apply basic formatting based on style type
            if style_type == 'heading':
                # Headings are already formatted by add_heading
                pass
            
            # Could extend this to apply more specific formatting from metadata
            
        except Exception as e:
            logger.warning(f"[WARNING] Error applying paragraph formatting: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(docx_content: bytes) -> str:
        """
        Extract text content from DOCX file bytes.
        
        Args:
            docx_content: Raw DOCX file content as bytes
            
        Returns:
            Extracted text content as string
        """
        try:
            # Create document from bytes
            doc = Document(io.BytesIO(docx_content))
            
            # Extract text from all paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n".join(text_content)
            logger.info(f"[SUCCESS] Extracted {len(full_text)} characters from DOCX")
            return full_text
            
        except Exception as e:
            logger.error(f"[ERROR] Failed to extract text from DOCX: {str(e)}")
            return ""
    
    @staticmethod
    def create_docx_from_text(text_content: str, title: str = "Document") -> bytes:
        """
        Create DOCX document from text content.
        
        Args:
            text_content: Text content to convert to DOCX
            title: Document title
            
        Returns:
            DOCX file content as bytes
        """
        try:
            # Create new document
            doc = Document()
            
            # Add title
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add content
            paragraphs = text_content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    # Check if it's a heading (starts with #)
                    if para_text.strip().startswith('#'):
                        level = len(para_text) - len(para_text.lstrip('#'))
                        heading_text = para_text.strip('#').strip()
                        doc.add_heading(heading_text, level=min(level, 6))
                    else:
                        doc.add_paragraph(para_text.strip())
            
            # Save to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            result_bytes = doc_bytes.getvalue()
            logger.info(f"[SUCCESS] Created DOCX document with {len(result_bytes)} bytes")
            return result_bytes
            
        except Exception as e:
            logger.error(f"[ERROR] Failed to create DOCX from text: {str(e)}")
            return b""
    
    @staticmethod
    def update_docx_content_enhanced(original_docx: bytes, updated_text: str) -> bytes:
        """
        Enhanced DOCX update using Markdown conversion for better LLM processing.
        
        Args:
            original_docx: Original DOCX file bytes
            updated_text: Updated text content from LLM (should be Markdown)
            
        Returns:
            Updated DOCX file bytes
        """
        try:
            logger.info("🚀 Using enhanced DOCX update with Markdown conversion")
            
            # Step 1: Convert original DOCX to Markdown with metadata
            markdown_content, formatting_metadata = DOCXHandler.docx_to_markdown_with_metadata(original_docx)
            
            # Step 2: The updated_text should already be Markdown from LLM
            # If it's not Markdown, we'll treat it as such anyway
            updated_markdown = updated_text
            
            # Step 3: Convert updated Markdown back to DOCX with formatting restoration
            updated_docx = DOCXHandler.markdown_to_docx_with_formatting(
                updated_markdown, formatting_metadata, original_docx
            )
            
            logger.info("[SUCCESS] Enhanced DOCX update completed successfully")
            return updated_docx
            
        except Exception as e:
            logger.error(f"[ERROR] Enhanced DOCX update failed: {str(e)}")
            logger.info("🔄 Falling back to legacy DOCX update method")
            return DOCXHandler.update_docx_content_legacy(original_docx, updated_text)
    
    @staticmethod
    def update_docx_content(original_docx: bytes, updated_text: str) -> bytes:
        """
        Main DOCX update method - uses enhanced method by default with fallback.
        
        Args:
            original_docx: Original DOCX file bytes
            updated_text: Updated text content from LLM
            
        Returns:
            Updated DOCX file bytes
        """
        # Try enhanced method first
        try:
            return DOCXHandler.update_docx_content_enhanced(original_docx, updated_text)
        except Exception as e:
            logger.warning(f"[WARNING] Enhanced method failed, using legacy: {str(e)}")
            return DOCXHandler.update_docx_content_legacy(original_docx, updated_text)
    
    @staticmethod
    def update_docx_content_legacy(original_docx: bytes, updated_text: str) -> bytes:
        """
        Update DOCX content using LLM-assisted intelligent editing.
        
        Args:
            original_docx: Original DOCX file bytes
            updated_text: Updated text content from LLM
            
        Returns:
            Updated DOCX file bytes
        """
        try:
            # Load original document
            doc = Document(io.BytesIO(original_docx))
            
            # Extract current document content for LLM analysis
            current_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    current_content.append(para.text.strip())
            
            current_doc_text = '\n'.join(current_content)
            
            # Use LLM to generate specific update instructions
            from services.llm_service import LLMService
            from utils.config import get_settings
            
            settings = get_settings()
            llm_service = LLMService(settings.ai.dict())
            
            # Create a prompt for the LLM to analyze and generate update instructions
            update_prompt = f"""
You are an expert document editor. I need you to analyze the current DOCX document content and the new information, then provide specific update instructions.

CURRENT DOCUMENT CONTENT:
{current_doc_text[:2000]}...

NEW INFORMATION TO INTEGRATE:
{updated_text[:1000]}...

Please analyze both and provide specific update instructions in this exact JSON format:
{{
    "updates": [
        {{
            "action": "replace_text",
            "search": "exact text to find",
            "replace": "new text to replace with"
        }},
        {{
            "action": "update_version",
            "old_version": "8.3.dev",
            "new_version": "8.4"
        }}
    ]
}}

CRITICAL INSTRUCTIONS:
1. Use ONLY "replace_text" and "update_version" actions - NO "add_after"
2. You MUST update ALL relevant sections - be comprehensive, not minimal
3. For Key Features: Find the existing features list and REPLACE it with updated list including new features
4. For Utility Functions: Find the existing functions list and REPLACE it with updated list including new functions  
5. For Examples: Find the existing examples list and REPLACE it with updated list including new examples
6. Always REPLACE entire sections, never append
7. Preserve existing formatting and structure
8. IMPORTANT: The commit adds Interactive CLI Builder - this MUST be reflected in Key Features, Utility Functions, and Examples sections

Return ONLY the JSON, no other text.
"""
            
            # Get LLM response
            llm_response = llm_service._call_llm(update_prompt)
            
            # Parse LLM response
            import json
            try:
                # Extract JSON from response
                json_start = llm_response.find('{')
                json_end = llm_response.rfind('}') + 1
                if json_start != -1 and json_end != -1:
                    json_str = llm_response[json_start:json_end]
                    update_instructions = json.loads(json_str)
                else:
                    logger.warning("[WARNING] Could not parse LLM response as JSON, using fallback")
                    update_instructions = {"updates": []}
            except Exception as e:
                logger.warning(f"[WARNING] Failed to parse LLM response: {str(e)}, using fallback")
                update_instructions = {"updates": []}
            
            # Apply the update instructions
            updates_applied = []
            for update in update_instructions.get("updates", []):
                action = update.get("action")
                
                if action == "replace_text":
                    search = update.get("search", "")
                    replace = update.get("replace", "")
                    if search and replace:
                        for para in doc.paragraphs:
                            if search in para.text:
                                para.text = para.text.replace(search, replace)
                                updates_applied.append(f"replaced: {search[:50]}...")
                                break
                
                
                elif action == "update_version":
                    old_version = update.get("old_version", "")
                    new_version = update.get("new_version", "")
                    if old_version and new_version:
                        for para in doc.paragraphs:
                            if old_version in para.text:
                                para.text = para.text.replace(old_version, new_version)
                                updates_applied.append(f"version: {old_version} → {new_version}")
                                break
            
            # Save updated document
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            result_bytes = doc_bytes.getvalue()
            logger.info(f"[SUCCESS] Updated DOCX document with {len(result_bytes)} bytes (LLM-assisted: {', '.join(updates_applied)})")
            return result_bytes
            
        except Exception as e:
            logger.error(f"[ERROR] Failed to update DOCX content: {str(e)}")
            return original_docx  # Return original if update fails
    
    @staticmethod
    def is_docx_file(file_path: str) -> bool:
        """
        Check if file is a DOCX file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if file is DOCX, False otherwise
        """
        return file_path.lower().endswith('.docx')
    
    @staticmethod
    def get_file_extension(file_path: str) -> str:
        """
        Get file extension from path.
        
        Args:
            file_path: Path to the file
            
        Returns:
            File extension (e.g., '.docx', '.md')
        """
        return file_path.lower().split('.')[-1] if '.' in file_path else ''
    
    # =============================================================================
    # SECTION-LEVEL EDITING METHODS
    # =============================================================================
    
    @staticmethod
    def extract_document_structure(docx_bytes: bytes) -> Dict[str, Any]:
        """
        Extract DOCX document structure with section IDs for targeted editing.
        
        Args:
            docx_bytes: Raw DOCX file content as bytes
            
        Returns:
            Dictionary containing document structure with paragraphs, tables, and metadata
        """
        try:
            doc = Document(io.BytesIO(docx_bytes))
            
            structure = {
                'file_type': 'docx',
                'paragraphs': [],
                'tables': [],
                'headings': [],
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables)
            }
            
            # Extract paragraphs with IDs and formatting
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():  # Skip empty paragraphs
                    para_info = {
                        'paragraph_id': f"para_{i}",
                        'index': i,
                        'text': paragraph.text,
                        'formatted_text': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else None,
                        'is_heading': False,
                        'heading_level': 0,
                        'formatting': DOCXHandler._extract_paragraph_formatting(paragraph)
                    }
                    
                    # Check if it's a heading
                    if paragraph.style and 'heading' in paragraph.style.name.lower():
                        para_info['is_heading'] = True
                        para_info['heading_level'] = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
                        structure['headings'].append(para_info)
                    
                    structure['paragraphs'].append(para_info)
            
            # Extract tables with IDs and structure
            for i, table in enumerate(doc.tables):
                table_info = {
                    'table_id': f"table_{i}",
                    'index': i,
                    'rows': len(table.rows),
                    'columns': len(table.columns) if table.rows else 0,
                    'cells': [],
                    'style': table.style.name if table.style else None
                }
                
                # Extract cell information
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_info = {
                            'row': row_idx,
                            'column': col_idx,
                            'cell_id': f"table_{i}_cell_{row_idx}_{col_idx}",
                            'value': cell.text,
                            'formatted_value': cell.text,
                            'formatting': DOCXHandler._extract_cell_formatting(cell)
                        }
                        table_info['cells'].append(cell_info)
                
                structure['tables'].append(table_info)
            
            logger.info(f"[SUCCESS] Extracted DOCX structure: {len(structure['paragraphs'])} paragraphs, {len(structure['tables'])} tables")
            return structure
            
        except Exception as e:
            logger.error(f"[ERROR] Failed to extract DOCX structure: {str(e)}")
            return {'file_type': 'docx', 'error': str(e)}
    
    @staticmethod
    def apply_targeted_updates(docx_bytes: bytes, updates: List[Dict[str, Any]]) -> bytes:
        """
        Apply targeted updates to specific sections while preserving formatting.
        
        Args:
            docx_bytes: Original DOCX file content as bytes
            updates: List of update dictionaries with section info
            
        Returns:
            Updated DOCX file content as bytes
        """
        try:
            doc = Document(io.BytesIO(docx_bytes))
            
            for update in updates:
                DOCXHandler._apply_single_targeted_update(doc, update)
            
            # Save to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            result_bytes = doc_bytes.getvalue()
            logger.info(f"[SUCCESS] Applied {len(updates)} targeted DOCX updates: {len(result_bytes)} bytes")
            return result_bytes
            
        except Exception as e:
            logger.error(f"[ERROR] Failed to apply targeted DOCX updates: {str(e)}")
            return docx_bytes  # Return original if update fails
    
    @staticmethod
    def _apply_single_targeted_update(doc: Document, update: Dict[str, Any]):
        """Apply a single targeted update to the document."""
        try:
            section_type = update.get('section_type')
            
            if section_type == 'paragraph':
                DOCXHandler._update_targeted_paragraph(doc, update)
            elif section_type == 'table':
                DOCXHandler._update_targeted_table(doc, update)
            elif section_type == 'heading':
                DOCXHandler._update_targeted_heading(doc, update)
            else:
                logger.warning(f"[WARNING] Unknown section type: {section_type}")
                
        except Exception as e:
            logger.error(f"[ERROR] Failed to apply targeted update: {str(e)}")
    
    @staticmethod
    def _update_targeted_paragraph(doc: Document, update: Dict[str, Any]):
        """Update a specific paragraph while preserving formatting."""
        try:
            paragraph_index = update.get('paragraph_index')
            new_content = update.get('new_content')
            
            if paragraph_index is None or paragraph_index >= len(doc.paragraphs):
                logger.error(f"[ERROR] Invalid paragraph index: {paragraph_index}")
                return
            
            paragraph = doc.paragraphs[paragraph_index]
            
            # Store original formatting
            original_runs = list(paragraph.runs)
            original_formatting = DOCXHandler._extract_paragraph_formatting(paragraph)
            
            # Clear paragraph content
            paragraph.clear()
            
            # Add new content while preserving formatting
            if original_runs:
                # Use formatting from first run as template
                template_run = original_runs[0]
                new_run = paragraph.add_run(new_content)
                
                # Apply template formatting
                new_run.font.name = template_run.font.name
                new_run.font.size = template_run.font.size
                new_run.font.bold = template_run.font.bold
                new_run.font.italic = template_run.font.italic
                if template_run.font.color and template_run.font.color.rgb:
                    new_run.font.color.rgb = template_run.font.color.rgb
            else:
                # No original runs, add plain text
                paragraph.add_run(new_content)
            
            logger.info(f"[SUCCESS] Updated paragraph {paragraph_index}")
            
        except Exception as e:
            logger.error(f"[ERROR] Failed to update paragraph: {str(e)}")
    
    @staticmethod
    def _update_targeted_table(doc: Document, update: Dict[str, Any]):
        """Update specific table cells while preserving formatting."""
        try:
            table_index = update.get('table_index')
            row_index = update.get('row_index')
            col_index = update.get('col_index')
            new_value = update.get('new_value')
            
            if table_index is None or table_index >= len(doc.tables):
                logger.error(f"[ERROR] Invalid table index: {table_index}")
                return
            
            table = doc.tables[table_index]
            
            if row_index is None or col_index is None:
                logger.error(f"[ERROR] Invalid cell coordinates: row={row_index}, col={col_index}")
                return
            
            if row_index >= len(table.rows) or col_index >= len(table.rows[row_index].cells):
                logger.error(f"[ERROR] Cell coordinates out of range")
                return
            
            cell = table.rows[row_index].cells[col_index]
            
            # Store original formatting
            original_formatting = DOCXHandler._extract_cell_formatting(cell)
            
            # Clear cell content
            cell.text = ""
            
            # Add new content while preserving formatting
            if cell.paragraphs:
                para = cell.paragraphs[0]
                para.clear()
                new_run = para.add_run(new_value)
                
                # Apply original formatting if available
                if original_formatting:
                    DOCXHandler._apply_cell_formatting(new_run, original_formatting)
            else:
                cell.text = new_value
            
            logger.info(f"[SUCCESS] Updated table cell [{table_index}][{row_index}][{col_index}]")
            
        except Exception as e:
            logger.error(f"[ERROR] Failed to update table cell: {str(e)}")
    
    @staticmethod
    def _update_targeted_heading(doc: Document, update: Dict[str, Any]):
        """Update a heading while preserving formatting."""
        try:
            paragraph_index = update.get('paragraph_index')
            new_content = update.get('new_content')
            heading_level = update.get('heading_level', 1)
            
            if paragraph_index is None or paragraph_index >= len(doc.paragraphs):
                logger.error(f"[ERROR] Invalid heading index: {paragraph_index}")
                return
            
            paragraph = doc.paragraphs[paragraph_index]
            
            # Store original formatting
            original_formatting = DOCXHandler._extract_paragraph_formatting(paragraph)
            
            # Update heading content
            paragraph.clear()
            new_run = paragraph.add_run(new_content)
            
            # Apply heading formatting
            new_run.font.size = Inches(18/72) if heading_level == 1 else Inches(16/72) if heading_level == 2 else Inches(14/72)
            new_run.bold = True
            
            # Apply original formatting if available
            if original_formatting:
                DOCXHandler._apply_paragraph_formatting(new_run, original_formatting)
            
            logger.info(f"[SUCCESS] Updated heading {paragraph_index}")
            
        except Exception as e:
            logger.error(f"[ERROR] Failed to update heading: {str(e)}")
    
    @staticmethod
    def _extract_paragraph_formatting(paragraph) -> Dict[str, Any]:
        """Extract formatting information from a paragraph."""
        try:
            formatting = {
                'alignment': paragraph.alignment,
                'style': paragraph.style.name if paragraph.style else None,
                'runs': []
            }
            
            for run in paragraph.runs:
                run_formatting = {
                    'font_name': run.font.name,
                    'font_size': run.font.size.pt if run.font.size else None,
                    'bold': run.font.bold,
                    'italic': run.font.italic,
                    'underline': run.underline,
                    'color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None
                }
                formatting['runs'].append(run_formatting)
            
            return formatting
        except Exception as e:
            logger.warning(f"[WARNING] Could not extract paragraph formatting: {str(e)}")
            return {}
    
    @staticmethod
    def _extract_cell_formatting(cell) -> Dict[str, Any]:
        """Extract formatting information from a table cell."""
        try:
            formatting = {}
            
            if cell.paragraphs:
                para = cell.paragraphs[0]
                formatting = DOCXHandler._extract_paragraph_formatting(para)
            
            return formatting
        except Exception as e:
            logger.warning(f"[WARNING] Could not extract cell formatting: {str(e)}")
            return {}
    
    @staticmethod
    def _apply_paragraph_formatting(run, formatting: Dict[str, Any]):
        """Apply formatting to a run."""
        try:
            if 'runs' in formatting and formatting['runs']:
                template_run = formatting['runs'][0]
                
                if template_run.get('font_name'):
                    run.font.name = template_run['font_name']
                if template_run.get('font_size'):
                    run.font.size = Inches(template_run['font_size']/72)
                if template_run.get('bold') is not None:
                    run.font.bold = template_run['bold']
                if template_run.get('italic') is not None:
                    run.font.italic = template_run['italic']
                if template_run.get('underline') is not None:
                    run.underline = template_run['underline']
                if template_run.get('color'):
                    run.font.color.rgb = RGBColor.from_string(template_run['color'].replace('#', ''))
        except Exception as e:
            logger.warning(f"[WARNING] Could not apply paragraph formatting: {str(e)}")
    
    @staticmethod
    def _apply_cell_formatting(run, formatting: Dict[str, Any]):
        """Apply formatting to a cell run."""
        try:
            DOCXHandler._apply_paragraph_formatting(run, formatting)
        except Exception as e:
            logger.warning(f"[WARNING] Could not apply cell formatting: {str(e)}")
    
    @staticmethod
    def create_paragraph_update(paragraph_index: int, new_content: str, 
                               reason: str = None) -> Dict[str, Any]:
        """Create a paragraph update dictionary."""
        return {
            'section_id': f"para_{paragraph_index}",
            'section_type': 'paragraph',
            'paragraph_index': paragraph_index,
            'new_content': new_content,
            'reason': reason
        }
    
    @staticmethod
    def create_table_update(table_index: int, row_index: int, col_index: int, 
                           new_value: str, reason: str = None) -> Dict[str, Any]:
        """Create a table cell update dictionary."""
        return {
            'section_id': f"table_{table_index}_cell_{row_index}_{col_index}",
            'section_type': 'table',
            'table_index': table_index,
            'row_index': row_index,
            'col_index': col_index,
            'new_value': new_value,
            'reason': reason
        }
    
    @staticmethod
    def create_heading_update(paragraph_index: int, new_content: str, 
                             heading_level: int = 1, reason: str = None) -> Dict[str, Any]:
        """Create a heading update dictionary."""
        return {
            'section_id': f"heading_{paragraph_index}",
            'section_type': 'heading',
            'paragraph_index': paragraph_index,
            'new_content': new_content,
            'heading_level': heading_level,
            'reason': reason
        }