"""
DOCX Line Editor Module

Handles editing existing lines/paragraphs in DOCX documents while preserving formatting.
Based on the workflow from document_editor (fixed).py but integrated with the LLM service.
"""

import io
import logging
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)


class DocxLineEditor:
    """Handles line/paragraph editing in DOCX documents with formatting preservation."""
    
    def __init__(self):
        self.logger = logger
    
    def edit_document_lines(self, docx_bytes: bytes, commit_context: Dict[str, Any], llm_service) -> bytes:
        """
        Main entry point for editing lines in a DOCX document.
        
        Args:
            docx_bytes: Original DOCX file as bytes
            commit_context: Commit information and context
            llm_service: LLM service instance for text generation
            
        Returns:
            Updated DOCX file as bytes
        """
        try:
            # Load document
            doc = Document(io.BytesIO(docx_bytes))
            
            # Get all content items (paragraphs, tables, etc.)
            content_items = self._get_all_content_items(doc)
            
            if not content_items:
                self.logger.warning("No content items found in document")
                return docx_bytes
            
            # Find the best paragraph to edit based on commit context
            target_paragraph = self._find_target_paragraph(content_items, commit_context, llm_service)
            
            if not target_paragraph:
                self.logger.warning("No suitable paragraph found for editing")
                return docx_bytes
            
            # Capture formatting of the target paragraph
            original_formats = self._capture_detailed_formatting(target_paragraph)
            
            # Generate updated text with formatting preservation
            updated_text = self._rephrase_text_with_formatting(
                target_paragraph.text, 
                original_formats, 
                llm_service, 
                commit_context
            )
            
            # Apply the updated text with formatting
            self._replace_paragraph_text(target_paragraph, updated_text, original_formats)
            
            # Save updated document
            output = io.BytesIO()
            doc.save(output)
            return output.getvalue()
            
        except Exception as e:
            self.logger.error(f"Error editing document lines: {e}")
            return docx_bytes
    
    def _get_all_content_items(self, doc: Document) -> List[Any]:
        """Get all content items from the document."""
        content_items = []
        
        # Add paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only non-empty paragraphs
                content_items.append(paragraph)
        
        # Add table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            content_items.append(paragraph)
        
        return content_items
    
    def _find_target_paragraph(self, content_items: List[Any], commit_context: Dict[str, Any], llm_service) -> Optional[Any]:
        """
        Find the best paragraph to edit based on commit context.
        Uses LLM to determine which paragraph is most relevant.
        """
        try:
            # Create a prompt to help LLM select the best paragraph
            paragraphs_text = []
            for i, item in enumerate(content_items):
                paragraphs_text.append(f"{i}: {item.text[:200]}...")  # First 200 chars
            
            prompt = f"""
            Based on this commit information, select the paragraph number that should be updated:
            
            Commit Message: {commit_context.get('message', '')}
            Files Changed: {', '.join(commit_context.get('files', []))}
            
            Available paragraphs:
            {chr(10).join(paragraphs_text)}
            
            Return only the paragraph number (0-based index) that should be updated.
            If no paragraph needs updating, return -1.
            """
            
            response = llm_service.generate_response(prompt)
            
            # Parse response to get paragraph index
            try:
                paragraph_index = int(response.strip())
                if 0 <= paragraph_index < len(content_items):
                    return content_items[paragraph_index]
            except ValueError:
                pass
            
            # Fallback: return first paragraph if LLM selection fails
            return content_items[0] if content_items else None
            
        except Exception as e:
            self.logger.error(f"Error finding target paragraph: {e}")
            return content_items[0] if content_items else None
    
    def _capture_detailed_formatting(self, paragraph) -> List[Dict[str, Any]]:
        """
        Capture detailed formatting information for each run in the paragraph.
        Returns a list of formatting dictionaries, one per run.
        """
        formats = []
        
        for run in paragraph.runs:
            format_info = {
                'text': run.text,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'color': None,
                'highlight': None
            }
            
            # Capture font color
            if run.font.color.rgb:
                format_info['color'] = run.font.color.rgb
            
            # Capture highlight color
            if hasattr(run.font, 'highlight_color') and run.font.highlight_color:
                format_info['highlight'] = run.font.highlight_color
            
            formats.append(format_info)
        
        return formats
    
    def _rephrase_text_with_formatting(self, text: str, original_formats: List[Dict[str, Any]], 
                                     llm_service, commit_context: Dict[str, Any]) -> str:
        """
        Use LLM to rephrase text while preserving formatting structure.
        Returns text with formatting tags that can be parsed.
        """
        try:
            # Create formatting-aware prompt
            prompt = f"""
            Update the following text based on the commit information while preserving the formatting structure.
            Use these formatting tags to indicate formatting:
            - <BOLD>text</BOLD> for bold text
            - <ITALIC>text</ITALIC> for italic text  
            - <UNDERLINE>text</UNDERLINE> for underlined text
            
            Original text: {text}
            
            Commit Message: {commit_context.get('message', '')}
            Files Changed: {', '.join(commit_context.get('files', []))}
            
            Update the text to reflect the changes described in the commit. 
            Preserve the original formatting by using the tags above.
            Return only the updated text with formatting tags.
            """
            
            response = llm_service.generate_response(prompt)
            return response.strip()
            
        except Exception as e:
            self.logger.error(f"Error generating formatted text: {e}")
            return text  # Return original text as fallback
    
    def _parse_formatted_response(self, llm_response: str, original_formats: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Parse LLM response with formatting tags into segments.
        Returns list of segments with text and formatting information.
        """
        segments = []
        text = llm_response
        
        # Simple parsing for formatting tags
        while text:
            # Find next formatting tag
            bold_start = text.find('<BOLD>')
            italic_start = text.find('<ITALIC>')
            underline_start = text.find('<UNDERLINE>')
            
            # Find the earliest tag
            tag_positions = []
            if bold_start != -1:
                tag_positions.append(('bold', bold_start))
            if italic_start != -1:
                tag_positions.append(('italic', italic_start))
            if underline_start != -1:
                tag_positions.append(('underline', underline_start))
            
            if not tag_positions:
                # No more tags, add remaining text as plain
                if text.strip():
                    segments.append({'text': text, 'formatting': {}})
                break
            
            # Sort by position
            tag_positions.sort(key=lambda x: x[1])
            tag_type, tag_start = tag_positions[0]
            
            # Add text before tag as plain
            if tag_start > 0:
                plain_text = text[:tag_start]
                if plain_text.strip():
                    segments.append({'text': plain_text, 'formatting': {}})
            
            # Extract formatted text
            if tag_type == 'bold':
                end_tag = '</BOLD>'
                tag_len = len('<BOLD>')
            elif tag_type == 'italic':
                end_tag = '</ITALIC>'
                tag_len = len('<ITALIC>')
            else:  # underline
                end_tag = '</UNDERLINE>'
                tag_len = len('<UNDERLINE>')
            
            end_pos = text.find(end_tag, tag_start + tag_len)
            if end_pos != -1:
                formatted_text = text[tag_start + tag_len:end_pos]
                formatting = {tag_type: True}
                segments.append({'text': formatted_text, 'formatting': formatting})
                
                # Continue with remaining text
                text = text[end_pos + len(end_tag):]
            else:
                # Malformed tag, treat as plain text
                text = text[tag_start + tag_len:]
        
        return segments
    
    def _replace_paragraph_text(self, paragraph, new_text_or_segments, original_formats: List[Dict[str, Any]]):
        """
        Replace paragraph text while preserving formatting.
        Uses safe clearing approach - clears runs but doesn't remove them.
        """
        try:
            # If new_text_or_segments is a string, parse it for formatting tags
            if isinstance(new_text_or_segments, str):
                segments = self._parse_formatted_response(new_text_or_segments, original_formats)
            else:
                segments = new_text_or_segments
            
            # Clear existing runs (safe clearing)
            for run in paragraph.runs:
                run.text = ""
            
            # Apply new text with formatting
            run_index = 0
            for segment in segments:
                text = segment['text']
                formatting = segment.get('formatting', {})
                
                # Use existing run if available, otherwise create new one
                if run_index < len(paragraph.runs):
                    run = paragraph.runs[run_index]
                else:
                    run = paragraph.add_run()
                
                # Set text
                run.text = text
                
                # Apply formatting
                if formatting.get('bold'):
                    run.bold = True
                if formatting.get('italic'):
                    run.italic = True
                if formatting.get('underline'):
                    run.underline = True
                
                # Apply color and other formatting from original if available
                if run_index < len(original_formats):
                    orig_format = original_formats[run_index]
                    if orig_format.get('color'):
                        run.font.color.rgb = orig_format['color']
                    if orig_format.get('highlight'):
                        run.font.highlight_color = orig_format['highlight']
                
                run_index += 1
            
            # Clear any remaining unused runs
            for i in range(run_index, len(paragraph.runs)):
                paragraph.runs[i].text = ""
                
        except Exception as e:
            self.logger.error(f"Error replacing paragraph text: {e}")
            # Fallback: apply text without formatting
            self._apply_proportional_formatting_safe(paragraph, new_text_or_segments, original_formats)
    
    def _apply_proportional_formatting_safe(self, paragraph, new_text: str, original_formats: List[Dict[str, Any]]):
        """
        Fallback method to apply text with proportional formatting distribution.
        Used when LLM formatting tags are not available or parsing fails.
        """
        try:
            # Clear existing runs
            for run in paragraph.runs:
                run.text = ""
            
            # Distribute text proportionally across runs
            if not original_formats:
                # No original formatting, just set text
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)
                return
            
            # Calculate proportional distribution
            total_original_length = sum(len(fmt['text']) for fmt in original_formats)
            if total_original_length == 0:
                total_original_length = 1  # Avoid division by zero
            
            text_position = 0
            for i, orig_format in enumerate(original_formats):
                if i >= len(paragraph.runs):
                    run = paragraph.add_run()
                else:
                    run = paragraph.runs[i]
                
                # Calculate proportional length
                proportion = len(orig_format['text']) / total_original_length
                segment_length = int(len(new_text) * proportion)
                
                # Get text segment
                if i == len(original_formats) - 1:
                    # Last segment gets remaining text
                    segment_text = new_text[text_position:]
                else:
                    segment_text = new_text[text_position:text_position + segment_length]
                
                # Set text and formatting
                run.text = segment_text
                run.bold = orig_format.get('bold', False)
                run.italic = orig_format.get('italic', False)
                run.underline = orig_format.get('underline', False)
                
                if orig_format.get('color'):
                    run.font.color.rgb = orig_format['color']
                if orig_format.get('highlight'):
                    run.font.highlight_color = orig_format['highlight']
                
                text_position += segment_length
                
        except Exception as e:
            self.logger.error(f"Error applying proportional formatting: {e}")
            # Ultimate fallback: just set the text
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)

