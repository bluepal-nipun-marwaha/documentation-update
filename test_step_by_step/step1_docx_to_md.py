#!/usr/bin/env python3
"""
Step 1: DOCX to Markdown Conversion with LLM Text Updates
- Converts DOCX to Markdown while preserving paragraph metadata
- Removes tables completely from the Markdown
- Uses LLM to update text content based on commit diff
- Outputs updated Markdown and metadata files
"""

import os
import sys
import json
import requests
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple

# Add parent directory to path to import our modules
sys.path.append(str(Path(__file__).parent.parent))

from services.llm_service import LLMService
from utils.config import get_settings
from utils.docx_handler import DOCXHandler
from docx.enum.style import WD_STYLE_TYPE
import re
import structlog

logger = structlog.get_logger(__name__)

def safe_print(text: str, prefix: str = ""):
    """Safely print text that might contain Unicode characters."""
    try:
        print(f"{prefix}{text}")
    except UnicodeEncodeError:
        safe_text = text.encode('ascii', 'replace').decode('ascii')
        print(f"{prefix}{safe_text}")

class Step1Processor:
    """Step 1: DOCX to Markdown with LLM text updates."""
    
    def __init__(self, input_dir=None, output_dir=None, llm_service=None):
        if input_dir and output_dir:
            # Use provided directories
            self.input_dir = Path(input_dir)
            self.output_dir = Path(output_dir)
        else:
            # Use default test directories
            self.test_dir = Path(__file__).parent
            self.input_dir = self.test_dir / "input"
            self.output_dir = self.test_dir / "output"
        
        # Create directories
        self.input_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize LLM service (use provided service or create new one)
        if llm_service:
            self.llm_service = llm_service
        else:
            settings = get_settings()
            self.llm_service = LLMService(settings.ai.model_dump())
        
        # Initialize DOCX handler
        self.docx_handler = DOCXHandler()
        
        print("Step 1: DOCX to Markdown Conversion with LLM Text Updates")
        print(f"Input Directory: {self.input_dir}")
        print(f"Output Directory: {self.output_dir}")
    
    def fetch_commit_data(self, repo_url: str, token: str, before_commit: str, after_commit: str) -> Dict[str, Any]:
        """Fetch commit data from GitHub API."""
        try:
            # Extract owner and repo from URL
            repo_path = repo_url.replace('https://github.com/', '').replace('.git', '')
            owner, repo = repo_path.split('/')
            
            print(f"Fetching commit data from {owner}/{repo}")
            
            # Fetch the after commit (the one we're analyzing)
            commit_url = f"https://api.github.com/repos/{owner}/{repo}/commits/{after_commit}"
            headers = {
                'Authorization': f'Bearer {token}',
                'Accept': 'application/vnd.github.v3+json'
            }
            
            response = requests.get(commit_url, headers=headers)
            if response.status_code != 200:
                raise Exception(f"Failed to fetch commit: {response.status_code}")
            
            commit_data = response.json()
            
            # Also fetch the diff
            diff_url = f"https://api.github.com/repos/{owner}/{repo}/commits/{after_commit}"
            diff_headers = {
                'Authorization': f'Bearer {token}',
                'Accept': 'application/vnd.github.v3.diff'
            }
            
            diff_response = requests.get(diff_url, headers=diff_headers)
            diff_content = diff_response.text if diff_response.status_code == 200 else ""
            
            # Extract file changes
            files = commit_data.get('files', [])
            modified_files = [f['filename'] for f in files if f['status'] == 'modified']
            added_files = [f['filename'] for f in files if f['status'] == 'added']
            removed_files = [f['filename'] for f in files if f['status'] == 'removed']
            
            commit_info = {
                'hash': commit_data['sha'],
                'message': commit_data['commit']['message'],
                'author': commit_data['commit']['author']['name'],
                'timestamp': commit_data['commit']['author']['date'],
                'url': commit_data['html_url'],
                'modified': modified_files,
                'added': added_files,
                'removed': removed_files,
                'diff': diff_content,
                'files': files
            }
            
            print(f"Commit data fetched successfully")
            print(f"Message: {commit_info['message']}")
            print(f"Author: {commit_info['author']}")
            
            return commit_info
            
        except Exception as e:
            print(f"Error fetching commit data: {str(e)}")
            return {}
    
    def convert_docx_to_markdown_with_metadata(self, docx_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Convert DOCX to Markdown while preserving paragraph metadata and removing tables."""
        try:
            print(f"Converting DOCX to Markdown: {docx_path.name}")
            
            # Read the DOCX file as bytes
            with open(docx_path, 'rb') as f:
                docx_content = f.read()
            
            # Use DOCX handler to convert to markdown
            markdown_content, metadata = self.docx_handler.docx_to_markdown_with_metadata(docx_content)
            
            # Extract detailed formatting information from the original DOCX
            formatting_info = self.extract_formatting_from_docx(docx_path)
            
            # Extract line-by-line metadata
            line_metadata = self.extract_line_by_line_metadata(docx_path)
            
            # Aggressively remove tables from markdown
            markdown_content = self.remove_tables_from_markdown(markdown_content)
            
            print(f"Markdown conversion completed")
            print(f"Original length: {len(markdown_content)} characters")
            
            return markdown_content, {
                'formatting_info': formatting_info,
                'line_metadata': line_metadata
            }
            
        except Exception as e:
            print(f"Error converting DOCX to Markdown: {str(e)}")
            return "", {}
    
    def extract_formatting_from_docx(self, docx_path: Path) -> Dict[str, Any]:
        """Extract detailed formatting information from the original DOCX."""
        try:
            print("Extracting formatting information from original DOCX...")
            
            from docx import Document
            
            doc = Document(docx_path)
            formatting_info = {
                'document_styles': {},
                'paragraph_formats': [],
                'heading_styles': {},
                'font_info': {},
                'colors': {},
                'layout_info': {}
            }
            
            # Extract document-level formatting
            formatting_info['font_info']['default_font'] = doc.styles['Normal'].font.name
            formatting_info['font_info']['default_size'] = doc.styles['Normal'].font.size.pt if doc.styles['Normal'].font.size else 11
            
            # Extract paragraph formatting
            for i, paragraph in enumerate(doc.paragraphs):
                para_info = {
                    'index': i,
                    'text': paragraph.text,
                    'style': paragraph.style.name if paragraph.style else 'Normal',
                    'alignment': str(paragraph.alignment) if paragraph.alignment else None,
                    'runs': []
                }
                
                # Extract run formatting
                for run in paragraph.runs:
                    run_info = {
                        'text': run.text,
                        'font_name': run.font.name,
                        'font_size': run.font.size.pt if run.font.size else None,
                        'bold': run.font.bold,
                        'italic': run.font.italic,
                        'color': str(run.font.color.rgb) if run.font.color.rgb else None,
                        'underline': run.font.underline
                    }
                    para_info['runs'].append(run_info)
                
                formatting_info['paragraph_formats'].append(para_info)
            
            # Extract heading styles
            for style in doc.styles:
                if style.type == WD_STYLE_TYPE.PARAGRAPH:
                    style_info = {
                        'name': style.name,
                        'font_name': style.font.name,
                        'font_size': style.font.size.pt if style.font.size else None,
                        'bold': style.font.bold,
                        'italic': style.font.italic,
                        'color': str(style.font.color.rgb) if style.font.color.rgb else None
                    }
                    formatting_info['heading_styles'][style.name] = style_info
            
            print(f"Extracted formatting for {len(formatting_info['paragraph_formats'])} paragraphs")
            print(f"Extracted {len(formatting_info['heading_styles'])} heading styles")
            
            return formatting_info
            
        except Exception as e:
            print(f"Error extracting formatting from DOCX: {str(e)}")
            return {}
    
    def extract_line_by_line_metadata(self, docx_path: Path) -> List[Dict[str, Any]]:
        """Extract line-by-line formatting metadata from DOCX."""
        try:
            print("Extracting line-by-line metadata from original DOCX...")
            
            from docx import Document
            
            doc = Document(docx_path)
            line_metadata = []
            
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if not paragraph.text.strip():
                    continue
                    
                # Extract paragraph-level formatting
                para_info = {
                    'paragraph_index': para_idx,
                    'text': paragraph.text,
                    'style_name': paragraph.style.name if paragraph.style else 'Normal',
                    'alignment': str(paragraph.alignment) if paragraph.alignment else None,
                    'runs': []
                }
                
                # Extract run-level formatting for each run in the paragraph
                for run_idx, run in enumerate(paragraph.runs):
                    if not run.text.strip():
                        continue
                        
                    run_info = {
                        'run_index': run_idx,
                        'text': run.text,
                        'font_name': run.font.name,
                        'font_size': run.font.size.pt if run.font.size else None,
                        'bold': run.font.bold,
                        'italic': run.font.italic,
                        'color': str(run.font.color.rgb) if run.font.color.rgb else None,
                        'underline': run.font.underline,
                        'strike': run.font.strike,
                        'subscript': run.font.subscript,
                        'superscript': run.font.superscript
                    }
                    para_info['runs'].append(run_info)
                
                line_metadata.append(para_info)
            
            print(f"Extracted metadata for {len(line_metadata)} lines")
            return line_metadata
            
        except Exception as e:
            print(f"Error extracting line-by-line metadata: {str(e)}")
            return []
    
    def remove_tables_from_markdown(self, markdown_content: str) -> str:
        """Aggressively remove tables from markdown content."""
        try:
            lines = markdown_content.split('\n')
            filtered_lines = []
            in_table = False
            
            for line in lines:
                # Check if line looks like a table row (contains |)
                if '|' in line and line.strip():
                    # Skip table lines
                    continue
                
                # Check if line looks like table header separator
                if re.match(r'^\s*\|?[\s\-:]+\|[\s\-:]*\|?\s*$', line):
                    # Skip table separator lines
                    continue
                
                # Check for table-related patterns
                if re.match(r'^\s*\|.*\|.*\|', line):
                    # Skip table rows
                    continue
                
                # Add non-table lines
                filtered_lines.append(line)
            
            # Join lines and clean up extra newlines
            result = '\n'.join(filtered_lines)
            result = re.sub(r'\n\s*\n\s*\n', '\n\n', result)  # Remove excessive newlines
            
            print(f"Tables removed from markdown")
            print(f"Filtered length: {len(result)} characters")
            
            return result
            
        except Exception as e:
            print(f"Error removing tables from markdown: {str(e)}")
            return markdown_content
    
    def extract_diff_summary(self, commit_info: Dict[str, Any]) -> str:
        """Extract and summarize the diff content."""
        try:
            print("Extracting and summarizing diff content...")
            
            # Create prompt for diff summarization
            prompt = f"""
You are a technical code reviewer. I need you to analyze this commit diff and provide a comprehensive summary.

COMMIT INFORMATION:
- Message: {commit_info['message']}
- Author: {commit_info['author']}
- Modified Files: {', '.join(commit_info['modified'])}
- Added Files: {', '.join(commit_info['added'])}
- Removed Files: {', '.join(commit_info['removed'])}

DIFF CONTENT:
{commit_info.get('diff', 'No diff available')}

INSTRUCTIONS:
1. **ANALYZE THE CHANGES**: Carefully examine what was added, modified, or removed
2. **IDENTIFY KEY CHANGES**: Focus on the most significant changes that would affect documentation
3. **CATEGORIZE CHANGES**: Group changes by type (new features, bug fixes, API changes, etc.)
4. **EXTRACT IMPACT**: Determine which parts of documentation would need updates
5. **PROVIDE SUMMARY**: Give a clear, structured summary of the changes

Please provide a detailed summary that will help determine which documentation sections need updates.
"""
            
            # Get diff summary from LLM
            diff_summary = self.llm_service._call_llm(prompt, temperature=0.2, max_tokens=2000)
            
            print("Diff summary extracted successfully")
            safe_print(diff_summary[:300] + "..." if len(diff_summary) > 300 else diff_summary)
            
            return diff_summary
            
        except Exception as e:
            print(f"Error extracting diff summary: {str(e)}")
            return commit_info.get('diff', 'No diff available')
    
    def identify_sections_to_update(self, markdown_content: str, diff_summary: str) -> List[str]:
        """Identify which sections/headings need updates based on diff summary."""
        try:
            print("Analyzing documentation structure for contextual updates...")
            
            # Extract headings from markdown
            headings = []
            lines = markdown_content.split('\n')
            for line in lines:
                if line.strip().startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    title = line.lstrip('#').strip()
                    headings.append({'level': level, 'title': title, 'line': line})
            
            print(f"Found {len(headings)} headings in documentation")
            
            # Create prompt to analyze where updates would naturally fit
            prompt = f"""
You are a documentation expert. Analyze this diff summary and existing documentation structure to understand where updates would naturally fit.

DIFF SUMMARY:
{diff_summary}

DOCUMENTATION STRUCTURE:
{chr(10).join([f"{'  ' * h['level']}- {h['title']}" for h in headings])}

INSTRUCTIONS:
1. **ANALYZE THE CHANGES**: Understand what new features or modifications were made
2. **MAP TO DOCUMENTATION**: Identify which existing sections would naturally contain information about these changes
3. **CONSIDER CONTEXT**: Think about where a reader would expect to find this information
4. **PRIORITIZE RELEVANCE**: Focus on sections that are most directly related to the changes

Please provide a brief analysis of where the changes would naturally fit in the existing documentation structure.
Focus on understanding the context rather than creating a rigid list.
"""
            
            # Get contextual analysis from LLM
            contextual_analysis = self.llm_service._call_llm(prompt, temperature=0.2, max_tokens=1000)
            
            print("Contextual analysis completed")
            safe_print(contextual_analysis[:300] + "..." if len(contextual_analysis) > 300 else contextual_analysis)
            
            # Return the analysis for reference (not used for rigid filtering)
            return [contextual_analysis]
            
        except Exception as e:
            print(f"Error analyzing documentation structure: {str(e)}")
            return []
    
    def update_markdown_with_llm(self, markdown_content: str, commit_info: Dict[str, Any], diff_summary: str, sections_to_update: List[str]) -> str:
        """Update markdown content using LLM based on diff summary and identified sections."""
        try:
            print("Updating markdown content with LLM...")
            
            # Create a much more conservative prompt that preserves structure
            prompt = f"""
You are a technical documentation expert. I need you to make MINIMAL, TARGETED updates to this existing documentation.

CRITICAL REQUIREMENTS:
1. **PRESERVE EXACT STRUCTURE**: Keep ALL headings, sections, and formatting exactly as they are
2. **MINIMAL CHANGES ONLY**: Only add or modify content where absolutely necessary
3. **NO RESTRUCTURING**: Do not reorganize, reorder, or change the document structure
4. **MAINTAIN ORIGINAL TONE**: Keep the same writing style and tone

COMMIT INFORMATION:
- Message: {commit_info['message']}
- Author: {commit_info['author']}

DIFF SUMMARY:
{diff_summary}

ORIGINAL DOCUMENTATION (PRESERVE STRUCTURE):
{markdown_content}

INSTRUCTIONS:
1. **IDENTIFY RELEVANT SECTIONS**: Find sections that relate to the new interactive builder feature
2. **MAKE TARGETED UPDATES**: Only add brief mentions or enhancements to existing content
3. **PRESERVE EVERYTHING ELSE**: Keep all other content exactly as it is
4. **MAINTAIN HEADINGS**: Do not change any heading text or structure
5. **FOCUS ON INTEGRATION**: Show how the new feature fits into existing sections

EXAMPLE OF WHAT TO DO:
- If there's a "Key Features" section, add a brief mention of the interactive builder
- If there's a "Examples" section, add the interactive builder as an example
- If there's a "Future Roadmap" section, mention the interactive builder

EXAMPLE OF WHAT NOT TO DO:
- Do not rewrite entire sections
- Do not change heading names
- Do not reorganize the document
- Do not add new major sections

Please provide the documentation with MINIMAL, TARGETED updates that preserve the exact original structure.
"""
            
            # Get updated content from LLM
            updated_content = self.llm_service._call_llm(prompt, temperature=0.1, max_tokens=2000)
            
            print("Markdown content updated with LLM")
            safe_print(updated_content[:200] + "..." if len(updated_content) > 200 else updated_content)
            
            return updated_content
            
        except Exception as e:
            print(f"Error updating markdown with LLM: {str(e)}")
            return markdown_content
    
    def save_step1_results(self, updated_markdown: str, formatting_info: Dict[str, Any], diff_summary: str = "", sections_to_update: List[str] = None):
        """Save Step 1 results."""
        try:
            # Save updated markdown
            markdown_path = self.output_dir / "updated_documentation.md"
            with open(markdown_path, 'w', encoding='utf-8') as f:
                f.write(updated_markdown)
            
            # Save formatting information
            formatting_path = self.output_dir / "original_formatting.json"
            with open(formatting_path, 'w', encoding='utf-8') as f:
                json.dump(formatting_info['formatting_info'], f, indent=2, ensure_ascii=False)
            
            # Save line-by-line metadata
            line_metadata_path = self.output_dir / "line_metadata.json"
            with open(line_metadata_path, 'w', encoding='utf-8') as f:
                json.dump(formatting_info['line_metadata'], f, indent=2, ensure_ascii=False)
            
            # Save diff summary
            if diff_summary:
                diff_summary_path = self.output_dir / "diff_summary.txt"
                with open(diff_summary_path, 'w', encoding='utf-8') as f:
                    f.write(diff_summary)
            
            # Save sections to update
            if sections_to_update:
                sections_path = self.output_dir / "contextual_analysis.txt"
                with open(sections_path, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(sections_to_update))
            
            print(f"Step 1 results saved successfully")
            print(f"Updated Markdown: {markdown_path}")
            print(f"Original Formatting: {formatting_path}")
            print(f"Line Metadata: {line_metadata_path}")
            if diff_summary:
                print(f"Diff Summary: {diff_summary_path}")
            if sections_to_update:
                print(f"Contextual Analysis: {sections_path}")
            
            return {
                'markdown_path': markdown_path,
                'formatting_path': formatting_path,
                'line_metadata_path': line_metadata_path,
                'diff_summary_path': diff_summary_path if diff_summary else None,
                'sections_path': sections_path if sections_to_update else None
            }
            
        except Exception as e:
            print(f"Error saving Step 1 results: {str(e)}")
            return {}
    
    def run_step1_integration(self, docx_path: str, commit_info: dict):
        """Run Step 1 for integration with enhanced workflow."""
        try:
            print("Starting Step 1 Integration: DOCX to Markdown Conversion with LLM Text Updates")
            print("=" * 80)
            
            # Step 1: Load DOCX file
            if not Path(docx_path).exists():
                print(f"DOCX file not found: {docx_path}")
                return False
            
            print(f"Loading DOCX file: {docx_path}")
            
            # Step 2: Convert DOCX to Markdown with metadata
            markdown_content, formatting_info = self.convert_docx_to_markdown_with_metadata(Path(docx_path))
            
            if not markdown_content:
                print("Failed to convert DOCX to Markdown")
                return False
            
            print(f"Converted DOCX to Markdown: {len(markdown_content)} characters")
            
            # Step 3: Remove tables from Markdown
            clean_markdown = self.remove_tables_from_markdown(markdown_content)
            print(f"Removed tables from Markdown: {len(clean_markdown)} characters")
            
            # Step 4: Summarize commit diff
            diff_summary = self.extract_diff_summary(commit_info)
            print(f"Generated diff summary: {len(diff_summary)} characters")
            
            # Step 5: Identify sections to update
            sections_to_update = self.identify_sections_to_update(clean_markdown, diff_summary)
            
            # Step 6: Update Markdown with LLM
            updated_markdown = self.update_markdown_with_llm(clean_markdown, commit_info, diff_summary, sections_to_update)
            
            if not updated_markdown:
                print("Failed to update Markdown with LLM")
                return False
            
            print(f"Updated Markdown with LLM: {len(updated_markdown)} characters")
            
            # Step 7: Save results
            self.save_step1_results(updated_markdown, formatting_info, diff_summary, sections_to_update)
            
            print("Step 1 Integration completed successfully!")
            return True
            
        except Exception as e:
            print(f"Step 1 Integration failed: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
    
    def run_step1(self, repo_url: str, token: str, before_commit: str, after_commit: str, docx_filename: str):
        """Run Step 1: DOCX to Markdown with LLM updates."""
        try:
            # Step 1: Load DOCX file
            docx_path = self.input_dir / docx_filename
            if not docx_path.exists():
                print(f"DOCX file not found: {docx_path}")
                return False
            
            # Step 2: Convert DOCX to Markdown with metadata
            markdown_content, formatting_info = self.convert_docx_to_markdown_with_metadata(docx_path)
            if not markdown_content:
                print("Failed to convert DOCX to Markdown")
                return False
            
            # Step 3: Fetch commit data
            commit_info = self.fetch_commit_data(repo_url, token, before_commit, after_commit)
            if not commit_info:
                print("Failed to fetch commit data")
                return False
            
            # Step 4: Extract and summarize diff
            diff_summary = self.extract_diff_summary(commit_info)
            
            # Step 5: Identify sections to update
            sections_to_update = self.identify_sections_to_update(markdown_content, diff_summary)
            
            # Step 6: Update markdown with LLM (targeted updates)
            updated_markdown = self.update_markdown_with_llm(markdown_content, commit_info, diff_summary, sections_to_update)
            
            # Step 7: Save results
            results = self.save_step1_results(updated_markdown, formatting_info, diff_summary, sections_to_update)
            
            print("=" * 80)
            print("Step 1 Completed Successfully!")
            print(f"Check the output directory: {self.output_dir}")
            
            return True
            
        except Exception as e:
            print(f"Step 1 failed: {str(e)}")
            return False

def main():
    """Main function to run Step 1."""
    print("Step 1: DOCX to Markdown Conversion with LLM Text Updates")
    print("=" * 80)
    
    # Configuration
    REPO_URL = "https://github.com/bluepal-nipun-marwaha/click"
    TOKEN = ""
    BEFORE_COMMIT = "39a5ba2ca67e994745150164bb44251a0d532c50"
    AFTER_COMMIT = "2ebf734ef773dda7e327fde803be922914a741a4"
    DOCX_FILENAME = "Click_Professional_Documentation.docx"
    
    # Initialize processor
    processor = Step1Processor()
    
    # Run Step 1
    success = processor.run_step1(REPO_URL, TOKEN, BEFORE_COMMIT, AFTER_COMMIT, DOCX_FILENAME)
    
    if success:
        print("\nStep 1 completed successfully!")
        print("Check the 'test_step_by_step/output' folder for results")
        print("\nGenerated files:")
        print("- updated_documentation.md: Updated markdown content")
        print("- paragraph_metadata.json: Paragraph formatting metadata")
    else:
        print("\nStep 1 failed!")
        print("Make sure the DOCX file is in the input directory")

if __name__ == "__main__":
    main()
