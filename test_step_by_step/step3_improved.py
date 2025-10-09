#!/usr/bin/env python3
"""
Step 3 Improved: Enhanced Logic with Proper Table Formatting
- Detailed commit summary
- Better context-aware table processing
- Proper table borders and formatting preservation
"""

import os
import sys
import json
import requests
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple

# Add parent directory to path to import our modules
sys.path.append(str(Path(__file__).parent.parent))

from services.llm_service import LLMService
from utils.config import get_settings
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.table import Table
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io
import re
import structlog

logger = structlog.get_logger(__name__)

def safe_print(text: str, prefix: str = ""):
    """Safely print text that might contain Unicode characters."""
    try:
        print(f"{prefix}{text}")
    except UnicodeEncodeError:
        safe_text = text.encode('ascii', 'replace').decode('ascii')
        print(f"{prefix}{safe_text}")

class ImprovedTableProcessor:
    """Improved processor with better logic and proper table formatting."""
    
    def __init__(self, input_dir=None, output_dir=None, llm_service=None):
        if input_dir and output_dir:
            # Use provided directories
            self.input_dir = Path(input_dir)
            self.output_dir = Path(output_dir)
        else:
            # Use default test directories
            self.test_dir = Path(__file__).parent
            self.input_dir = self.test_dir / "input"
            self.output_dir = self.test_dir / "output"
        
        # Create directories
        self.input_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize LLM service (use provided service or create new one)
        if llm_service:
            self.llm_service = llm_service
        else:
            settings = get_settings()
            self.llm_service = LLMService(settings.ai.model_dump())
        
        print("Step 3 Improved: Enhanced Logic with Proper Table Formatting")
        print(f"Input Directory: {self.input_dir}")
        print(f"Output Directory: {self.output_dir}")
    
    def load_extracted_tables(self) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """Load the previously extracted tables."""
        try:
            tables_json_path = self.output_dir / "extracted_tables.json"
            
            if not tables_json_path.exists():
                print("No extracted_tables.json found. Run extract_tables.py first.")
                return [], {}
            
            print(f"Loading extracted tables from: {tables_json_path.name}")
            
            with open(tables_json_path, 'r', encoding='utf-8') as f:
                tables_data = json.load(f)
            
            print(f"Loaded {len(tables_data)} tables")
            return tables_data, {}
            
        except Exception as e:
            print(f"Error loading extracted tables: {str(e)}")
            return [], {}
    
    def fetch_commit_data(self, repo_url: str, token: str, before_commit: str, after_commit: str) -> Dict[str, Any]:
        """Fetch commit data from GitHub API."""
        try:
            # Extract owner and repo from URL
            repo_path = repo_url.replace('https://github.com/', '').replace('.git', '')
            owner, repo = repo_path.split('/')
            
            print(f"Fetching commit data from {owner}/{repo}")
            
            # Fetch the after commit (the one we're analyzing)
            commit_url = f"https://api.github.com/repos/{owner}/{repo}/commits/{after_commit}"
            headers = {
                'Authorization': f'Bearer {token}',
                'Accept': 'application/vnd.github.v3+json'
            }
            
            response = requests.get(commit_url, headers=headers)
            if response.status_code != 200:
                raise Exception(f"Failed to fetch commit: {response.status_code}")
            
            commit_data = response.json()
            
            # Also fetch the diff
            diff_url = f"https://api.github.com/repos/{owner}/{repo}/commits/{after_commit}"
            diff_headers = {
                'Authorization': f'Bearer {token}',
                'Accept': 'application/vnd.github.v3.diff'
            }
            
            diff_response = requests.get(diff_url, headers=diff_headers)
            diff_content = diff_response.text if diff_response.status_code == 200 else ""
            
            # Extract file changes
            files = commit_data.get('files', [])
            modified_files = [f['filename'] for f in files if f['status'] == 'modified']
            added_files = [f['filename'] for f in files if f['status'] == 'added']
            removed_files = [f['filename'] for f in files if f['status'] == 'removed']
            
            commit_info = {
                'hash': commit_data['sha'],
                'message': commit_data['commit']['message'],
                'author': commit_data['commit']['author']['name'],
                'timestamp': commit_data['commit']['author']['date'],
                'url': commit_data['html_url'],
                'modified': modified_files,
                'added': added_files,
                'removed': removed_files,
                'diff': diff_content,
                'files': files
            }
            
            print(f"Commit data fetched successfully")
            print(f"Message: {commit_info['message']}")
            print(f"Author: {commit_info['author']}")
            
            return commit_info
            
        except Exception as e:
            print(f"Error fetching commit data: {str(e)}")
            return {}
    
    def get_detailed_commit_summary(self, commit_info: Dict[str, Any]) -> str:
        """Get a very detailed summary of the commit changes."""
        try:
            print("Getting detailed commit summary from LLM...")
            
            # Create a comprehensive prompt for detailed summary
            summary_prompt = f"""
You are a technical documentation expert. Please provide a VERY DETAILED analysis of these code changes:

COMMIT INFORMATION:
- Message: {commit_info['message']}
- Author: {commit_info['author']}
- Files Changed: {len(commit_info['files'])} files
- Modified Files: {', '.join(commit_info['modified'])}
- Added Files: {', '.join(commit_info['added'])}
- Removed Files: {', '.join(commit_info['removed'])}

FULL DIFF CONTENT:
{commit_info.get('diff', 'No diff available')}

Please provide a comprehensive analysis covering:

1. **OVERVIEW**: What is the main purpose of this commit?

2. **NEW FEATURES**: What new functionality was added? Be specific about:
   - Feature names and descriptions
   - User-facing capabilities
   - Technical implementation details
   - Any new APIs, functions, or classes

3. **MODIFICATIONS**: What existing features were changed? Include:
   - What was modified and why
   - Breaking changes (if any)
   - Improvements or enhancements
   - Bug fixes

4. **TECHNICAL DETAILS**: 
   - Key files and their roles
   - Architecture changes
   - Dependencies or requirements
   - Performance implications

5. **USER IMPACT**: How does this affect end users?
   - New commands or options
   - Changed behavior
   - New capabilities
   - Documentation needs

6. **DOCUMENTATION IMPLICATIONS**: What documentation needs to be updated?
   - New sections needed
   - Existing sections to modify
   - Examples to add
   - API references to update

Please be very detailed and specific. This summary will be used to update technical documentation tables.
"""
            
            # Get detailed summary from LLM
            detailed_summary = self.llm_service._call_llm(summary_prompt, temperature=0.2, max_tokens=2000)
            
            print("Detailed commit summary obtained:")
            safe_print(detailed_summary[:300] + "..." if len(detailed_summary) > 300 else detailed_summary)
            
            return detailed_summary
            
        except Exception as e:
            print(f"Error getting detailed commit summary: {str(e)}")
            # Fallback to basic info
            return f"Commit: {commit_info['message']}\nFiles: {', '.join(commit_info['modified'] + commit_info['added'])}"
    
    def get_table_context_summary(self, table_data: Dict[str, Any]) -> str:
        """Get a detailed context summary of the table structure and content."""
        try:
            summary = f"TABLE ANALYSIS:\n"
            summary += f"Table ID: {table_data['id']}\n"
            summary += f"Dimensions: {table_data['rows']} rows × {table_data['columns']} columns\n\n"
            
            # Analyze column headers and their purpose
            if table_data['columns_data']:
                summary += "COLUMN STRUCTURE:\n"
                for col_idx, col_data in enumerate(table_data['columns_data']):
                    summary += f"  Column {col_idx + 1}: '{col_data['content']}' - "
                    # Analyze column purpose based on content
                    if 'name' in col_data['content'].lower():
                        summary += "Contains names/identifiers\n"
                    elif 'description' in col_data['content'].lower():
                        summary += "Contains descriptions/explanations\n"
                    elif 'status' in col_data['content'].lower():
                        summary += "Contains status/state information\n"
                    elif 'version' in col_data['content'].lower():
                        summary += "Contains version information\n"
                    elif 'url' in col_data['content'].lower() or 'link' in col_data['content'].lower():
                        summary += "Contains URLs/links\n"
                    else:
                        summary += "Contains general data\n"
            
            # Analyze existing content patterns
            summary += "\nEXISTING CONTENT PATTERNS:\n"
            for row_idx, row_data in enumerate(table_data['rows_data'][:5]):  # First 5 rows
                cell_contents = [cell['content'] for cell in row_data['cells']]
                summary += f"  Row {row_idx + 1}: {' | '.join(cell_contents)}\n"
            
            if len(table_data['rows_data']) > 5:
                summary += f"  ... and {len(table_data['rows_data']) - 5} more rows\n"
            
            return summary
            
        except Exception as e:
            print(f"Error creating table context summary: {str(e)}")
            return str(table_data)
    
    def process_single_table_with_context(self, table_data: Dict[str, Any], commit_info: Dict[str, Any], detailed_summary: str, table_index: int) -> Dict[str, Any]:
        """Process a single table with enhanced context awareness."""
        try:
            print(f"\n--- Processing Table {table_index + 1} with Enhanced Context ---")
            
            # Get detailed table context
            table_context = self.get_table_context_summary(table_data)
            
            # Create enhanced prompt with better context awareness
            prompt = f"""
You are a technical documentation expert. I need you to analyze this table and determine how to update it based on the commit changes.

COMMIT INFORMATION:
- Message: {commit_info['message']}
- Author: {commit_info['author']}
- Modified Files: {', '.join(commit_info['modified'])}
- Added Files: {', '.join(commit_info['added'])}

DETAILED COMMIT ANALYSIS:
{detailed_summary}

TABLE CONTEXT:
{table_context}

ANALYSIS INSTRUCTIONS:
1. **CONTEXT ANALYSIS**: Understand what this table represents and its current structure
2. **RELEVANCE CHECK**: Determine if this table is relevant to the commit changes
3. **UPDATE LEVEL**: Decide if updates are:
   - LOW LEVEL: Minor additions, corrections, or small enhancements
   - HIGH LEVEL: Major new features, significant changes, or new sections
4. **CONTENT STRATEGY**: Plan what specific content should be added/modified

RESPONSE FORMAT:
ANALYSIS: [Brief analysis of table relevance and update level]

UPDATES NEEDED: YES/NO

If YES:
UPDATE LEVEL: LOW/HIGH

NEW ROWS TO ADD:
- [Specific content for new row, separated by | for each column]
- [Another specific content for new row, separated by | for each column]

ROWS TO UPDATE:
- Row [NUMBER]: [Specific new content, separated by | for each column]
- Row [NUMBER]: [Specific new content, separated by | for each column]

COLUMN HEADERS TO CHANGE:
- [Current Header Name]: [New Header Name]

RATIONALE: [Brief explanation of why these updates are needed and how they fit the table's purpose]

IMPORTANT RULES:
- Use actual content, NOT placeholders
- Use actual row numbers like "Row 1", "Row 2", NOT "Row X" or "Row Y"
- Ensure content fits the table's existing structure and purpose
- Be specific and meaningful in your updates
- Consider the table's role in the overall documentation

Please provide your response in the exact format above.
"""
            
            # Generate enhanced response
            llm_response = self.llm_service._call_llm(prompt, temperature=0.3, max_tokens=1500)
            
            print(f"LLM response for Table {table_index + 1}:")
            safe_print(llm_response[:300] + "..." if len(llm_response) > 300 else llm_response)
            
            # Parse and apply updates with enhanced context
            updated_table = self.apply_enhanced_updates(table_data, llm_response, table_index)
            
            return updated_table
            
        except Exception as e:
            print(f"Error processing Table {table_index + 1}: {str(e)}")
            return table_data
    
    def apply_enhanced_updates(self, table_data: Dict[str, Any], llm_response: str, table_index: int) -> Dict[str, Any]:
        """Apply updates with enhanced context awareness."""
        try:
            # Create a copy of the original table
            updated_table = json.loads(json.dumps(table_data))  # Deep copy
            
            # Ensure llm_response is safe for processing
            try:
                test_upper = llm_response.upper()
                safe_llm_response = llm_response
            except UnicodeEncodeError:
                safe_llm_response = llm_response.encode('ascii', 'replace').decode('ascii')
                print(f"  Unicode characters replaced with ASCII equivalents")
            
            # Check if updates are needed
            if "UPDATES NEEDED: NO" in safe_llm_response.upper():
                print(f"Table {table_index + 1}: No updates needed")
                return updated_table
            
            if "UPDATES NEEDED: YES" not in safe_llm_response.upper():
                print(f"Table {table_index + 1}: Could not determine if updates needed")
                return updated_table
            
            # Extract analysis and update level
            if "ANALYSIS:" in safe_llm_response:
                analysis_section = safe_llm_response.split("ANALYSIS:")[1].split("UPDATES NEEDED:")[0].strip()
                print(f"Table {table_index + 1} Analysis: {analysis_section}")
            
            if "UPDATE LEVEL:" in safe_llm_response:
                level_section = safe_llm_response.split("UPDATE LEVEL:")[1].split("NEW ROWS TO ADD:")[0].strip()
                print(f"Table {table_index + 1} Update Level: {level_section}")
            
            print(f"Table {table_index + 1}: Applying enhanced updates...")
            
            # Extract new rows to add
            if "NEW ROWS TO ADD:" in safe_llm_response:
                new_rows_section = safe_llm_response.split("NEW ROWS TO ADD:")[1].split("ROWS TO UPDATE:")[0]
                new_rows = [line.strip("- ").strip() for line in new_rows_section.split("\n") if line.strip().startswith("-")]
                
                for new_row_content in new_rows:
                    if new_row_content:
                        # Clean up generic placeholders, row numbers, and handle Unicode
                        try:
                            clean_content = new_row_content.replace("Row content here:", "").replace("Another row content here:", "").strip()
                            # Remove row number prefixes like "Row 6:", "Row 7:", etc.
                            import re
                            clean_content = re.sub(r'^Row \d+:\s*', '', clean_content)
                        except UnicodeEncodeError:
                            clean_content = new_row_content.encode('ascii', 'replace').decode('ascii').replace("Row content here:", "").replace("Another row content here:", "").strip()
                            # Remove row number prefixes like "Row 6:", "Row 7:", etc.
                            import re
                            clean_content = re.sub(r'^Row \d+:\s*', '', clean_content)
                        
                        if not clean_content:
                            continue
                            
                        # Create new row data
                        new_row = {
                            'index': len(updated_table['rows_data']),
                            'is_header': False,
                            'cells': []
                        }
                        
                        # Split content by | and create cells
                        cell_contents = [content.strip() for content in clean_content.split("|")]
                        for cell_idx, cell_content in enumerate(cell_contents):
                            if cell_idx < updated_table['columns']:
                                # Handle Unicode characters safely
                                try:
                                    safe_content = cell_content
                                except UnicodeEncodeError:
                                    safe_content = cell_content.encode('ascii', 'replace').decode('ascii')
                                
                                new_cell = {
                                    'index': cell_idx,
                                    'content': safe_content,
                                    'is_header': False,
                                    'formatting': {
                                        'font': {'name': 'Calibri', 'size': 11, 'bold': False, 'italic': False, 'color': '000000'},
                                        'alignment': {'horizontal': 'left', 'vertical': 'top'},
                                        'background': {},
                                        'borders': {}
                                    }
                                }
                                new_row['cells'].append(new_cell)
                        
                        # Add remaining empty cells if needed
                        while len(new_row['cells']) < updated_table['columns']:
                            empty_cell = {
                                'index': len(new_row['cells']),
                                'content': '',
                                'is_header': False,
                                'formatting': {
                                    'font': {'name': 'Calibri', 'size': 11, 'bold': False, 'italic': False, 'color': '000000'},
                                    'alignment': {'horizontal': 'left', 'vertical': 'top'},
                                    'background': {},
                                    'borders': {}
                                }
                            }
                            new_row['cells'].append(empty_cell)
                        
                        updated_table['rows_data'].append(new_row)
                        updated_table['rows'] += 1
                        
                        print(f"  Added new row: {clean_content}")
            
            # Extract rows to update
            if "ROWS TO UPDATE:" in safe_llm_response:
                update_rows_section = safe_llm_response.split("ROWS TO UPDATE:")[1].split("COLUMN HEADERS TO CHANGE:")[0]
                update_lines = [line.strip("- ").strip() for line in update_rows_section.split("\n") if line.strip().startswith("-")]
                
                for update_line in update_lines:
                    if ":" in update_line:
                        row_info, new_content = update_line.split(":", 1)
                        row_info = row_info.strip()
                        new_content = new_content.strip()
                        
                        # Try to extract row number - handle "Row X" placeholders
                        if "Row" in row_info:
                            try:
                                # Handle "Row X" and "Row Y" placeholders by skipping them
                                if row_info in ["Row X", "Row Y"]:
                                    print(f"  Skipping placeholder: {row_info}")
                                    continue
                                    
                                row_num = int(row_info.replace("Row", "").strip()) - 1  # Convert to 0-based index
                                if 0 <= row_num < len(updated_table['rows_data']):
                                    # Update the row content - clean up any row number prefixes
                                    cell_contents = [content.strip() for content in new_content.split("|")]
                                    for cell_idx, cell_content in enumerate(cell_contents):
                                        if cell_idx < len(updated_table['rows_data'][row_num]['cells']):
                                            # Clean up row number prefixes and handle Unicode
                                            try:
                                                safe_content = cell_content
                                                # Remove row number prefixes like "Row 6:", "Row 7:", etc.
                                                import re
                                                safe_content = re.sub(r'^Row \d+:\s*', '', safe_content)
                                            except UnicodeEncodeError:
                                                safe_content = cell_content.encode('ascii', 'replace').decode('ascii')
                                                # Remove row number prefixes like "Row 6:", "Row 7:", etc.
                                                import re
                                                safe_content = re.sub(r'^Row \d+:\s*', '', safe_content)
                                            updated_table['rows_data'][row_num]['cells'][cell_idx]['content'] = safe_content
                                    
                                    print(f"  Updated row {row_num + 1}: {new_content}")
                            except ValueError:
                                print(f"  Could not parse row number from: {row_info}")
            
            # Extract column header changes
            if "COLUMN HEADERS TO CHANGE:" in safe_llm_response:
                header_section = safe_llm_response.split("COLUMN HEADERS TO CHANGE:")[1].split("RATIONALE:")[0]
                header_lines = [line.strip("- ").strip() for line in header_section.split("\n") if line.strip().startswith("-")]
                
                for header_line in header_lines:
                    if ":" in header_line:
                        old_header, new_header = header_line.split(":", 1)
                        old_header = old_header.strip()
                        new_header = new_header.strip()
                        
                        # Find and update the column header
                        for col_idx, col_data in enumerate(updated_table['columns_data']):
                            if col_data['content'] == old_header:
                                updated_table['columns_data'][col_idx]['content'] = new_header
                                # Also update the header row
                                if updated_table['rows_data']:
                                    updated_table['rows_data'][0]['cells'][col_idx]['content'] = new_header
                                print(f"  Updated column header: {old_header} -> {new_header}")
                                break
            
            # Extract rationale
            if "RATIONALE:" in safe_llm_response:
                rationale_section = safe_llm_response.split("RATIONALE:")[1].strip()
                print(f"Table {table_index + 1} Rationale: {rationale_section}")
            
            return updated_table
            
        except Exception as e:
            print(f"Error applying enhanced updates to table: {str(e)}")
            return table_data
    
    def create_docx_with_proper_formatting(self, updated_tables: List[Dict[str, Any]]) -> bytes:
        """Create a DOCX file with proper table borders and formatting."""
        try:
            print("Creating DOCX with proper table formatting...")
            
            # Create a new document
            doc = Document()
            
            # Add a title
            title = doc.add_heading('Updated Documentation Tables (Improved)', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Process each table
            for table_idx, table_data in enumerate(updated_tables):
                # Add table title
                table_title = doc.add_heading(f'Table {table_idx + 1}', level=2)
                
                # Create the table
                rows = table_data['rows']
                cols = table_data['columns']
                table = doc.add_table(rows=rows, cols=cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Apply proper table borders
                self.apply_table_borders(table)
                
                # Populate the table with data
                for row_idx, row_data in enumerate(table_data['rows_data']):
                    table_row = table.rows[row_idx]
                    
                    for cell_idx, cell_data in enumerate(row_data['cells']):
                        if cell_idx < len(table_row.cells):
                            cell = table_row.cells[cell_idx]
                            
                            # Set cell content
                            cell.text = cell_data['content']
                            
                            # Apply cell formatting
                            self.apply_cell_formatting(cell, cell_data)
            
            # Save to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            result_bytes = doc_bytes.getvalue()
            print(f"DOCX with proper formatting created: {len(result_bytes)} bytes")
            return result_bytes
            
        except Exception as e:
            print(f"Error creating DOCX with proper formatting: {str(e)}")
            return b""
    
    def apply_table_borders(self, table: Table):
        """Apply proper borders to the table."""
        try:
            # Define border style
            border_xml = '''
            <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            </w:tcBorders>
            '''
            
            # Apply borders to all cells
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcPr.append(parse_xml(border_xml))
                    
        except Exception as e:
            print(f"Error applying table borders: {str(e)}")
    
    def apply_cell_formatting(self, cell, cell_data: Dict[str, Any]):
        """Apply formatting to a cell."""
        try:
            if 'formatting' in cell_data and 'font' in cell_data['formatting']:
                font_info = cell_data['formatting']['font']
                
                if cell.paragraphs:
                    para = cell.paragraphs[0]
                    if para.runs:
                        run = para.runs[0]
                        
                        if 'name' in font_info:
                            run.font.name = font_info['name']
                        if 'size' in font_info:
                            run.font.size = Pt(font_info['size'])
                        if 'bold' in font_info:
                            run.bold = font_info['bold']
                        if 'italic' in font_info:
                            run.italic = font_info['italic']
                            
        except Exception as e:
            print(f"Error applying cell formatting: {str(e)}")
    
    def process_all_tables_enhanced(self, tables_data: List[Dict[str, Any]], commit_info: Dict[str, Any], detailed_summary: str) -> List[Dict[str, Any]]:
        """Process all tables with enhanced logic."""
        try:
            print(f"Processing {len(tables_data)} tables with enhanced logic...")
            
            updated_tables = []
            
            for table_index, table_data in enumerate(tables_data):
                print(f"\n--- Processing Table {table_index + 1}/{len(tables_data)} ---")
                
                # Process single table with enhanced context
                updated_table = self.process_single_table_with_context(table_data, commit_info, detailed_summary, table_index)
                updated_tables.append(updated_table)
                
                # Small delay between requests
                import time
                time.sleep(1)
            
            print(f"\nAll {len(tables_data)} tables processed successfully!")
            return updated_tables
            
        except Exception as e:
            print(f"Error processing tables with enhanced logic: {str(e)}")
            return tables_data
    
    def save_improved_results(self, updated_tables: List[Dict[str, Any]], 
                             table_metadata: Dict[str, Any], 
                             updated_tables_docx: bytes):
        """Save improved results."""
        try:
            # Save updated tables JSON
            tables_json_path = self.output_dir / "updated_tables_improved.json"
            with open(tables_json_path, 'w', encoding='utf-8') as f:
                json.dump(updated_tables, f, indent=2, ensure_ascii=False)
            
            # Save DOCX with updated tables
            docx_path = self.output_dir / "updated_tables_improved.docx"
            with open(docx_path, 'wb') as f:
                f.write(updated_tables_docx)
            
            print(f"Improved results saved successfully")
            print(f"Updated Tables JSON: {tables_json_path}")
            print(f"Updated Tables DOCX: {docx_path}")
            
            return {
                'tables_json_path': tables_json_path,
                'docx_path': docx_path
            }
            
        except Exception as e:
            print(f"Error saving improved results: {str(e)}")
            return {}
    
    def run_improved_processing(self, repo_url: str, token: str, 
                               before_commit: str, after_commit: str):
        """Run improved processing with enhanced logic."""
        try:
            print("Starting Step 3 Improved: Enhanced Logic with Proper Table Formatting")
            print("=" * 80)
            
            # Step 1: Load extracted tables
            tables_data, table_metadata = self.load_extracted_tables()
            if not tables_data:
                print("No tables found. Run step3_extract_only.py first.")
                return False
            
            # Step 2: Fetch commit data
            commit_info = self.fetch_commit_data(repo_url, token, before_commit, after_commit)
            if not commit_info:
                print("Failed to fetch commit data")
                return False
            
            # Step 3: Get detailed commit summary
            detailed_summary = self.get_detailed_commit_summary(commit_info)
            
            # Step 4: Process all tables with enhanced logic
            updated_tables = self.process_all_tables_enhanced(tables_data, commit_info, detailed_summary)
            
            # Step 5: Create DOCX with proper formatting
            updated_tables_docx = self.create_docx_with_proper_formatting(updated_tables)
            
            if not updated_tables_docx:
                print("Failed to create DOCX with proper formatting")
                return False
            
            # Step 6: Save results
            results = self.save_improved_results(updated_tables, table_metadata, updated_tables_docx)
            
            print("=" * 80)
            print("Step 3 Improved Completed Successfully!")
            print(f"Check the output directory: {self.output_dir}")
            
            return True
            
        except Exception as e:
            print(f"Step 3 Improved failed: {str(e)}")
            return False

def main():
    """Main function to run Step 3 Improved."""
    print("Step 3 Improved: Enhanced Logic with Proper Table Formatting")
    print("=" * 80)
    
    # Configuration
    REPO_URL = "https://github.com/bluepal-nipun-marwaha/click"
    TOKEN = ""
    BEFORE_COMMIT = "39a5ba2ca67e994745150164bb44251a0d532c50"
    AFTER_COMMIT = "2ebf734ef773dda7e327fde803be922914a741a4"
    
    # Initialize processor
    processor = ImprovedTableProcessor()
    
    # Run improved processing
    success = processor.run_improved_processing(
        REPO_URL, TOKEN, BEFORE_COMMIT, AFTER_COMMIT
    )
    
    if success:
        print("\nStep 3 Improved completed successfully!")
        print("Check the 'test_step_by_step/output' folder for results")
        print("\nGenerated files:")
        print("- updated_tables_improved.json: Updated table data (enhanced logic)")
        print("- updated_tables_improved.docx: DOCX with proper table formatting")
    else:
        print("\nStep 3 Improved failed!")
        print("Make sure step3_extract_only.py has been run first")

if __name__ == "__main__":
    main()
