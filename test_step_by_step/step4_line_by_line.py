#!/usr/bin/env python3
"""
Step 4: Line-by-Line Documentation Merge
Compares original and updated documentation line by line, using original if no changes,
updated if changes detected, and smart table replacement.
"""

import os
import sys
import json
from pathlib import Path
from typing import Dict, List, Any, Tuple, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# Add the parent directory to the path to import services
sys.path.append(str(Path(__file__).parent.parent))

from services.llm_service import LLMService
from utils.logger import get_logger

logger = get_logger(__name__)

class LineByLineDocumentMerger:
    """Line-by-line document merger that compares original and updated content."""
    
    def __init__(self, input_dir, output_dir):
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.llm_service = LLMService()
        
    def run_step4(self):
        """Run Step 4: Line-by-line documentation merge."""
        try:
            print("Step 4: Line-by-Line Documentation Merge")
            print("=" * 80)
            print(f"Input Directory: {self.input_dir}")
            print(f"Output Directory: {self.output_dir}")
            print("Starting Step 4: Line-by-Line Documentation Merge")
            print("=" * 80)
            
            # Load all documentation components
            original_doc = self.load_original_documentation()
            updated_text_doc = self.load_updated_text_documentation()
            updated_tables = self.load_updated_tables()
            
            # Perform line-by-line merge
            merged_doc = self.merge_line_by_line(original_doc, updated_text_doc, updated_tables)
            
            # Save the merged document
            self.save_merged_document(merged_doc)
            
            print("=" * 80)
            print("Step 4: Line-by-Line Documentation Merge Completed Successfully!")
            print(f"Check the output directory: {self.output_dir}")
            print("=" * 80)
            print("Step 4 completed successfully!")
            print("Check the 'test_step_by_step/output' folder for results")
            print("\nGenerated files:")
            print("- merged_documentation.docx: Line-by-line merged DOCX")
            
        except Exception as e:
            print(f"Step 4 failed: {str(e)}")
            logger.error(f"Step 4 failed: {str(e)}")
            return False
        
        return True
    
    def load_original_documentation(self) -> Document:
        """Load the original documentation."""
        try:
            print("Loading original documentation...")
            # Try both possible filenames
            original_path = self.input_dir / "original.docx"
            if not original_path.exists():
                # Fallback to the original filename
                original_path = self.input_dir / "Click_Professional_Documentation.docx"
                if not original_path.exists():
                    raise FileNotFoundError(f"Original documentation not found in {self.input_dir}. Tried: original.docx and Click_Professional_Documentation.docx")
            
            doc = Document(original_path)
            print(f"Original documentation loaded: {original_path.name}")
            return doc
            
        except Exception as e:
            print(f"Error loading original documentation: {str(e)}")
            raise
    
    def load_updated_text_documentation(self) -> Document:
        """Load the updated text documentation from Step 2."""
        try:
            print("Loading updated text documentation...")
            updated_path = self.output_dir / "updated_documentation.docx"
            if not updated_path.exists():
                raise FileNotFoundError(f"Updated text documentation not found: {updated_path}")
            
            doc = Document(updated_path)
            print(f"Updated text documentation loaded: {updated_path.name}")
            return doc
            
        except Exception as e:
            print(f"Error loading updated text documentation: {str(e)}")
            raise
    
    def load_updated_tables(self) -> List[Dict[str, Any]]:
        """Load the updated tables from Step 3."""
        try:
            print("Loading updated tables...")
            tables_path = self.output_dir / "updated_tables_improved.json"
            if not tables_path.exists():
                raise FileNotFoundError(f"Updated tables not found: {tables_path}")
            
            with open(tables_path, 'r', encoding='utf-8') as f:
                updated_tables = json.load(f)
            
            print(f"Updated tables loaded: {len(updated_tables)} tables")
            return updated_tables
            
        except Exception as e:
            print(f"Error loading updated tables: {str(e)}")
            raise
    
    def merge_line_by_line(self, original_doc: Document, updated_text_doc: Document, updated_tables: List[Dict[str, Any]]) -> Document:
        """Merge documents line by line."""
        try:
            print("Performing line-by-line merge...")
            
            # Create new document
            merged_doc = Document()
            
            # Extract lines from both documents
            original_lines = self.extract_lines_from_doc(original_doc)
            updated_lines = self.extract_lines_from_doc(updated_text_doc)
            
            print(f"Original document: {len(original_lines)} lines")
            print(f"Updated document: {len(updated_lines)} lines")
            
            # Perform line-by-line comparison and merge
            merged_lines = self.compare_and_merge_lines(original_lines, updated_lines, updated_tables)
            
            # Rebuild document from merged lines
            self.rebuild_document_from_lines(merged_doc, merged_lines)
            
            print("Line-by-line merge completed")
            return merged_doc
            
        except Exception as e:
            print(f"Error in line-by-line merge: {str(e)}")
            raise
    
    def extract_lines_from_doc(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract all lines (paragraphs and tables) from a document."""
        try:
            lines = []
            
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    paragraph = None
                    for para in doc.paragraphs:
                        if para._element == element:
                            paragraph = para
                            break
                    
                    if paragraph and paragraph.text.strip():
                        lines.append({
                            'type': 'paragraph',
                            'text': paragraph.text.strip(),
                            'original_element': paragraph,
                            'element': element
                        })
                
                elif element.tag.endswith('tbl'):  # Table
                    table = None
                    for tbl in doc.tables:
                        if tbl._element == element:
                            table = tbl
                            break
                    
                    if table:
                        lines.append({
                            'type': 'table',
                            'table': table,
                            'element': element,
                            'rows': len(table.rows),
                            'cols': len(table.columns)
                        })
            
            return lines
            
        except Exception as e:
            print(f"Error extracting lines from document: {str(e)}")
            return []
    
    def compare_and_merge_lines(self, original_lines: List[Dict[str, Any]], updated_lines: List[Dict[str, Any]], updated_tables: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Compare lines and merge, choosing the best version for each line."""
        try:
            merged_lines = []
            updated_index = 0
            
            for orig_line in original_lines:
                if orig_line['type'] == 'paragraph':
                    # Find matching updated line
                    matching_updated = self.find_matching_paragraph(orig_line, updated_lines)
                    
                    if matching_updated and self.has_paragraph_changes(orig_line, matching_updated):
                        # Use updated version
                        merged_lines.append({
                            'type': 'paragraph',
                            'text': matching_updated['text'],
                            'original_element': matching_updated['original_element'],
                            'source': 'updated'
                        })
                        print(f"  Using updated paragraph: '{orig_line['text'][:50]}...'")
                    else:
                        # Use original version
                        merged_lines.append({
                            'type': 'paragraph',
                            'text': orig_line['text'],
                            'original_element': orig_line['original_element'],
                            'source': 'original'
                        })
                
                elif orig_line['type'] == 'table':
                    # Find matching updated table
                    print(f"  Processing table: {orig_line['rows']}x{orig_line['cols']}")
                    matching_table = self.find_matching_table(orig_line, updated_tables)
                    
                    if matching_table and self.has_table_changes(orig_line, matching_table):
                        # Use updated table
                        merged_lines.append({
                            'type': 'table',
                            'table_data': matching_table,
                            'original_table': orig_line['table'],  # Store original table for formatting
                            'source': 'updated'
                        })
                        print(f"  Using updated table: {orig_line['rows']}x{orig_line['cols']}")
                    else:
                        # Use original table
                        merged_lines.append({
                            'type': 'table',
                            'table': orig_line['table'],
                            'source': 'original'
                        })
                        print(f"  Using original table: {orig_line['rows']}x{orig_line['cols']}")
            
            print(f"Merged {len(merged_lines)} lines")
            return merged_lines
            
        except Exception as e:
            print(f"Error comparing and merging lines: {str(e)}")
            return original_lines
    
    def find_matching_paragraph(self, orig_para: Dict[str, Any], updated_lines: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
        """Find matching paragraph in updated lines."""
        try:
            orig_text = orig_para['text'].strip().lower()
            
            for updated_line in updated_lines:
                if updated_line['type'] == 'paragraph':
                    updated_text = updated_line['text'].strip().lower()
                    
                    # Check for exact match or high similarity
                    if orig_text == updated_text:
                        return updated_line
                    
                    # Check for partial match (for headings or key phrases)
                    if len(orig_text) > 10 and (orig_text in updated_text or updated_text in orig_text):
                        return updated_line
            
            return None
            
        except Exception as e:
            print(f"Error finding matching paragraph: {str(e)}")
            return None
    
    def has_paragraph_changes(self, orig_para: Dict[str, Any], updated_para: Dict[str, Any]) -> bool:
        """Check if paragraph has meaningful changes."""
        try:
            orig_text = orig_para['text'].strip()
            updated_text = updated_para['text'].strip()
            
            # If texts are identical, no changes
            if orig_text == updated_text:
                return False
            
            # Check for meaningful differences (not just whitespace)
            orig_words = set(orig_text.lower().split())
            updated_words = set(updated_text.lower().split())
            
            # If more than 20% of words are different, consider it changed
            if len(orig_words) > 0:
                word_diff = len(orig_words.symmetric_difference(updated_words))
                change_ratio = word_diff / len(orig_words)
                return change_ratio > 0.2
            
            return True
            
        except Exception as e:
            print(f"Error checking paragraph changes: {str(e)}")
            return False
    
    def find_matching_table(self, orig_table: Dict[str, Any], updated_tables: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
        """Find matching table in updated tables."""
        try:
            orig_rows = orig_table['rows']
            orig_cols = orig_table['cols']
            
            print(f"    Looking for table match: {orig_cols} cols, {orig_rows} rows")
            
            # Extract original headers for better matching
            orig_table_obj = orig_table['table']
            orig_headers = []
            if orig_table_obj and len(orig_table_obj.rows) > 0:
                orig_headers = [cell.text.strip() for cell in orig_table_obj.rows[0].cells]
                print(f"    Original headers: {orig_headers}")
            
            best_match = None
            best_score = 0
            
            for i, updated_table in enumerate(updated_tables):
                score = self.calculate_table_match_score(orig_table, updated_table)
                
                # Get table dimensions from the correct structure
                updated_rows = updated_table.get('rows', 0)
                updated_cols = updated_table.get('columns', 0)
                print(f"      Checking table {i}: {updated_cols} cols, {updated_rows} rows")
                
                # Get headers from columns_data structure
                updated_headers = []
                if updated_table.get('columns_data'):
                    updated_headers = [col.get('content', '') for col in updated_table['columns_data'] if col.get('is_header', False)]
                    print(f"      Updated headers: {updated_headers}")
                
                print(f"      Match score: {score}")
                
                if score > best_score and score > 0.3:  # Lower threshold for better matching
                    best_match = updated_table
                    best_score = score
            
            if best_match:
                print(f"    Found best match with score: {best_score}")
            else:
                print(f"    No suitable match found")
            
            return best_match
            
        except Exception as e:
            print(f"Error finding matching table: {str(e)}")
            return None
    
    def calculate_table_match_score(self, orig_table: Dict[str, Any], updated_table: Dict[str, Any]) -> float:
        """Calculate match score between original and updated table."""
        try:
            score = 0.0
            
            # Size similarity (40% weight)
            orig_rows = orig_table['rows']
            orig_cols = orig_table['cols']
            updated_rows = updated_table.get('rows', 0)
            updated_cols = updated_table.get('columns', 0)
            
            if orig_rows > 0 and orig_cols > 0:
                row_similarity = 1.0 - abs(orig_rows - updated_rows) / max(orig_rows, updated_rows)
                col_similarity = 1.0 - abs(orig_cols - updated_cols) / max(orig_cols, updated_cols)
                score += 0.4 * (row_similarity + col_similarity) / 2
            
            # Header similarity (60% weight)
            if updated_table.get('columns_data'):
                updated_headers = [col.get('content', '') for col in updated_table['columns_data'] if col.get('is_header', False)]
                # Extract original headers from table
                orig_table_obj = orig_table['table']
                if orig_table_obj and len(orig_table_obj.rows) > 0:
                    orig_headers = [cell.text.strip() for cell in orig_table_obj.rows[0].cells]
                    
                    if orig_headers and updated_headers:
                        header_matches = sum(1 for h1, h2 in zip(orig_headers, updated_headers) 
                                           if h1.lower() == h2.lower())
                        header_similarity = header_matches / max(len(orig_headers), len(updated_headers))
                        score += 0.6 * header_similarity
            
            return score
            
        except Exception as e:
            print(f"Error calculating table match score: {str(e)}")
            return 0.0
    
    def has_table_changes(self, orig_table: Dict[str, Any], updated_table: Dict[str, Any]) -> bool:
        """Check if table has meaningful changes."""
        try:
            # For now, assume updated tables always have changes
            # This could be enhanced to do content comparison
            return True
            
        except Exception as e:
            print(f"Error checking table changes: {str(e)}")
            return False
    
    def rebuild_document_from_lines(self, doc: Document, merged_lines: List[Dict[str, Any]]):
        """Rebuild document from merged lines."""
        try:
            print("Rebuilding document from merged lines...")
            
            for line in merged_lines:
                if line['type'] == 'paragraph':
                    # Add paragraph
                    para = doc.add_paragraph()
                    para.text = line['text']
                    
                    # Copy formatting from original element
                    if 'original_element' in line:
                        self.copy_paragraph_formatting(line['original_element'], para)
                
                elif line['type'] == 'table':
                    if line['source'] == 'updated':
                        # Add updated table with original formatting
                        original_table = line.get('original_table')
                        self.add_updated_table_to_doc(doc, line['table_data'], original_table)
                    else:
                        # Add original table
                        self.add_original_table_to_doc(doc, line['table'])
            
            print("Document rebuilt successfully")
            
        except Exception as e:
            print(f"Error rebuilding document: {str(e)}")
            raise
    
    def copy_paragraph_formatting(self, original_para, new_para):
        """Copy formatting from original paragraph to new paragraph."""
        try:
            # Copy style
            new_para.style = original_para.style
            
            # Copy run formatting
            for run in original_para.runs:
                new_run = new_para.add_run()
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                if run.font.name:
                    new_run.font.name = run.font.name
                if run.font.size:
                    new_run.font.size = run.font.size
                if run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb
            
        except Exception as e:
            print(f"Error copying paragraph formatting: {str(e)}")
    
    def add_updated_table_to_doc(self, doc: Document, table_data: Dict[str, Any], original_table=None):
        """Add updated table to document with original document's formatting."""
        try:
            if not table_data.get('rows_data'):
                print(f"    No rows_data found in table")
                return
            
            # Get table dimensions
            rows = table_data.get('rows', 0)
            cols = table_data.get('columns', 0)
            
            if rows == 0 or cols == 0:
                print(f"    Invalid table dimensions: {rows} rows x {cols} cols")
                return
            
            print(f"    Creating table: {rows} rows x {cols} cols")
            
            # Create table
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            # Fill table with data from rows_data structure
            rows_data = table_data['rows_data']
            
            # Fill table cells from rows_data with ORIGINAL document formatting
            for row_idx, row_data in enumerate(rows_data):
                if row_idx < len(table.rows):
                    cells = row_data.get('cells', [])
                    for cell_data in cells:
                        col_idx = cell_data.get('index', 0)
                        content = cell_data.get('content', '')
                        
                        if col_idx < len(table.columns):
                            cell = table.cell(row_idx, col_idx)
                            # Handle Unicode characters safely
                            safe_content = str(content).encode('ascii', 'replace').decode('ascii')
                            cell.text = safe_content
                            
                            # Apply ORIGINAL document formatting instead of Step 3 formatting
                            if original_table and row_idx < len(original_table.rows) and col_idx < len(original_table.columns):
                                original_cell = original_table.cell(row_idx, col_idx)
                                self.copy_original_cell_formatting(original_cell, cell)
                                print(f"      Applied original formatting to cell [{row_idx}][{col_idx}]")
                            else:
                                # Fallback: apply basic formatting
                                self.apply_basic_cell_formatting(cell)
                                print(f"      Applied basic formatting to cell [{row_idx}][{col_idx}]")
                            
                            print(f"      Cell [{row_idx}][{col_idx}]: '{safe_content}'")
            
            # Apply borders
            self.apply_table_borders(table)
            
        except Exception as e:
            print(f"Error adding updated table: {str(e)}")
    
    def apply_basic_cell_formatting(self, cell):
        """Apply basic formatting to a cell when original formatting is not available."""
        try:
            # Get the paragraph in the cell
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            
            # Get or create a run
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            
            # Apply basic formatting
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.bold = False
            run.font.italic = False
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
            
            # Set paragraph alignment
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Set cell vertical alignment
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
        except Exception as e:
            print(f"Error applying basic cell formatting: {str(e)}")
    
    def ensure_text_visibility(self, cell):
        """Ensure text in a cell is visible regardless of formatting."""
        try:
            # Get the paragraph in the cell
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            
            # Get or create a run
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            
            # Always ensure text is visible by setting a visible color
            # Check if we can determine background color
            try:
                # Try to get background color from cell
                tcPr = cell._tc.get_or_add_tcPr()
                shading = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                if shading is not None:
                    fill = shading.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                    if fill:
                        # Parse background color
                        bg_color_hex = fill
                        if bg_color_hex.startswith('#'):
                            bg_color_hex = bg_color_hex[1:]
                        
                        try:
                            r = int(bg_color_hex[0:2], 16)
                            g = int(bg_color_hex[2:4], 16)
                            b = int(bg_color_hex[4:6], 16)
                            
                            # If background is dark, use white text
                            if r + g + b < 200:
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                print(f"      Set white text for dark background (RGB: {r},{g},{b})")
                            else:
                                run.font.color.rgb = RGBColor(0, 0, 0)
                                print(f"      Set black text for light background (RGB: {r},{g},{b})")
                        except (ValueError, IndexError):
                            # Default to black text
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            print(f"      Set black text (invalid background color)")
                    else:
                        # No background color, use black text
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        print(f"      Set black text (no background color)")
                else:
                    # No background color, use black text
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    print(f"      Set black text (no shading element)")
            except Exception as e:
                # Fallback to black text
                run.font.color.rgb = RGBColor(0, 0, 0)
                print(f"      Set black text (fallback: {str(e)})")
            
        except Exception as e:
            print(f"Error ensuring text visibility: {str(e)}")
    
    def apply_cell_formatting(self, cell, formatting: Dict[str, Any]):
        """Apply formatting to a table cell."""
        try:
            if not formatting:
                return
            
            # Apply font formatting
            font_info = formatting.get('font', {})
            if font_info:
                # Get the paragraph in the cell
                paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                
                # Apply font properties to the run
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                
                # Font name
                if font_info.get('name'):
                    run.font.name = font_info['name']
                
                # Font size
                if font_info.get('size'):
                    run.font.size = Pt(font_info['size'])
                
                # Bold
                if font_info.get('bold') is not None:
                    run.font.bold = font_info['bold']
                
                # Italic
                if font_info.get('italic') is not None:
                    run.font.italic = font_info['italic']
                
                # Font color with visibility check
                if font_info.get('color'):
                    color_hex = font_info['color']
                    if color_hex.startswith('#'):
                        color_hex = color_hex[1:]
                    elif len(color_hex) == 6:
                        pass  # Already hex without #
                    else:
                        # Default to white for visibility
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        return
                    
                    try:
                        # Convert hex to RGB
                        r = int(color_hex[0:2], 16)
                        g = int(color_hex[2:4], 16)
                        b = int(color_hex[4:6], 16)
                        
                        # Ensure text is visible - if color is too dark, use white
                        if r + g + b < 100:  # Dark color
                            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                        else:
                            run.font.color.rgb = RGBColor(r, g, b)
                    except (ValueError, IndexError):
                        # Default to white for visibility
                        run.font.color.rgb = RGBColor(255, 255, 255)
                else:
                    # Default to white for visibility
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            # Apply background color
            background_info = formatting.get('background', {})
            if background_info.get('color'):
                bg_color_hex = background_info['color']
                if bg_color_hex.startswith('#'):
                    bg_color_hex = bg_color_hex[1:]
                elif len(bg_color_hex) == 6:
                    pass  # Already hex without #
                else:
                    return  # Skip invalid color format
                
                try:
                    # Convert hex to RGB
                    r = int(bg_color_hex[0:2], 16)
                    g = int(bg_color_hex[2:4], 16)
                    b = int(bg_color_hex[4:6], 16)
                    
                    # Apply background color to cell
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:val'), 'clear')
                    shading.set(qn('w:color'), 'auto')
                    shading.set(qn('w:fill'), bg_color_hex)
                    
                    cell._tc.get_or_add_tcPr().append(shading)
                    
                    # Ensure text color contrasts with background
                    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    
                    # If background is dark, ensure text is light
                    if r + g + b < 200:  # Dark background
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                    else:  # Light background
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
                        
                except (ValueError, IndexError):
                    pass  # Skip invalid color values
            
            # Apply alignment
            alignment_info = formatting.get('alignment', {})
            if alignment_info:
                paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                
                # Horizontal alignment
                horizontal = alignment_info.get('horizontal', 'left')
                if horizontal == 'center':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif horizontal == 'right':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif horizontal == 'justify':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:  # left
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Vertical alignment
                vertical = alignment_info.get('vertical', 'top')
                if vertical == 'center':
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                elif vertical == 'bottom':
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                else:  # top
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
        except Exception as e:
            print(f"Error applying cell formatting: {str(e)}")
    
    def add_original_table_to_doc(self, doc: Document, original_table):
        """Add original table to document."""
        try:
            # Create new table with same dimensions
            rows = len(original_table.rows)
            cols = len(original_table.columns)
            
            if rows == 0 or cols == 0:
                return
            
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            # Copy content and preserve original formatting including background colors
            for row_idx, row in enumerate(original_table.rows):
                for col_idx, cell in enumerate(row.cells):
                    if row_idx < len(table.rows) and col_idx < len(table.columns):
                        new_cell = table.cell(row_idx, col_idx)
                        new_cell.text = cell.text
                        
                        # Copy original cell formatting including background color
                        self.copy_original_cell_formatting(cell, new_cell)
            
            # Apply borders
            self.apply_table_borders(table)
            
        except Exception as e:
            print(f"Error adding original table: {str(e)}")
    
    def copy_original_cell_formatting(self, original_cell, new_cell):
        """Copy original cell formatting including background color to preserve visual consistency."""
        try:
            # Copy background color from original cell
            try:
                original_tcPr = original_cell._tc.get_or_add_tcPr()
                original_shading = original_tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                
                if original_shading is not None:
                    # Get the background color from original
                    fill = original_shading.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                    if fill:
                        # Apply the same background color to new cell
                        new_tcPr = new_cell._tc.get_or_add_tcPr()
                        new_shading = OxmlElement('w:shd')
                        new_shading.set(qn('w:val'), 'clear')
                        new_shading.set(qn('w:color'), 'auto')
                        new_shading.set(qn('w:fill'), fill)
                        new_tcPr.append(new_shading)
                        
                        print(f"      Copied background color: {fill}")
            except Exception as e:
                print(f"      Could not copy background color: {str(e)}")
            
            # Copy text formatting from original cell
            try:
                if original_cell.paragraphs and new_cell.paragraphs:
                    original_para = original_cell.paragraphs[0]
                    new_para = new_cell.paragraphs[0]
                    
                    # Copy paragraph formatting
                    if original_para.alignment:
                        new_para.alignment = original_para.alignment
                    
                    # Copy run formatting
                    if original_para.runs and new_para.runs:
                        original_run = original_para.runs[0]
                        new_run = new_para.runs[0]
                        
                        # Copy font properties
                        if original_run.font.name:
                            new_run.font.name = original_run.font.name
                        if original_run.font.size:
                            new_run.font.size = original_run.font.size
                        if original_run.font.bold is not None:
                            new_run.font.bold = original_run.font.bold
                        if original_run.font.italic is not None:
                            new_run.font.italic = original_run.font.italic
                        
                        # Copy text color (preserve original text color)
                        if original_run.font.color.rgb:
                            new_run.font.color.rgb = original_run.font.color.rgb
                            print(f"      Copied text color: {original_run.font.color.rgb}")
                        
            except Exception as e:
                print(f"      Could not copy text formatting: {str(e)}")
                
        except Exception as e:
            print(f"Error copying original cell formatting: {str(e)}")
    
    def apply_table_borders(self, table):
        """Apply borders to table."""
        try:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
            
        except Exception as e:
            print(f"Error applying table borders: {str(e)}")
    
    def save_merged_document(self, merged_doc: Document):
        """Save the merged document."""
        try:
            print("Saving merged document...")
            
            output_path = self.output_dir / "merged_documentation_fixed.docx"
            merged_doc.save(output_path)
            
            print(f"Merged documentation saved successfully")
            print(f"Merged DOCX: {output_path}")
            
        except Exception as e:
            print(f"Error saving merged document: {str(e)}")
            raise

def main():
    """Main function to run Step 4."""
    try:
        # Set up paths
        script_dir = Path(__file__).parent
        input_dir = script_dir / "input"
        output_dir = script_dir / "output"
        
        # Ensure output directory exists
        output_dir.mkdir(exist_ok=True)
        
        # Create merger and run Step 4
        merger = LineByLineDocumentMerger(input_dir, output_dir)
        success = merger.run_step4()
        
        if success:
            print("Step 4 completed successfully!")
            print("Check the 'test_step_by_step/output' folder for results")
        else:
            print("Step 4 failed!")
            sys.exit(1)
            
    except Exception as e:
        print(f"Step 4 failed: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()
