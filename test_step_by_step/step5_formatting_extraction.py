#!/usr/bin/env python3
"""
Step 5: Formatting Extraction and Application
Extracts exact formatting from original document and applies it to the updated document.
"""

import os
import sys
import json
from pathlib import Path
from typing import Dict, List, Any, Tuple, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

class FormattingExtractor:
    """Extracts and applies formatting from original to updated document."""
    
    def __init__(self, input_dir=None, output_dir=None):
        """Initialize the formatting extractor."""
        if input_dir is None:
            self.input_dir = Path(__file__).parent / "input"
        else:
            self.input_dir = Path(input_dir)
            
        if output_dir is None:
            self.output_dir = Path(__file__).parent / "output"
        else:
            self.output_dir = Path(output_dir)
        
        # Ensure directories exist
        self.input_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)
    
    def run_step5(self):
        """Run Step 5: Formatting extraction and application."""
        try:
            print("Step 5: Formatting Extraction and Application")
            print("=" * 80)
            print(f"Input Directory: {self.input_dir}")
            print(f"Output Directory: {self.output_dir}")
            print("Starting Step 5: Formatting Extraction and Application")
            print("=" * 80)
            
            # Load documents
            original_doc = self.load_original_documentation()
            updated_doc = self.load_updated_documentation()
            
            if not original_doc or not updated_doc:
                print("Failed to load documents")
                return False
            
            # Extract formatting from original document
            original_formatting = self.extract_complete_formatting(original_doc)
            
            # Apply original formatting to updated document
            self.apply_original_formatting(updated_doc, original_formatting)
            
            # Save the formatted document
            self.save_formatted_document(updated_doc)
            
            print("=" * 80)
            print("Step 5: Formatting Extraction and Application Completed Successfully!")
            print(f"Check the output directory: {self.output_dir}")
            print("=" * 80)
            
            return True
            
        except Exception as e:
            print(f"Step 5 failed: {str(e)}")
            return False
    
    def load_original_documentation(self) -> Optional[Document]:
        """Load the original documentation."""
        try:
            print("Loading original documentation...")
            docx_files = list(self.input_dir.glob("*.docx"))
            if not docx_files:
                print("No DOCX files found in input directory")
                return None
            
            original_file = docx_files[0]  # Take the first DOCX file
            doc = Document(original_file)
            print(f"Original documentation loaded: {original_file.name}")
            return doc
            
        except Exception as e:
            print(f"Error loading original documentation: {str(e)}")
            return None
    
    def load_updated_documentation(self) -> Optional[Document]:
        """Load the updated documentation."""
        try:
            print("Loading updated documentation...")
            updated_file = self.output_dir / "merged_documentation_fixed.docx"
            if not updated_file.exists():
                print("Updated documentation file not found")
                return None
            
            doc = Document(updated_file)
            print(f"Updated documentation loaded: {updated_file.name}")
            return doc
            
        except Exception as e:
            print(f"Error loading updated documentation: {str(e)}")
            return None
    
    def extract_complete_formatting(self, doc: Document) -> Dict[str, Any]:
        """Extract complete formatting information from the original document."""
        try:
            print("Extracting complete formatting from original document...")
            
            formatting_info = {
                "document_styles": {},
                "paragraphs": [],
                "tables": [],
                "sections": []
            }
            
            # Extract document styles
            for style in doc.styles:
                try:
                    if hasattr(style, 'font') and style.font:
                        formatting_info["document_styles"][style.name] = {
                            "type": str(style.type),
                            "font_name": style.font.name if style.font.name else None,
                            "font_size": style.font.size.pt if style.font.size else None,
                            "bold": style.font.bold,
                            "italic": style.font.italic,
                            "color": str(style.font.color.rgb) if style.font.color.rgb else None
                        }
                    else:
                        formatting_info["document_styles"][style.name] = {
                            "type": str(style.type),
                            "font_name": None,
                            "font_size": None,
                            "bold": None,
                            "italic": None,
                            "color": None
                        }
                except Exception as e:
                    print(f"Error extracting style {style.name}: {str(e)}")
                    continue
            
            # Extract paragraph formatting with content matching
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    para_info = {
                        "index": i,
                        "text": paragraph.text.strip(),
                        "style_name": paragraph.style.name,
                        "alignment": str(paragraph.alignment) if paragraph.alignment else None,
                        "space_before": paragraph.paragraph_format.space_before.pt if paragraph.paragraph_format.space_before else None,
                        "space_after": paragraph.paragraph_format.space_after.pt if paragraph.paragraph_format.space_after else None,
                        "runs": []
                    }
                    
                    # Extract run formatting
                    for run in paragraph.runs:
                        if run.text.strip():
                            run_info = {
                                "text": run.text.strip(),
                                "font_name": run.font.name,
                                "font_size": run.font.size.pt if run.font.size else None,
                                "bold": run.font.bold,
                                "italic": run.font.italic,
                                "color": str(run.font.color.rgb) if run.font.color.rgb else None,
                                "underline": run.font.underline
                            }
                            para_info["runs"].append(run_info)
                    
                    formatting_info["paragraphs"].append(para_info)
            
            # Extract table formatting with content matching
            for i, table in enumerate(doc.tables):
                table_info = {
                    "index": i,
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                    "alignment": str(table.alignment) if table.alignment else None,
                    "cells": []
                }
                
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_info = {
                            "row": row_idx,
                            "column": col_idx,
                            "text": cell.text.strip(),
                            "background_color": self.extract_cell_background(cell),
                            "vertical_alignment": str(cell.vertical_alignment) if cell.vertical_alignment else None,
                            "paragraphs": []
                        }
                        
                        # Extract cell paragraph formatting
                        for para in cell.paragraphs:
                            if para.text.strip():
                                para_info = {
                                    "text": para.text.strip(),
                                    "alignment": str(para.alignment) if para.alignment else None,
                                    "runs": []
                                }
                                
                                for run in para.runs:
                                    if run.text.strip():
                                        run_info = {
                                            "text": run.text.strip(),
                                            "font_name": run.font.name,
                                            "font_size": run.font.size.pt if run.font.size else None,
                                            "bold": run.font.bold,
                                            "italic": run.font.italic,
                                            "color": str(run.font.color.rgb) if run.font.color.rgb else None
                                        }
                                        para_info["runs"].append(run_info)
                                
                                cell_info["paragraphs"].append(para_info)
                        
                        table_info["cells"].append(cell_info)
                
                formatting_info["tables"].append(table_info)
            
            print(f"Extracted formatting from {len(formatting_info['paragraphs'])} paragraphs and {len(formatting_info['tables'])} tables")
            return formatting_info
            
        except Exception as e:
            print(f"Error extracting formatting: {str(e)}")
            return {}
    
    def extract_cell_background(self, cell) -> Optional[str]:
        """Extract background color from a table cell."""
        try:
            tcPr = cell._tc.get_or_add_tcPr()
            shading = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
            if shading is not None:
                fill = shading.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                if fill:
                    return fill
            return None
        except Exception:
            return None
    
    def apply_original_formatting(self, doc: Document, original_formatting: Dict[str, Any]):
        """Apply original formatting to the updated document."""
        try:
            print("Applying original formatting to updated document...")
            
            # Apply document styles
            self.apply_document_styles(doc, original_formatting.get("document_styles", {}))
            
            # Apply paragraph formatting
            self.apply_paragraph_formatting(doc, original_formatting.get("paragraphs", []))
            
            # Apply table formatting
            self.apply_table_formatting(doc, original_formatting.get("tables", []))
            
            print("Original formatting applied successfully")
            
        except Exception as e:
            print(f"Error applying original formatting: {str(e)}")
    
    def apply_document_styles(self, doc: Document, styles: Dict[str, Any]):
        """Apply document styles."""
        try:
            print("  Applying document styles...")
            
            for style_name, style_info in styles.items():
                try:
                    if style_name in doc.styles:
                        style = doc.styles[style_name]
                        if hasattr(style, 'font') and style.font:
                            if style_info.get("font_name"):
                                style.font.name = style_info["font_name"]
                            if style_info.get("font_size"):
                                style.font.size = Pt(style_info["font_size"])
                            if style_info.get("bold") is not None:
                                style.font.bold = style_info["bold"]
                            if style_info.get("italic") is not None:
                                style.font.italic = style_info["italic"]
                            if style_info.get("color"):
                                color_hex = style_info["color"]
                                if color_hex.startswith('#'):
                                    color_hex = color_hex[1:]
                                try:
                                    r = int(color_hex[0:2], 16)
                                    g = int(color_hex[2:4], 16)
                                    b = int(color_hex[4:6], 16)
                                    style.font.color.rgb = RGBColor(r, g, b)
                                except (ValueError, IndexError):
                                    pass
                except Exception as e:
                    print(f"    Error applying style {style_name}: {str(e)}")
                    continue
            
        except Exception as e:
            print(f"    Error applying document styles: {str(e)}")
    
    def apply_paragraph_formatting(self, doc: Document, paragraphs: List[Dict[str, Any]]):
        """Apply paragraph formatting by matching content."""
        try:
            print("  Applying paragraph formatting...")
            
            applied_count = 0
            
            for original_para in paragraphs:
                original_text = original_para["text"]
                
                # Find matching paragraph in updated document
                matching_para = self.find_matching_paragraph(doc, original_text)
                if matching_para:
                    self.apply_paragraph_format(matching_para, original_para)
                    applied_count += 1
            
            print(f"    Applied formatting to {applied_count} paragraphs")
            
        except Exception as e:
            print(f"    Error applying paragraph formatting: {str(e)}")
    
    def find_matching_paragraph(self, doc: Document, target_text: str) -> Optional[Any]:
        """Find a paragraph in the document that matches the target text."""
        try:
            # Try exact match first
            for paragraph in doc.paragraphs:
                if paragraph.text.strip() == target_text:
                    return paragraph
            
            # Try partial match
            for paragraph in doc.paragraphs:
                if target_text in paragraph.text.strip() or paragraph.text.strip() in target_text:
                    return paragraph
            
            return None
            
        except Exception:
            return None
    
    def apply_paragraph_format(self, paragraph, original_format: Dict[str, Any]):
        """Apply formatting to a paragraph."""
        try:
            # Apply paragraph-level formatting
            if original_format.get("alignment"):
                alignment = original_format["alignment"]
                if "CENTER" in alignment:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif "RIGHT" in alignment:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif "JUSTIFY" in alignment:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            if original_format.get("space_before"):
                paragraph.paragraph_format.space_before = Pt(original_format["space_before"])
            if original_format.get("space_after"):
                paragraph.paragraph_format.space_after = Pt(original_format["space_after"])
            
            # Apply run formatting
            original_runs = original_format.get("runs", [])
            if original_runs and paragraph.runs:
                for i, run in enumerate(paragraph.runs):
                    if i < len(original_runs):
                        original_run = original_runs[i]
                        self.apply_run_format(run, original_run)
            
        except Exception as e:
            print(f"      Error applying paragraph format: {str(e)}")
    
    def apply_run_format(self, run, original_format: Dict[str, Any]):
        """Apply formatting to a run with proper text visibility."""
        try:
            if original_format.get("font_name"):
                run.font.name = original_format["font_name"]
            if original_format.get("font_size"):
                run.font.size = Pt(original_format["font_size"])
            if original_format.get("bold") is not None:
                run.font.bold = original_format["bold"]
            if original_format.get("italic") is not None:
                run.font.italic = original_format["italic"]
            
            # Apply color with visibility check
            if original_format.get("color"):
                color_hex = original_format["color"]
                if color_hex.startswith('#'):
                    color_hex = color_hex[1:]
                try:
                    r = int(color_hex[0:2], 16)
                    g = int(color_hex[2:4], 16)
                    b = int(color_hex[4:6], 16)
                    
                    # Ensure text is visible - if color is too dark, use white
                    if r + g + b < 100:  # Dark color
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                    else:
                        run.font.color.rgb = RGBColor(r, g, b)
                except (ValueError, IndexError):
                    # Default to white text for visibility
                    run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                # Default to white text for visibility
                run.font.color.rgb = RGBColor(255, 255, 255)
            
        except Exception as e:
            print(f"        Error applying run format: {str(e)}")
            # Ensure text is visible even if there's an error
            try:
                run.font.color.rgb = RGBColor(255, 255, 255)
            except:
                pass
    
    def apply_paragraph_formatting_only(self, paragraph, original_format: Dict[str, Any]):
        """Apply formatting to a paragraph without changing content."""
        try:
            # Apply paragraph-level formatting
            if original_format.get("alignment"):
                alignment = original_format["alignment"]
                if "CENTER" in alignment:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif "RIGHT" in alignment:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif "JUSTIFY" in alignment:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            if original_format.get("space_before"):
                paragraph.paragraph_format.space_before = Pt(original_format["space_before"])
            if original_format.get("space_after"):
                paragraph.paragraph_format.space_after = Pt(original_format["space_after"])
            
            # Apply run formatting to existing runs without changing text
            original_runs = original_format.get("runs", [])
            if original_runs and paragraph.runs:
                for i, run in enumerate(paragraph.runs):
                    if i < len(original_runs):
                        original_run = original_runs[i]
                        self.apply_run_format(run, original_run)
            
        except Exception as e:
            print(f"      Error applying paragraph formatting only: {str(e)}")
    
    def apply_table_formatting(self, doc: Document, tables: List[Dict[str, Any]]):
        """Apply table formatting by matching content."""
        try:
            print("  Applying table formatting...")
            
            applied_count = 0
            
            for original_table in tables:
                # Find matching table in updated document
                matching_table = self.find_matching_table(doc, original_table)
                if matching_table:
                    self.apply_table_format(matching_table, original_table)
                    applied_count += 1
            
            print(f"    Applied formatting to {applied_count} tables")
            
        except Exception as e:
            print(f"    Error applying table formatting: {str(e)}")
    
    def find_matching_table(self, doc: Document, target_table: Dict[str, Any]) -> Optional[Any]:
        """Find a table in the document that matches the target table."""
        try:
            target_rows = target_table["rows"]
            target_cols = target_table["columns"]
            
            for table in doc.tables:
                if len(table.rows) == target_rows and len(table.columns) == target_cols:
                    # Check if header content matches
                    if table.rows and target_table.get("cells"):
                        original_header = target_table["cells"][0] if target_table["cells"] else None
                        if original_header:
                            table_header = table.rows[0].cells[0].text.strip()
                            if original_header["text"] == table_header:
                                return table
            
            return None
            
        except Exception:
            return None
    
    def apply_table_format(self, table, original_format: Dict[str, Any]):
        """Apply formatting to a table."""
        try:
            # Apply table alignment
            if original_format.get("alignment"):
                alignment = original_format["alignment"]
                if "CENTER" in alignment:
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                elif "RIGHT" in alignment:
                    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
                else:
                    table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # Apply cell formatting
            original_cells = original_format.get("cells", [])
            for original_cell in original_cells:
                row_idx = original_cell["row"]
                col_idx = original_cell["column"]
                
                if row_idx < len(table.rows) and col_idx < len(table.columns):
                    cell = table.cell(row_idx, col_idx)
                    self.apply_cell_format(cell, original_cell)
            
        except Exception as e:
            print(f"      Error applying table format: {str(e)}")
    
    def apply_cell_format(self, cell, original_format: Dict[str, Any]):
        """Apply formatting to a table cell without changing content."""
        try:
            # Apply background color
            if original_format.get("background_color"):
                self.set_cell_background(cell, original_format["background_color"])
            
            # Apply vertical alignment
            if original_format.get("vertical_alignment"):
                alignment = original_format["vertical_alignment"]
                if "CENTER" in alignment:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                elif "BOTTOM" in alignment:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                else:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
            # Apply paragraph formatting WITHOUT changing content
            original_paragraphs = original_format.get("paragraphs", [])
            for i, paragraph in enumerate(cell.paragraphs):
                if i < len(original_paragraphs):
                    original_para = original_paragraphs[i]
                    # Only apply formatting, not content
                    self.apply_paragraph_formatting_only(paragraph, original_para)
            
        except Exception as e:
            print(f"        Error applying cell format: {str(e)}")
    
    def set_cell_background(self, cell, color_hex: str):
        """Set background color for a table cell with proper contrast."""
        try:
            if color_hex.startswith('#'):
                color_hex = color_hex[1:]
            
            # Ensure we have a valid hex color
            if len(color_hex) != 6:
                return
            
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), color_hex)
            
            cell._tc.get_or_add_tcPr().append(shading)
            
            # Also ensure text color is visible against this background
            try:
                r = int(color_hex[0:2], 16)
                g = int(color_hex[2:4], 16)
                b = int(color_hex[4:6], 16)
                
                # If background is dark, ensure text is light
                if r + g + b < 200:  # Dark background
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                else:  # Light background
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
            except (ValueError, IndexError):
                pass
            
        except Exception as e:
            print(f"          Error setting cell background: {str(e)}")
    
    def save_formatted_document(self, doc: Document):
        """Save the formatted document."""
        try:
            print("Saving formatted document...")
            
            output_path = self.output_dir / "merged_documentation_formatting_applied.docx"
            doc.save(output_path)
            
            print(f"Formatted documentation saved successfully")
            print(f"Formatted DOCX: {output_path}")
            
        except Exception as e:
            print(f"Error saving formatted document: {str(e)}")
            raise

def main():
    """Main function to run Step 5."""
    extractor = FormattingExtractor()
    success = extractor.run_step5()
    
    if success:
        print("Step 5 completed successfully!")
        print("Check the 'test_step_by_step/output' folder for results")
        print("\nGenerated files:")
        print("- merged_documentation_formatting_applied.docx: Original formatting applied")
    else:
        print("Step 5 failed!")
        sys.exit(1)

if __name__ == "__main__":
    main()
