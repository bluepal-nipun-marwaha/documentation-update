#!/usr/bin/env python3
"""
Step 5: LLM-Powered Formatting Fix
Uses LLM to analyze formatting differences between original and merged documents,
then applies fixes to make the formatting accurate.
"""

import os
import sys
import json
from pathlib import Path
from typing import Dict, List, Any, Tuple, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# Set environment variables for Ollama
os.environ['LLM_PROVIDER'] = 'ollama'
os.environ['LLM_MODEL'] = 'qwen2.5:7b'
os.environ['EMBEDDINGS_PROVIDER'] = 'ollama'
os.environ['EMBEDDINGS_MODEL'] = 'nomic-embed-text'
os.environ['OLLAMA_BASE_URL'] = 'http://localhost:11434'

# Add the parent directory to the path to import services
sys.path.append(str(Path(__file__).parent.parent))

from services.llm_service import LLMService
from utils.logger import get_logger

logger = get_logger(__name__)

class LLMFormattingFixer:
    """LLM-powered formatting fixer for documentation."""
    
    def __init__(self):
        """Initialize the formatting fixer."""
        self.input_dir = Path(__file__).parent / "input"
        self.output_dir = Path(__file__).parent / "output"
        self.llm_service = LLMService()
        
        # Ensure directories exist
        self.input_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)
    
    def run_step5(self):
        """Run Step 5: LLM-powered formatting fix."""
        try:
            print("Step 5: LLM-Powered Formatting Fix")
            print("=" * 80)
            print(f"Input Directory: {self.input_dir}")
            print(f"Output Directory: {self.output_dir}")
            print("Starting Step 5: LLM-Powered Formatting Fix")
            print("=" * 80)
            
            # Load documents
            original_doc = self.load_original_documentation()
            merged_doc = self.load_merged_documentation()
            
            if not original_doc or not merged_doc:
                print("Failed to load documents")
                return False
            
            # Extract formatting information from both documents
            original_formatting = self.extract_document_formatting(original_doc, "original")
            merged_formatting = self.extract_document_formatting(merged_doc, "merged")
            
            # Use LLM to analyze formatting differences
            formatting_analysis = self.analyze_formatting_with_llm(original_formatting, merged_formatting)
            
            # Apply LLM-suggested fixes
            fixed_doc = self.apply_formatting_fixes(merged_doc, formatting_analysis)
            
            # Save the fixed document
            self.save_fixed_document(fixed_doc)
            
            print("=" * 80)
            print("Step 5: LLM-Powered Formatting Fix Completed Successfully!")
            print(f"Check the output directory: {self.output_dir}")
            print("=" * 80)
            
            return True
            
        except Exception as e:
            print(f"Step 5 failed: {str(e)}")
            logger.error(f"Step 5 failed: {str(e)}")
            return False
    
    def load_original_documentation(self) -> Optional[Document]:
        """Load the original documentation."""
        try:
            print("Loading original documentation...")
            docx_files = list(self.input_dir.glob("*.docx"))
            if not docx_files:
                print("No DOCX files found in input directory")
                return None
            
            original_file = docx_files[0]  # Take the first DOCX file
            doc = Document(original_file)
            print(f"Original documentation loaded: {original_file.name}")
            return doc
            
        except Exception as e:
            print(f"Error loading original documentation: {str(e)}")
            return None
    
    def load_merged_documentation(self) -> Optional[Document]:
        """Load the merged documentation."""
        try:
            print("Loading merged documentation...")
            merged_file = self.output_dir / "merged_documentation_fixed.docx"
            if not merged_file.exists():
                print("Merged documentation file not found")
                return None
            
            doc = Document(merged_file)
            print(f"Merged documentation loaded: {merged_file.name}")
            return doc
            
        except Exception as e:
            print(f"Error loading merged documentation: {str(e)}")
            return None
    
    def extract_document_formatting(self, doc: Document, doc_type: str) -> Dict[str, Any]:
        """Extract formatting information from a document."""
        try:
            print(f"Extracting formatting from {doc_type} document...")
            
            formatting_info = {
                "document_styles": {},
                "paragraphs": [],
                "tables": [],
                "headings": []
            }
            
            # Extract document styles
            for style in doc.styles:
                try:
                    # Skip numbering styles that don't have font attributes
                    if hasattr(style, 'font') and style.font:
                        formatting_info["document_styles"][style.name] = {
                            "type": str(style.type),
                            "font_name": style.font.name if style.font.name else None,
                            "font_size": style.font.size.pt if style.font.size else None,
                            "bold": style.font.bold,
                            "italic": style.font.italic,
                            "color": str(style.font.color.rgb) if style.font.color.rgb else None
                        }
                    else:
                        # For styles without font, just store basic info
                        formatting_info["document_styles"][style.name] = {
                            "type": str(style.type),
                            "font_name": None,
                            "font_size": None,
                            "bold": None,
                            "italic": None,
                            "color": None
                        }
                except Exception as e:
                    print(f"Error extracting style {style.name}: {str(e)}")
                    continue
            
            # Extract paragraph formatting
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    para_info = {
                        "index": i,
                        "text": paragraph.text.strip(),
                        "style": paragraph.style.name,
                        "alignment": str(paragraph.alignment) if paragraph.alignment else None,
                        "runs": []
                    }
                    
                    # Extract run formatting
                    for run in paragraph.runs:
                        if run.text.strip():
                            run_info = {
                                "text": run.text.strip(),
                                "font_name": run.font.name,
                                "font_size": run.font.size.pt if run.font.size else None,
                                "bold": run.font.bold,
                                "italic": run.font.italic,
                                "color": str(run.font.color.rgb) if run.font.color.rgb else None
                            }
                            para_info["runs"].append(run_info)
                    
                    formatting_info["paragraphs"].append(para_info)
            
            # Extract table formatting
            for i, table in enumerate(doc.tables):
                table_info = {
                    "index": i,
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                    "cells": []
                }
                
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_info = {
                            "row": row_idx,
                            "column": col_idx,
                            "text": cell.text.strip(),
                            "paragraphs": []
                        }
                        
                        # Extract cell paragraph formatting
                        for para in cell.paragraphs:
                            if para.text.strip():
                                para_info = {
                                    "text": para.text.strip(),
                                    "alignment": str(para.alignment) if para.alignment else None,
                                    "runs": []
                                }
                                
                                for run in para.runs:
                                    if run.text.strip():
                                        run_info = {
                                            "text": run.text.strip(),
                                            "font_name": run.font.name,
                                            "font_size": run.font.size.pt if run.font.size else None,
                                            "bold": run.font.bold,
                                            "italic": run.font.italic,
                                            "color": str(run.font.color.rgb) if run.font.color.rgb else None
                                        }
                                        para_info["runs"].append(run_info)
                                
                                cell_info["paragraphs"].append(para_info)
                        
                        table_info["cells"].append(cell_info)
                
                formatting_info["tables"].append(table_info)
            
            print(f"Extracted formatting from {len(formatting_info['paragraphs'])} paragraphs and {len(formatting_info['tables'])} tables")
            return formatting_info
            
        except Exception as e:
            print(f"Error extracting formatting: {str(e)}")
            return {}
    
    def analyze_formatting_with_llm(self, original_formatting: Dict[str, Any], merged_formatting: Dict[str, Any]) -> Dict[str, Any]:
        """Use LLM to analyze formatting differences and suggest fixes."""
        try:
            print("Analyzing formatting differences with LLM...")
            
            # Create a summary of formatting differences
            original_summary = self.create_formatting_summary(original_formatting)
            merged_summary = self.create_formatting_summary(merged_formatting)
            
            prompt = f"""
You are a document formatting expert. I need you to analyze the formatting differences between an original document and a merged document, then provide specific fixes to make the formatting accurate.

ORIGINAL DOCUMENT FORMATTING:
{original_summary}

MERGED DOCUMENT FORMATTING:
{merged_summary}

Please analyze the differences and provide specific formatting fixes. Focus on:

1. **Font Consistency**: Ensure fonts match (name, size, bold, italic)
2. **Color Consistency**: Ensure colors match (text colors, background colors)
3. **Alignment Consistency**: Ensure text alignment matches
4. **Table Formatting**: Ensure table cell formatting matches
5. **Heading Styles**: Ensure heading styles are consistent

Provide your response in this JSON format:
{{
    "analysis": "Brief analysis of main formatting issues",
    "fixes": [
        {{
            "type": "paragraph|table_cell|heading",
            "element_index": 0,
            "description": "What needs to be fixed",
            "font_name": "Calibri",
            "font_size": 10.5,
            "bold": true,
            "italic": false,
            "color": "FFFFFF",
            "alignment": "left|center|right|justify",
            "background_color": "34495E"
        }}
    ]
}}

Be specific and provide exact values for each fix needed.
"""
            
            # Get LLM analysis
            llm_response = self.llm_service._call_llm(prompt, temperature=0.1, max_tokens=3000)
            
            # Parse LLM response
            try:
                # Extract JSON from response
                json_start = llm_response.find('{')
                json_end = llm_response.rfind('}') + 1
                if json_start != -1 and json_end != -1:
                    json_str = llm_response[json_start:json_end]
                    analysis = json.loads(json_str)
                    print(f"LLM Analysis: {analysis.get('analysis', 'No analysis provided')}")
                    print(f"Found {len(analysis.get('fixes', []))} formatting fixes")
                    return analysis
                else:
                    print("Could not extract JSON from LLM response")
                    return {"analysis": "Failed to parse LLM response", "fixes": []}
                    
            except json.JSONDecodeError as e:
                print(f"Error parsing LLM JSON response: {str(e)}")
                return {"analysis": "Failed to parse LLM response", "fixes": []}
            
        except Exception as e:
            print(f"Error analyzing formatting with LLM: {str(e)}")
            return {"analysis": f"Error: {str(e)}", "fixes": []}
    
    def create_formatting_summary(self, formatting: Dict[str, Any]) -> str:
        """Create a summary of formatting information for LLM analysis."""
        try:
            summary_parts = []
            
            # Document styles summary
            if formatting.get("document_styles"):
                summary_parts.append("DOCUMENT STYLES:")
                for style_name, style_info in list(formatting["document_styles"].items())[:5]:  # Limit to first 5
                    summary_parts.append(f"  {style_name}: {style_info.get('font_name', 'N/A')} {style_info.get('font_size', 'N/A')}pt")
            
            # Paragraph formatting summary
            if formatting.get("paragraphs"):
                summary_parts.append(f"\nPARAGRAPHS ({len(formatting['paragraphs'])}):")
                for i, para in enumerate(formatting["paragraphs"][:10]):  # Limit to first 10
                    summary_parts.append(f"  Para {i}: '{para['text'][:50]}...' - Style: {para.get('style', 'N/A')}")
                    if para.get("runs"):
                        for j, run in enumerate(para["runs"][:3]):  # Limit to first 3 runs
                            summary_parts.append(f"    Run {j}: {run.get('font_name', 'N/A')} {run.get('font_size', 'N/A')}pt Bold:{run.get('bold', False)}")
            
            # Table formatting summary
            if formatting.get("tables"):
                summary_parts.append(f"\nTABLES ({len(formatting['tables'])}):")
                for i, table in enumerate(formatting["tables"][:5]):  # Limit to first 5 tables
                    summary_parts.append(f"  Table {i}: {table['rows']}x{table['columns']} cells")
                    if table.get("cells"):
                        for j, cell in enumerate(table["cells"][:3]):  # Limit to first 3 cells
                            summary_parts.append(f"    Cell {j}: '{cell['text'][:30]}...'")
            
            return "\n".join(summary_parts)
            
        except Exception as e:
            return f"Error creating summary: {str(e)}"
    
    def apply_formatting_fixes(self, doc: Document, analysis: Dict[str, Any]) -> Document:
        """Apply the LLM-suggested formatting fixes to the document."""
        try:
            print("Applying formatting fixes...")
            
            fixes = analysis.get("fixes", [])
            if not fixes:
                print("No fixes to apply")
                return doc
            
            applied_fixes = 0
            
            for fix in fixes:
                try:
                    fix_type = fix.get("type", "")
                    element_index = fix.get("element_index", 0)
                    
                    if fix_type == "paragraph":
                        if element_index < len(doc.paragraphs):
                            self.apply_paragraph_fix(doc.paragraphs[element_index], fix)
                            applied_fixes += 1
                    
                    elif fix_type == "table_cell":
                        # Find table and cell by index
                        table_idx = fix.get("table_index", 0)
                        cell_idx = fix.get("cell_index", 0)
                        if table_idx < len(doc.tables):
                            table = doc.tables[table_idx]
                            if cell_idx < len(table._cells):
                                cell = table._cells[cell_idx]
                                self.apply_cell_fix(cell, fix)
                                applied_fixes += 1
                    
                    elif fix_type == "heading":
                        if element_index < len(doc.paragraphs):
                            self.apply_paragraph_fix(doc.paragraphs[element_index], fix)
                            applied_fixes += 1
                
                except Exception as e:
                    print(f"Error applying fix {fix.get('description', 'Unknown')}: {str(e)}")
                    continue
            
            print(f"Applied {applied_fixes} formatting fixes")
            return doc
            
        except Exception as e:
            print(f"Error applying formatting fixes: {str(e)}")
            return doc
    
    def apply_paragraph_fix(self, paragraph, fix: Dict[str, Any]):
        """Apply formatting fix to a paragraph."""
        try:
            # Apply to the first run or create one
            if paragraph.runs:
                run = paragraph.runs[0]
            else:
                run = paragraph.add_run(paragraph.text)
            
            # Apply font formatting
            if fix.get("font_name"):
                run.font.name = fix["font_name"]
            if fix.get("font_size"):
                run.font.size = Pt(fix["font_size"])
            if fix.get("bold") is not None:
                run.font.bold = fix["bold"]
            if fix.get("italic") is not None:
                run.font.italic = fix["italic"]
            if fix.get("color"):
                color_hex = fix["color"]
                if color_hex.startswith('#'):
                    color_hex = color_hex[1:]
                try:
                    r = int(color_hex[0:2], 16)
                    g = int(color_hex[2:4], 16)
                    b = int(color_hex[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
                except (ValueError, IndexError):
                    pass
            
            # Apply alignment
            if fix.get("alignment"):
                alignment = fix["alignment"]
                if alignment == "center":
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment == "right":
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignment == "justify":
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        except Exception as e:
            print(f"Error applying paragraph fix: {str(e)}")
    
    def apply_cell_fix(self, cell, fix: Dict[str, Any]):
        """Apply formatting fix to a table cell."""
        try:
            # Apply background color
            if fix.get("background_color"):
                bg_color_hex = fix["background_color"]
                if bg_color_hex.startswith('#'):
                    bg_color_hex = bg_color_hex[1:]
                
                try:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:val'), 'clear')
                    shading.set(qn('w:color'), 'auto')
                    shading.set(qn('w:fill'), bg_color_hex)
                    cell._tc.get_or_add_tcPr().append(shading)
                except Exception:
                    pass
            
            # Apply paragraph formatting to cell paragraphs
            for paragraph in cell.paragraphs:
                self.apply_paragraph_fix(paragraph, fix)
            
        except Exception as e:
            print(f"Error applying cell fix: {str(e)}")
    
    def save_fixed_document(self, doc: Document):
        """Save the fixed document."""
        try:
            print("Saving fixed document...")
            
            output_path = self.output_dir / "merged_documentation_llm_fixed.docx"
            doc.save(output_path)
            
            print(f"Fixed documentation saved successfully")
            print(f"Fixed DOCX: {output_path}")
            
        except Exception as e:
            print(f"Error saving fixed document: {str(e)}")
            raise

def main():
    """Main function to run Step 5."""
    fixer = LLMFormattingFixer()
    success = fixer.run_step5()
    
    if success:
        print("Step 5 completed successfully!")
        print("Check the 'test_step_by_step/output' folder for results")
        print("\nGenerated files:")
        print("- merged_documentation_llm_fixed.docx: LLM-fixed DOCX")
    else:
        print("Step 5 failed!")
        sys.exit(1)

if __name__ == "__main__":
    main()
