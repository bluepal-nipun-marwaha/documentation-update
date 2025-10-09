#!/usr/bin/env python3
"""
Step 5B: Targeted Formatting Fix
Applies specific formatting fixes based on visual comparison between original and merged documents.
"""

import os
import sys
from pathlib import Path
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

class TargetedFormattingFixer:
    """Targeted formatting fixer based on visual comparison."""
    
    def __init__(self):
        """Initialize the formatting fixer."""
        self.input_dir = Path(__file__).parent / "input"
        self.output_dir = Path(__file__).parent / "output"
        
        # Ensure directories exist
        self.input_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)
    
    def run_step5b(self):
        """Run Step 5B: Targeted formatting fix."""
        try:
            print("Step 5B: Targeted Formatting Fix")
            print("=" * 80)
            print(f"Input Directory: {self.input_dir}")
            print(f"Output Directory: {self.output_dir}")
            print("Starting Step 5B: Targeted Formatting Fix")
            print("=" * 80)
            
            # Load the merged document
            merged_doc = self.load_merged_documentation()
            if not merged_doc:
                print("Failed to load merged document")
                return False
            
            # Apply targeted formatting fixes
            self.apply_targeted_fixes(merged_doc)
            
            # Save the fixed document
            self.save_fixed_document(merged_doc)
            
            print("=" * 80)
            print("Step 5B: Targeted Formatting Fix Completed Successfully!")
            print(f"Check the output directory: {self.output_dir}")
            print("=" * 80)
            
            return True
            
        except Exception as e:
            print(f"Step 5B failed: {str(e)}")
            return False
    
    def load_merged_documentation(self) -> Document:
        """Load the merged documentation."""
        try:
            print("Loading merged documentation...")
            merged_file = self.output_dir / "merged_documentation_fixed.docx"
            if not merged_file.exists():
                print("Merged documentation file not found")
                return None
            
            doc = Document(merged_file)
            print(f"Merged documentation loaded: {merged_file.name}")
            return doc
            
        except Exception as e:
            print(f"Error loading merged documentation: {str(e)}")
            return None
    
    def apply_targeted_fixes(self, doc: Document):
        """Apply targeted formatting fixes based on visual comparison."""
        try:
            print("Applying targeted formatting fixes...")
            
            # Fix 1: Document-level formatting
            self.fix_document_formatting(doc)
            
            # Fix 2: Title and header formatting
            self.fix_title_formatting(doc)
            
            # Fix 3: Paragraph formatting
            self.fix_paragraph_formatting(doc)
            
            # Fix 4: Table formatting
            self.fix_table_formatting(doc)
            
            # Fix 5: Spacing and alignment
            self.fix_spacing_alignment(doc)
            
            print("Targeted formatting fixes applied successfully")
            
        except Exception as e:
            print(f"Error applying targeted fixes: {str(e)}")
    
    def fix_document_formatting(self, doc: Document):
        """Fix document-level formatting."""
        try:
            print("  Fixing document-level formatting...")
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(10.5)
            font.color.rgb = RGBColor(255, 255, 255)  # White text
            
        except Exception as e:
            print(f"    Error fixing document formatting: {str(e)}")
    
    def fix_title_formatting(self, doc: Document):
        """Fix title and header formatting."""
        try:
            print("  Fixing title and header formatting...")
            
            # Fix main title
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                # Main title
                if "Click Program Documentation" in text and len(text) < 50:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if paragraph.runs:
                        run = paragraph.runs[0]
                        run.font.name = 'Calibri'
                        run.font.size = Pt(18)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White
                
                # Subtitle
                elif "Comprehensive Technical Analysis Report" in text:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if paragraph.runs:
                        run = paragraph.runs[0]
                        run.font.name = 'Calibri'
                        run.font.size = Pt(12)
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White
                
                # Version info
                elif "Version" in text and "Generated" in text:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if paragraph.runs:
                        run = paragraph.runs[0]
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White
                
                # Section headings
                elif any(heading in text for heading in [
                    "Executive Summary", "Table of Contents", "Program Overview",
                    "Architecture Analysis", "Module Structure", "Core Classes",
                    "Decorators", "Exception Handling", "Utility Functions",
                    "Dependencies", "Testing Framework", "Examples",
                    "Performance Analysis", "Future Roadmap", "Conclusion"
                ]):
                    if paragraph.runs:
                        run = paragraph.runs[0]
                        run.font.name = 'Calibri'
                        run.font.size = Pt(12)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0, 120, 215)  # Blue
                
                # Sub-headings
                elif any(subheading in text for subheading in [
                    "Basic Information", "Program Statistics", "Key Features"
                ]):
                    if paragraph.runs:
                        run = paragraph.runs[0]
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White
            
        except Exception as e:
            print(f"    Error fixing title formatting: {str(e)}")
    
    def fix_paragraph_formatting(self, doc: Document):
        """Fix paragraph formatting."""
        try:
            print("  Fixing paragraph formatting...")
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                # Skip empty paragraphs
                if not text:
                    continue
                
                # Regular paragraphs
                if not any(heading in text for heading in [
                    "Click Program Documentation", "Comprehensive Technical Analysis Report",
                    "Version", "Executive Summary", "Table of Contents", "Program Overview",
                    "Architecture Analysis", "Module Structure", "Core Classes",
                    "Decorators", "Exception Handling", "Utility Functions",
                    "Dependencies", "Testing Framework", "Examples",
                    "Performance Analysis", "Future Roadmap", "Conclusion",
                    "Basic Information", "Program Statistics", "Key Features"
                ]):
                    # Apply regular paragraph formatting
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragraph.paragraph_format.space_after = Pt(6)
                    
                    if paragraph.runs:
                        for run in paragraph.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(10.5)
                            run.font.color.rgb = RGBColor(255, 255, 255)  # White
            
        except Exception as e:
            print(f"    Error fixing paragraph formatting: {str(e)}")
    
    def fix_table_formatting(self, doc: Document):
        """Fix table formatting."""
        try:
            print("  Fixing table formatting...")
            
            for table in doc.tables:
                # Set table alignment
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Fix header row formatting
                if table.rows:
                    header_row = table.rows[0]
                    for cell in header_row.cells:
                        # Set background color for header cells
                        self.set_cell_background(cell, "34495E")  # Dark blue
                        
                        # Format header text
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if paragraph.runs:
                                for run in paragraph.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(10.5)
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)  # White
                
                # Fix data row formatting
                for row in table.rows[1:]:  # Skip header row
                    for cell in row.cells:
                        # Set background color for data cells
                        self.set_cell_background(cell, "2D2D2D")  # Dark gray
                        
                        # Format data text
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            if paragraph.runs:
                                for run in paragraph.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(10.5)
                                    run.font.bold = False
                                    run.font.color.rgb = RGBColor(255, 255, 255)  # White
                
                # Apply table borders
                self.apply_table_borders(table)
            
        except Exception as e:
            print(f"    Error fixing table formatting: {str(e)}")
    
    def set_cell_background(self, cell, color_hex: str):
        """Set background color for a table cell."""
        try:
            if color_hex.startswith('#'):
                color_hex = color_hex[1:]
            
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), color_hex)
            
            cell._tc.get_or_add_tcPr().append(shading)
            
        except Exception as e:
            print(f"      Error setting cell background: {str(e)}")
    
    def apply_table_borders(self, table):
        """Apply borders to table."""
        try:
            from docx.oxml import OxmlElement
            
            tbl = table._tbl
            tblBorders = OxmlElement('w:tblBorders')
            
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '808080')  # Gray
                tblBorders.append(border)
            
            tblPr = tbl.tblPr
            tblPr.append(tblBorders)
            
        except Exception as e:
            print(f"      Error applying table borders: {str(e)}")
    
    def fix_spacing_alignment(self, doc: Document):
        """Fix spacing and alignment issues."""
        try:
            print("  Fixing spacing and alignment...")
            
            # Add proper spacing between sections
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                
                # Add extra space before major headings
                if any(heading in text for heading in [
                    "Executive Summary", "Table of Contents", "Program Overview",
                    "Architecture Analysis", "Module Structure", "Core Classes",
                    "Decorators", "Exception Handling", "Utility Functions",
                    "Dependencies", "Testing Framework", "Examples",
                    "Performance Analysis", "Future Roadmap", "Conclusion"
                ]):
                    paragraph.paragraph_format.space_before = Pt(12)
                    paragraph.paragraph_format.space_after = Pt(6)
                
                # Add space before sub-headings
                elif any(subheading in text for subheading in [
                    "Basic Information", "Program Statistics", "Key Features"
                ]):
                    paragraph.paragraph_format.space_before = Pt(8)
                    paragraph.paragraph_format.space_after = Pt(4)
            
        except Exception as e:
            print(f"    Error fixing spacing: {str(e)}")
    
    def save_fixed_document(self, doc: Document):
        """Save the fixed document."""
        try:
            print("Saving fixed document...")
            
            output_path = self.output_dir / "merged_documentation_targeted_fixed.docx"
            doc.save(output_path)
            
            print(f"Fixed documentation saved successfully")
            print(f"Fixed DOCX: {output_path}")
            
        except Exception as e:
            print(f"Error saving fixed document: {str(e)}")
            raise

def main():
    """Main function to run Step 5B."""
    fixer = TargetedFormattingFixer()
    success = fixer.run_step5b()
    
    if success:
        print("Step 5B completed successfully!")
        print("Check the 'test_step_by_step/output' folder for results")
        print("\nGenerated files:")
        print("- merged_documentation_targeted_fixed.docx: Targeted formatting fixes applied")
    else:
        print("Step 5B failed!")
        sys.exit(1)

if __name__ == "__main__":
    main()

