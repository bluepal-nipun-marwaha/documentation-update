#!/usr/bin/env python3
"""
Step 4 Improved: Fresh DOCX Creation with Smart Section Matching
- Creates a completely new DOCX file
- Matches sections between original, updated text, and updated tables
- Recreates sections from original, then applies updates intelligently
- Handles both text and table sections properly
"""

import os
import sys
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple

# Set environment variables to use Ollama
os.environ['LLM_PROVIDER'] = 'ollama'
os.environ['LLM_MODEL'] = 'qwen2.5:7b'
os.environ['EMBEDDINGS_PROVIDER'] = 'nomic'
os.environ['EMBEDDINGS_MODEL'] = 'nomic-embed-text'
os.environ['OLLAMA_BASE_URL'] = 'http://localhost:11434'

# Add parent directory to path to import our modules
sys.path.append(str(Path(__file__).parent.parent))

from services.llm_service import LLMService
from utils.config import get_settings
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.table import Table
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from utils.docx_handler import DOCXHandler
import io
import re
import structlog

logger = structlog.get_logger(__name__)

def safe_print(text: str, prefix: str = ""):
    """Safely print text that might contain Unicode characters."""
    try:
        print(f"{prefix}{text}")
    except UnicodeEncodeError:
        safe_text = text.encode('ascii', 'replace').decode('ascii')
        print(f"{prefix}{safe_text}")

class SmartDocumentationCreator:
    """Creates fresh DOCX with smart section matching and recreation."""
    
    def __init__(self):
        self.test_dir = Path(__file__).parent
        self.input_dir = self.test_dir / "input"
        self.output_dir = self.test_dir / "output"
        
        # Create directories
        self.input_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)
        
        # Initialize LLM service
        settings = get_settings()
        self.llm_service = LLMService(settings.ai.model_dump())
        
        # Initialize DOCX handler
        self.docx_handler = DOCXHandler()
        
        print("Step 4 Improved: Fresh DOCX Creation with Smart Section Matching")
        print(f"Input Directory: {self.input_dir}")
        print(f"Output Directory: {self.output_dir}")
    
    def load_all_documentation(self) -> Tuple[Document, Document, List[Dict[str, Any]]]:
        """Load original, updated text, and updated tables."""
        try:
            print("Loading all documentation components...")
            
            # Load original documentation
            original_docx_path = self.input_dir / "Click_Professional_Documentation.docx"
            if not original_docx_path.exists():
                raise FileNotFoundError(f"Original documentation not found: {original_docx_path}")
            
            original_doc = Document(original_docx_path)
            safe_print(f"Original documentation loaded: {original_docx_path.name}")
            
            # Load updated text documentation
            updated_text_docx_path = self.output_dir / "updated_documentation.docx"
            if not updated_text_docx_path.exists():
                raise FileNotFoundError(f"Updated text documentation not found: {updated_text_docx_path}")
            
            updated_text_doc = Document(updated_text_docx_path)
            safe_print(f"Updated text documentation loaded: {updated_text_docx_path.name}")
            
            # Load updated tables
            updated_tables_json_path = self.output_dir / "updated_tables_improved.json"
            if not updated_tables_json_path.exists():
                raise FileNotFoundError(f"Updated tables not found: {updated_tables_json_path}")
            
            with open(updated_tables_json_path, 'r', encoding='utf-8') as f:
                updated_tables = json.load(f)
            
            safe_print(f"Updated tables loaded: {len(updated_tables)} tables")
            
            return original_doc, updated_text_doc, updated_tables
            
        except Exception as e:
            print(f"Error loading documentation: {str(e)}")
            return None, None, []
    
    def extract_document_structure(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract complete document structure with sections, paragraphs, and tables."""
        try:
            print("Extracting document structure...")
            
            structure = []
            current_section = None
            table_counter = 0
            
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    paragraph = None
                    for para in doc.paragraphs:
                        if para._element == element:
                            paragraph = para
                            break
                    
                    if paragraph:
                        text = paragraph.text.strip()
                        if not text:
                            continue
                        
                        # Check if this is a major heading (not sub-headings)
                        is_heading = self.is_major_heading(paragraph)
                        
                        if is_heading:
                            # Save previous section
                            if current_section:
                                structure.append(current_section)
                            
                            # Start new section
                            current_section = {
                                'type': 'section',
                                'title': text,
                                'paragraphs': [],
                                'tables': [],
                                'original_paragraph': paragraph
                            }
                        else:
                            # Add paragraph to current section
                            if current_section:
                                current_section['paragraphs'].append({
                                    'text': text,
                                    'original_paragraph': paragraph
                                })
                            else:
                                # Create a default section if none exists
                                current_section = {
                                    'type': 'section',
                                    'title': 'Introduction',
                                    'paragraphs': [{
                                        'text': text,
                                        'original_paragraph': paragraph
                                    }],
                                    'tables': []
                                }
                
                elif element.tag.endswith('tbl'):  # Table
                    table = None
                    for tbl in doc.tables:
                        if tbl._element == element:
                            table = tbl
                            break
                    
                    if table:
                        table_counter += 1
                        table_data = {
                            'type': 'table',
                            'index': table_counter - 1,
                            'rows': len(table.rows),
                            'cols': len(table.columns),
                            'data': [],
                            'original_table': table
                        }
                        
                        # Extract table data
                        for row in table.rows:
                            row_data = []
                            for cell in row.cells:
                                row_data.append(cell.text.strip())
                            table_data['data'].append(row_data)
                        
                        # Add table to current section or create new section
                        if current_section:
                            current_section['tables'].append(table_data)
                        else:
                            # Create a section for standalone table
                            structure.append({
                                'type': 'section',
                                'title': f'Table {table_counter}',
                                'paragraphs': [],
                                'tables': [table_data]
                            })
            
            # Add the last section
            if current_section:
                structure.append(current_section)
            
            print(f"Extracted structure with {len(structure)} sections")
            return structure
            
        except Exception as e:
            print(f"Error extracting document structure: {str(e)}")
            return []
    
    def is_major_heading(self, paragraph) -> bool:
        """Determine if a paragraph is a major heading (not sub-headings)."""
        try:
            text = paragraph.text.strip()
            
            # Skip very short text that might be sub-headings
            if len(text) < 3:
                return False
            
            # Check style - only Heading 1 and 2 are major
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name.replace('Heading ', '')
                if level.isdigit() and int(level) <= 2:
                    return True
            
            # Check formatting - only bold text that's not too long
            if paragraph.runs:
                first_run = paragraph.runs[0]
                if first_run.bold and len(text) < 80:
                    # Additional checks to avoid sub-headings
                    if not text.endswith(':') and not text.islower():
                        return True
            
            # Check content patterns - only major sections
            major_patterns = [
                'Program Overview', 'Architecture', 'Module Structure', 
                'Core Classes', 'Decorators', 'Exception Handling',
                'Utility Functions', 'Dependencies', 'Testing Framework',
                'Examples', 'Performance Analysis', 'Future Roadmap',
                'Conclusion', 'Executive Summary', 'Table of Contents'
            ]
            
            for pattern in major_patterns:
                if pattern.lower() in text.lower():
                    return True
            
            # Check for numbered sections (1., 2., etc.)
            if text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.', '11.', '12.', '13.')):
                return True
            
            return False
            
        except Exception as e:
            return False
    
    
    def match_sections(self, original_structure: List[Dict[str, Any]], 
                      updated_text_doc: Document, 
                      updated_tables: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Match sections between original, updated text, and updated tables."""
        try:
            print("Matching sections across all documentation...")
            
            # Extract updated text structure
            updated_text_structure = self.extract_document_structure(updated_text_doc)
            
            matched_sections = []
            processed_titles = set()  # Track processed section titles to avoid duplicates
            
            for orig_section in original_structure:
                # Skip if we've already processed this section title
                if orig_section['title'] in processed_titles:
                    print(f"  Skipping duplicate section: '{orig_section['title']}'")
                    continue
                
                section_info = {
                    'original': orig_section,
                    'updated_text': None,
                    'updated_tables': [],
                    'has_text_changes': False,
                    'has_table_changes': False
                }
                
                # Find matching section in updated text
                for updated_section in updated_text_structure:
                    if self.sections_match(orig_section['title'], updated_section['title']):
                        section_info['updated_text'] = updated_section
                        section_info['has_text_changes'] = self.has_text_changes(
                            orig_section, updated_section
                        )
                        break
                
                # Find matching tables by column structure and content
                if orig_section['tables']:
                    section_info['has_table_changes'] = True
                    for orig_table in orig_section['tables']:
                        # Find matching table by columns and content
                        matching_table = self.find_matching_table(orig_table, updated_tables)
                        if matching_table:
                            section_info['updated_tables'].append(matching_table)
                
                matched_sections.append(section_info)
                processed_titles.add(orig_section['title'])
            
            print(f"Matched {len(matched_sections)} unique sections")
            return matched_sections
            
        except Exception as e:
            print(f"Error matching sections: {str(e)}")
            return []
    
    def sections_match(self, title1: str, title2: str) -> bool:
        """Check if two section titles match."""
        try:
            # Normalize titles
            norm1 = title1.lower().strip()
            norm2 = title2.lower().strip()
            
            # Exact match
            if norm1 == norm2:
                return True
            
            # Partial match (one contains the other)
            if norm1 in norm2 or norm2 in norm1:
                return True
            
            # Check for common variations
            variations = [
                ('table of contents', 'contents'),
                ('key features', 'features'),
                ('example applications', 'examples'),
                ('performance optimizations', 'performance'),
            ]
            
            for var1, var2 in variations:
                if (var1 in norm1 and var2 in norm2) or (var2 in norm1 and var1 in norm2):
                    return True
            
            return False
            
        except Exception as e:
            return False
    
    def has_text_changes(self, original_section: Dict[str, Any], 
                        updated_section: Dict[str, Any]) -> bool:
        """Check if there are text changes between sections."""
        try:
            orig_text = ' '.join([p['text'] for p in original_section['paragraphs']])
            updated_text = ' '.join([p['text'] for p in updated_section['paragraphs']])
            
            return orig_text != updated_text
            
        except Exception as e:
            return False
    
    def find_matching_table(self, original_table: Dict[str, Any], 
                           updated_tables: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
        """Find matching table by column structure and content."""
        try:
            orig_cols = original_table['cols']
            orig_rows = original_table['rows']
            
            # Get original table headers (first row)
            orig_headers = []
            if original_table['data'] and len(original_table['data']) > 0:
                orig_headers = original_table['data'][0]
            
            print(f"    Looking for table match: {orig_cols} cols, {orig_rows} rows")
            print(f"    Original headers: {orig_headers}")
            
            best_match = None
            best_score = 0
            
            for updated_table in updated_tables:
                updated_cols = updated_table.get('columns', 0)
                updated_rows = updated_table.get('rows', 0)
                
                # Get updated table headers
                updated_headers = []
                if 'rows_data' in updated_table and len(updated_table['rows_data']) > 0:
                    first_row = updated_table['rows_data'][0]
                    if 'cells' in first_row:
                        updated_headers = [cell.get('content', '') for cell in first_row['cells']]
                
                print(f"      Checking table: {updated_cols} cols, {updated_rows} rows")
                print(f"      Updated headers: {updated_headers}")
                
                # Calculate match score
                score = self.calculate_table_match_score(
                    orig_cols, orig_rows, orig_headers,
                    updated_cols, updated_rows, updated_headers
                )
                
                print(f"      Match score: {score}")
                
                if score > best_score:
                    best_score = score
                    best_match = updated_table
            
            if best_match and best_score > 0.3:  # Minimum threshold for match
                print(f"    Found best match with score: {best_score}")
                return best_match
            else:
                print(f"    No suitable match found (best score: {best_score})")
                return None
                
        except Exception as e:
            print(f"Error finding matching table: {str(e)}")
            return None
    
    def calculate_table_match_score(self, orig_cols: int, orig_rows: int, orig_headers: List[str],
                                  updated_cols: int, updated_rows: int, updated_headers: List[str]) -> float:
        """Calculate match score between original and updated table."""
        try:
            score = 0.0
            
            # Column count match (40% weight)
            if orig_cols == updated_cols:
                score += 0.4
            elif abs(orig_cols - updated_cols) <= 1:
                score += 0.2
            
            # Row count similarity (20% weight)
            if orig_rows == updated_rows:
                score += 0.2
            elif abs(orig_rows - updated_rows) <= 2:
                score += 0.1
            
            # Header content match (40% weight)
            if orig_headers and updated_headers:
                header_matches = 0
                for orig_header in orig_headers:
                    orig_header_lower = orig_header.lower().strip()
                    for updated_header in updated_headers:
                        updated_header_lower = updated_header.lower().strip()
                        
                        # Exact match
                        if orig_header_lower == updated_header_lower:
                            header_matches += 1
                            break
                        # Partial match
                        elif (orig_header_lower in updated_header_lower or 
                              updated_header_lower in orig_header_lower):
                            header_matches += 0.5
                            break
                
                if len(orig_headers) > 0:
                    header_score = header_matches / len(orig_headers)
                    score += header_score * 0.4
            
            return score
            
        except Exception as e:
            print(f"Error calculating table match score: {str(e)}")
            return 0.0
    
    def create_fresh_document(self, matched_sections: List[Dict[str, Any]]) -> Document:
        """Create a fresh DOCX document with smart section recreation."""
        try:
            print("Creating fresh DOCX document...")
            
            # Create new document
            new_doc = Document()
            
            # Initialize duplicate tracking
            self._last_section_title = None
            
            # Add title
            title = new_doc.add_heading('Updated Click Documentation', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Process each section
            for section_info in matched_sections:
                self.recreate_section(new_doc, section_info)
            
            print("Fresh DOCX document created successfully")
            return new_doc
            
        except Exception as e:
            print(f"Error creating fresh document: {str(e)}")
            return Document()
    
    def recreate_section(self, doc: Document, section_info: Dict[str, Any]):
        """Recreate a section with smart updates - only add content once."""
        try:
            original_section = section_info['original']
            updated_text = section_info['updated_text']
            updated_tables = section_info['updated_tables']
            
            # Skip if this section is a duplicate (same title as previous)
            if hasattr(self, '_last_section_title') and self._last_section_title == original_section['title']:
                print(f"  Skipping duplicate section '{original_section['title']}'")
                return
            
            # Add section title only if it's not empty and not a duplicate
            if original_section['title'].strip():
                title_para = doc.add_heading(original_section['title'], level=1)
                self._last_section_title = original_section['title']
            
            # Choose the best content version - only add ONE version
            if section_info['has_text_changes'] and updated_text and updated_text['paragraphs']:
                # Use updated text if available and has content
                print(f"  Adding updated content for section '{original_section['title']}'")
                for para_info in updated_text['paragraphs']:
                    if para_info['text'].strip():  # Only add non-empty paragraphs
                        new_para = doc.add_paragraph()
                        self.copy_paragraph_formatting(para_info['original_paragraph'], new_para)
                        new_para.text = para_info['text']
            elif original_section['paragraphs']:
                # Use original text if no updates or updates are empty
                print(f"  Adding original content for section '{original_section['title']}'")
                for para_info in original_section['paragraphs']:
                    if para_info['text'].strip():  # Only add non-empty paragraphs
                        new_para = doc.add_paragraph()
                        self.copy_paragraph_formatting(para_info['original_paragraph'], new_para)
                        new_para.text = para_info['text']
            
            # Add tables - only add ONE version
            if section_info['has_table_changes'] and updated_tables:
                print(f"  Adding {len(updated_tables)} updated tables to section")
                for table_data in updated_tables:
                    self.add_updated_table(doc, table_data)
            elif original_section['tables']:
                print(f"  Adding {len(original_section['tables'])} original tables to section")
                for table_info in original_section['tables']:
                    self.add_original_table(doc, table_info)
            
        except Exception as e:
            print(f"Error recreating section: {str(e)}")
    
    def copy_paragraph_formatting(self, original_para, new_para):
        """Copy formatting from original paragraph to new paragraph."""
        try:
            # Copy style
            new_para.style = original_para.style
            
            # Copy run formatting
            for run in original_para.runs:
                new_run = new_para.add_run()
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                
        except Exception as e:
            print(f"Error copying paragraph formatting: {str(e)}")
    
    def add_updated_table(self, doc: Document, table_data: Dict[str, Any]):
        """Add an updated table to the document."""
        try:
            # Create table
            table = doc.add_table(rows=table_data['rows'], cols=table_data['columns'])
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Apply borders
            self.apply_table_borders(table)
            
            # Populate with updated data
            for row_idx, row_data in enumerate(table_data['rows_data']):
                if row_idx < len(table.rows):
                    table_row = table.rows[row_idx]
                    for cell_idx, cell_data in enumerate(row_data['cells']):
                        if cell_idx < len(table_row.cells):
                            cell = table_row.cells[cell_idx]
                            cell.text = cell_data['content']
                            self.apply_cell_formatting(cell, cell_data)
            
        except Exception as e:
            print(f"Error adding updated table: {str(e)}")
    
    def add_original_table(self, doc: Document, table_info: Dict[str, Any]):
        """Add an original table to the document."""
        try:
            # Create table
            table = doc.add_table(rows=table_info['rows'], cols=table_info['cols'])
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Apply borders
            self.apply_table_borders(table)
            
            # Populate with original data
            for row_idx, row_data in enumerate(table_info['data']):
                if row_idx < len(table.rows):
                    table_row = table.rows[row_idx]
                    for cell_idx, cell_text in enumerate(row_data):
                        if cell_idx < len(table_row.cells):
                            cell = table_row.cells[cell_idx]
                            cell.text = cell_text
            
        except Exception as e:
            print(f"Error adding original table: {str(e)}")
    
    def apply_table_borders(self, table: Table):
        """Apply proper borders to a table."""
        try:
            # Define border style
            border_xml = '''
            <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            </w:tcBorders>
            '''
            
            # Apply borders to all cells
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcPr.append(parse_xml(border_xml))
                    
        except Exception as e:
            print(f"Error applying table borders: {str(e)}")
    
    def apply_cell_formatting(self, cell, cell_data: Dict[str, Any]):
        """Apply formatting to a cell."""
        try:
            if 'formatting' in cell_data and 'font' in cell_data['formatting']:
                font_info = cell_data['formatting']['font']
                
                if cell.paragraphs:
                    para = cell.paragraphs[0]
                    if para.runs:
                        run = para.runs[0]
                        
                        if 'name' in font_info:
                            run.font.name = font_info['name']
                        if 'size' in font_info:
                            run.font.size = Pt(font_info['size'])
                        if 'bold' in font_info:
                            run.bold = font_info['bold']
                        if 'italic' in font_info:
                            run.italic = font_info['italic']
                            
        except Exception as e:
            print(f"Error applying cell formatting: {str(e)}")
    
    def save_fresh_documentation(self, fresh_doc: Document):
        """Save the fresh documentation."""
        try:
            # Save the fresh DOCX
            fresh_docx_path = self.output_dir / "fresh_documentation.docx"
            fresh_doc.save(fresh_docx_path)
            
            print(f"Fresh documentation saved successfully")
            print(f"Fresh DOCX: {fresh_docx_path}")
            
            return {
                'fresh_docx_path': fresh_docx_path
            }
            
        except Exception as e:
            print(f"Error saving fresh documentation: {str(e)}")
            return {}
    
    def run_smart_creation(self):
        """Run the smart documentation creation process."""
        try:
            print("Starting Step 4 Improved: Fresh DOCX Creation with Smart Section Matching")
            print("=" * 80)
            
            # Step 1: Load all documentation components
            original_doc, updated_text_doc, updated_tables = self.load_all_documentation()
            if not original_doc or not updated_text_doc or not updated_tables:
                print("Failed to load all documentation components")
                return False
            
            # Step 2: Extract document structures
            original_structure = self.extract_document_structure(original_doc)
            
            # Step 3: Match sections across all documentation
            matched_sections = self.match_sections(original_structure, updated_text_doc, updated_tables)
            
            # Step 4: Create fresh document with smart recreation
            fresh_doc = self.create_fresh_document(matched_sections)
            
            # Step 5: Save fresh documentation
            results = self.save_fresh_documentation(fresh_doc)
            
            print("=" * 80)
            print("Step 4 Improved: Fresh DOCX Creation Completed Successfully!")
            print(f"Check the output directory: {self.output_dir}")
            
            return True
            
        except Exception as e:
            print(f"Step 4 Improved failed: {str(e)}")
            return False

def main():
    """Main function to run Step 4 Improved."""
    print("Step 4 Improved: Fresh DOCX Creation with Smart Section Matching")
    print("=" * 80)
    
    # Initialize creator
    creator = SmartDocumentationCreator()
    
    # Run smart creation
    success = creator.run_smart_creation()
    
    if success:
        print("\nStep 4 Improved completed successfully!")
        print("Check the 'test_step_by_step/output' folder for results")
        print("\nGenerated files:")
        print("- fresh_documentation.docx: Fresh DOCX with smart section matching")
    else:
        print("\nStep 4 Improved failed!")
        print("Make sure Steps 1-3 have been completed first")

if __name__ == "__main__":
    main()
